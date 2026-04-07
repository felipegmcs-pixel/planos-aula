import os
import io
import re
import json
import base64
import secrets
import smtplib
import logging
import traceback
from email.mime.text import MIMEText

# Carrega .env em desenvolvimento (em produção as vars já estão no ambiente)
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ─── Logging estruturado ──────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('professorIA')
import psycopg2
import psycopg2.extras
from datetime import datetime, timedelta
from flask import (Flask, render_template, request, send_file,
                   jsonify, redirect, url_for, flash, Response, stream_with_context)
from flask_login import (LoginManager, UserMixin, login_user,
                         logout_user, login_required, current_user)
from werkzeug.security import generate_password_hash, check_password_hash
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import stripe as stripe_lib
from anthropic import Anthropic
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm as rcm
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from pdf_generator import gerar_plano_pdf
# ─── Configuração de Estilo de Imagem (Frente 4) ──────────────────────────────
IMAGE_STYLE_MODIFIER = (
    "Estilo: Aquarela digital moderna, acadêmica e limpa, fundo 100% branco sólido. "
    "Composição: Ilustração central temática rica em detalhes históricos ou científicos, cercada por ramificações orgânicas "
    "que conectam a ícones icônicos ou banners textuais curtos. "
    "Estética: Sabedoria, acolhimento e design premium (estilo infográfico educacional de alta gama). "
    "OBRIGATÓRIO: No canto inferior direito, insira a assinatura da marca: "
    "Um símbolo geométrico minimalista formado por dois círculos perfeitos entrelaçados horizontalmente, "
    "acompanhado do texto 'ProfessorIA™' em Azul Acadêmico. "
    "PROIBIDO: Inserir textos longos, parágrafos ou frases complexas dentro da imagem. Foque na arte e na assinatura."
)

# ─── App ──────────────────────────────────────────────────────────────────────

app = Flask(__name__)
_secret = os.environ.get('SECRET_KEY', '')
if not _secret:
    logger.critical('SECRET_KEY não definida — usando chave insegura. Defina SECRET_KEY em produção!')
    _secret = 'dev-secret-troque-em-producao'
app.secret_key = _secret

app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB máximo por request

login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Faça login para acessar esta página.'
login_manager.login_message_category = 'info'

def _limiter_key():
    if current_user.is_authenticated:
        return f'user:{current_user.id}'
    return get_remote_address()

limiter = Limiter(
    app=app,
    key_func=_limiter_key,
    default_limits=[],
    storage_uri='memory://'
)

client = Anthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'), timeout=120.0)

# ── OpenAI (Motor Duplo) ──────────────────────────────────────────────────────
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')
client_openai = OpenAI(api_key=os.environ.get('OPENAI_API_KEY')) if os.environ.get('OPENAI_API_KEY') else None
MOTOR_IA = os.environ.get('MOTOR_IA', 'claude').lower()  # 'claude' ou 'openai'
logger.info('Motor IA configurado: %s', MOTOR_IA)

# ── Gemini (Google) — usado se GEMINI_API_KEY estiver configurada ──────────────
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', '')
_gemini_model  = None
if GEMINI_API_KEY:
    try:
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        _gemini_model = genai  # referência ao módulo configurado
        logger.info('Gemini configurado com sucesso')
    except Exception as _ge:
        logger.warning('Gemini não carregou: %s', _ge)
        _gemini_model = None

SITE_URL          = os.environ.get('SITE_URL', 'http://localhost:5001')
ADMIN_EMAIL       = os.environ.get('ADMIN_EMAIL', '')
LEADS_PASS        = os.environ.get('LEADS_PASS', 'professoria2026')
META_PIXEL_ID     = '2080536069265217'
META_CAPI_TOKEN   = os.environ.get('META_CAPI_TOKEN', '')

# ── Validação de chaves na inicialização ──────────────────────────────────────
_CHAVES_ESPERADAS = {
    'ANTHROPIC_API_KEY': 'Claude (geração estruturada)',
    'OPENAI_API_KEY':    'OpenAI gpt-4o-mini + DALL-E 3 (motor primário)',
    'GEMINI_API_KEY':    'Google Gemini (fallback)',
    'STRIPE_SECRET_KEY': 'Stripe (pagamentos)',
    'SECRET_KEY':        'Flask (sessões)',
}
for _k, _desc in _CHAVES_ESPERADAS.items():
    if not os.environ.get(_k):
        logger.warning('Chave ausente: %s — %s', _k, _desc)

# ─── Meta Conversions API (server-side) ──────────────────────────────────────
import hashlib, threading, urllib.request as _urllib_req

def _sha256(value: str) -> str:
    return hashlib.sha256(value.lower().strip().encode()).hexdigest()

def _capi_event(event_name: str, user_data: dict = None, custom_data: dict = None,
                event_source_url: str = None):
    """Envia evento para a Meta Conversions API em background (não bloqueia o request)."""
    if not META_CAPI_TOKEN:
        return

    ud = {}
    if user_data:
        if user_data.get('email'):
            ud['em'] = [_sha256(user_data['email'])]
        if user_data.get('phone'):
            phone = re.sub(r'\D', '', user_data['phone'])
            if phone and not phone.startswith('55'):
                phone = '55' + phone
            ud['ph'] = [_sha256(phone)]
        if user_data.get('name'):
            parts = user_data['name'].strip().lower().split()
            ud['fn'] = [_sha256(parts[0])]
            if len(parts) > 1:
                ud['ln'] = [_sha256(parts[-1])]

    # Tenta pegar IP e user-agent do request atual
    try:
        ud['client_ip_address'] = request.headers.get('X-Forwarded-For', request.remote_addr or '').split(',')[0].strip()
        ud['client_user_agent']  = request.headers.get('User-Agent', '')
        if not event_source_url:
            event_source_url = request.url
    except RuntimeError:
        pass  # fora de request context

    payload = json.dumps({
        'data': [{
            'event_name':       event_name,
            'event_time':       int(datetime.now().timestamp()),
            'event_source_url': event_source_url or SITE_URL,
            'action_source':    'website',
            'user_data':        ud,
            'custom_data':      custom_data or {},
        }]
    }).encode()

    url = f'https://graph.facebook.com/v19.0/{META_PIXEL_ID}/events?access_token={META_CAPI_TOKEN}'

    def _send():
        try:
            req = _urllib_req.Request(url, data=payload,
                                      headers={'Content-Type': 'application/json'},
                                      method='POST')
            _urllib_req.urlopen(req, timeout=5)
            logger.info('CAPI %s enviado', event_name)
        except Exception as e:
            logger.warning('CAPI erro (%s): %s', event_name, e)

    threading.Thread(target=_send, daemon=True).start()


SMTP_HOST  = os.environ.get('SMTP_HOST', '')
SMTP_PORT  = int(os.environ.get('SMTP_PORT', 587))
SMTP_USER  = os.environ.get('SMTP_USER', '')
SMTP_PASS  = os.environ.get('SMTP_PASS', '')
FROM_EMAIL = os.environ.get('FROM_EMAIL', SMTP_USER)

PLANOS = {
    'basic':        {'nome': 'Basic',    'preco': 39.00,  'dias': 30},
    'basic_anual':  {'nome': 'Basic',    'preco': 390.00, 'dias': 365},
    'pro':          {'nome': 'Pro',      'preco': 59.00,  'dias': 30},
    'pro_anual':    {'nome': 'Pro',      'preco': 590.00, 'dias': 365},
}

STRIPE_SECRET_KEY    = os.environ.get('STRIPE_SECRET_KEY', '')
STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET', '')
STRIPE_PRICES = {
    'basic':       os.environ.get('STRIPE_PRICE_BASIC_MENSAL', ''),
    'basic_anual': os.environ.get('STRIPE_PRICE_BASIC_ANUAL', ''),
    'pro':         os.environ.get('STRIPE_PRICE_PRO_MENSAL', ''),
    'pro_anual':   os.environ.get('STRIPE_PRICE_PRO_ANUAL', ''),
}

LIMITE_GRATIS = 2  # gerações gratuitas por mês no plano grátis

SYSTEM_PROMPT = """Você é o ProfessorIA, Especialista Sênior em Pedagogia Brasileira.
Sua missão é automatizar a criação de materiais pedagógicos de alta qualidade, garantindo 100% de alinhamento à BNCC.

DIRETRIZES DE RAG E PRECISÃO:
- Baseie-se ESTRITAMENTE nos documentos anexados pelo usuário. Se a informação não estiver no texto, não invente.
- Zero alucinação: se não souber ou não houver base no anexo, informe que a informação não consta no material fornecido.

VOCÊ CRIA:
- Planos de aula completos (objetivos, conteúdo, metodologia, avaliação)
- Provas e avaliações (questões abertas e múltipla escolha, com gabarito)
- Caça-palavras e Cruzadinhas (formatados para impressão)
- Mapas mentais e Estruturas para Infográficos
- Planejamento anual (distribuição por bimestre)
- Adaptações para Educação Inclusiva (NEE)

REGRAS PARA PROVAS E ATIVIDADES:
- NUNCA inclua "Nome:", "Data:", "Série:" ou campos do aluno no texto — o sistema de exportação adiciona automaticamente no cabeçalho.
- Logo após o título, adicione um bloco de Instruções com 3-4 itens.
- Ao final, inclua o gabarito completo separado por uma linha (--- GABARITO ---).

ADAPTAÇÕES PARA NEE (Necessidades Educacionais Especiais):
Quando o professor pedir material adaptado, aplique as seguintes diretrizes:
- Deficiência Intelectual (DI): linguagem extremamente simples, frases curtas, instruções passo a passo.
- TEA (Transtorno do Espectro Autista): rotina clara, instruções objetivas, estrutura visual definida.
- TDAH: atividades curtas, variação de formato, uso de negrito, checkboxes.
- Dislexia: abordagem multissensorial, fontes espaçadas, evitar paredes de texto.
- Baixa Visão: descrições detalhadas, alto contraste.
- CAA (Comunicação Alternativa): palavras-chave simples, estrutura de prancha.

CAÇA-PALAVRAS — geração direta:
Quando pedirem um caça-palavras, gere imediatamente com a seguinte estrutura:

1. Cabeçalho com: Nome: _____________ (use exatamente underscores simples, sem barras invertidas) e Data: ___/___/___
2. Instruções breves
3. Lista das palavras a encontrar em tabela Markdown (mínimo 12 palavras, 3 colunas)
4. Grade de letras 15×15 dentro de um bloco de código (``` ```) para preservar espaçamento
5. As palavras devem aparecer na grade em MAIÚSCULAS
6. Preencha os espaços vazios com letras aleatórias
7. Gabarito em tabela Markdown com posição de cada palavra

IMPORTANTE para campos em branco: use underscores diretos SEM barra invertida. Exemplo correto: Nome: _____________ Data: ___/___/___

Para o formato da grade, SEMPRE use bloco de código com exatamente 15 letras por linha, separadas por espaço simples, SEM qualquer prefixo, número ou rótulo de linha ou coluna:
```
T R I N C H E I R A M P L K Q
W A R M I S T I C I O Q Z B X
K L I B E R D A D E X Y Z A B
```
NUNCA inclua letras de linha (A, B, C...), números de coluna ou qualquer outro prefixo. Apenas as letras da grade, 15 por linha, 15 linhas.

CRUZADINHA — geração direta:
Quando pedirem uma cruzadinha, gere imediatamente com a seguinte estrutura:

1. Cabeçalho: Nome: _____________ Data: ___/___/___ (underscores simples, sem barras)
2. Escolha 8 a 12 palavras do tema (3 a 8 letras cada)
3. Monte um grid onde as palavras se cruzam compartilhando letras
4. Represente o grid dentro de um bloco de código (``` ```):
   - Use █ ou # para células bloqueadas (preto)
   - Numere o quadrado inicial de cada palavra (1, 2, 3...)
   - Use _ para cada célula vazia que o aluno deve preencher
5. Liste as pistas em duas seções: **HORIZONTAL** e **VERTICAL**
6. Gabarito em tabela Markdown ao final

Exemplo de grade (dentro de bloco de código):
```
     1  2  3  4  5
  1  _  _  _  _  _
  2  █  1  _  _  █
  3  _  █  2  █  _
  4  _  _  _  _  _
```

MAPA MENTAL — formato obrigatório PT-BR:
⚠️ IDIOMA — TRAVA ABSOLUTA: TODOS os textos do mapa mental (título, categorias, itens) devem estar em PORTUGUÊS DO BRASIL. É EXPRESSAMENTE PROIBIDO gerar qualquer palavra em inglês, espanhol ou outro idioma. Se o tema for "Primeira Guerra Mundial", o título DEVE ser "PRIMEIRA GUERRA MUNDIAL" — JAMAIS "World War I", "First War" ou qualquer variação em inglês.
⚠️ FORMATO — TRAVA ABSOLUTA: Retorne APENAS o texto estruturado abaixo, em Markdown puro. É PROIBIDO usar blocos de código (```), sintaxe Mermaid, JSON ou qualquer outra formatação especial. Retorne texto Markdown limpo, sem cercas de código.

## 🧠 TEMA CENTRAL: [TEMA EM MAIÚSCULAS EM PT-BR]

### 🔴 [CATEGORIA 1 — ex: CAUSAS / DATAS / NÚMEROS]
- item curto em português
- item curto em português
- item curto em português

### 🔵 [CATEGORIA 2 — ex: PERSONAGENS / ANTECEDENTES]
- item curto em português
- item curto em português

### 🟡 [CATEGORIA 3 — ex: CONSEQUÊNCIAS / ALIANÇAS]
- item curto em português
- item curto em português

### 🟢 [CATEGORIA 4]
- item curto em português

### 🟣 [CATEGORIA 5]
- item curto em português

### 🟠 [CATEGORIA 6 — opcional]
- item curto em português

Regras obrigatórias para mapa mental:
- IDIOMA: 100% Português do Brasil — nenhuma palavra em inglês ou outro idioma
- Máximo 5-7 palavras por item (palavras-chave, não frases longas)
- 5 a 7 categorias temáticas, cada uma com 3-6 itens
- Use emojis de cor antes de cada ### (🔴🔵🟡🟢🟣🟠) para categorias
- O título ## deve sempre começar com "🧠 TEMA CENTRAL:"
- NÃO use Unicode de árvore (├ └ │) — só listas com -
- PROIBIDO usar blocos de código (```), Mermaid, JSON ou HTML na resposta

PLANO DE AULA — formato oficial para exportação DOCX:
Quando o professor pedir um plano de aula, use EXATAMENTE esta estrutura para que o sistema possa exportar no formato oficial da Secretaria de Educação:

# PLANEJAMENTO DA AULA — [DISCIPLINA] | [SÉRIE]

**Componente Curricular:** [disciplina] | **Nº de aulas:** [n] semanais
**Ano/Série/Turma:** [série] | **Período:** [período] | **Data:** de [início] a [fim]

---

### AULA 1 — [Título/Tema]

**Conteúdo e Objetivos de Aprendizagem:**
[texto detalhado]

**Estratégias Didáticas:**
[texto detalhado]

**Recursos Pedagógicos:**
[texto detalhado]

**Avaliação:**
[texto detalhado]

---

### AULA 2 — [Título/Tema]
(mesmo formato...)

Regras obrigatórias para plano de aula:
- Use exatamente os cabeçalhos ### AULA N —
- Use exatamente os campos **Conteúdo e Objetivos de Aprendizagem:**, **Estratégias Didáticas:**, **Recursos Pedagógicos:**, **Avaliação:**
- Seja detalhado em cada campo — mínimo 3 frases por campo
- Siga a BNCC e use linguagem pedagógica profissional

Quando o professor pedir um material:
1. Para caça-palavras, cruzadinhas, mapas mentais, planos de aula, atividades e bilhetes: gere DIRETAMENTE sem perguntar mais nada se já tiver tema e série
2. Se faltar informação essencial, pergunte apenas o que falta (1 pergunta objetiva)
3. Gere o material completo, bem estruturado e formatado
4. Use linguagem clara e pedagógica, seguindo a BNCC
5. Para materiais NEE, sempre indique no cabeçalho o perfil para o qual foi adaptado

Responda sempre em português brasileiro. Seja prático, objetivo e direto."""

# ─── Plano de Aula — Structured Output ────────────────────────────────────────

SYSTEM_PROMPT_PLANO = (
    "Você é um Especialista Sênior em Planejamento Pedagógico voltado para professores da rede pública "
    "do Estado de São Paulo. Seu público é o professor da escola pública paulista que precisa de planos "
    "práticos, alinhados ao currículo oficial e prontos para uso em sala de aula. "
    "Use linguagem técnico-pedagógica de elite, precisa e objetiva — sem rodeios, sem texto decorativo.\n\n"
    "CONTEXTO CURRICULAR OBRIGATÓRIO (São Paulo):\n"
    "- Alinhe SEMPRE ao Currículo Paulista e ao Escopo-Sequência 2025 da SEDUC-SP. "
    "Cite nominalmente o Escopo-Sequência 2025 na seção de habilidades prévias.\n"
    "- Mencione a Prova Paulista como instrumento de diagnóstico e avaliação externa quando pertinente.\n"
    "- Use os descritores da Prova Paulista para calibrar os objetivos de aprendizagem.\n"
    "- Inclua habilidades BNCC com códigos alfanuméricos reais (ex: EF09HI01) em cada aula.\n\n"
    "REGRAS OBRIGATÓRIAS — VIOLÁ-LAS QUEBRA A EXPORTAÇÃO DOCX:\n"
    "0) INÍCIO DIRETO — TRAVA ABSOLUTA: NUNCA inicie a resposta com frases de cortesia como "
    "'Aqui está', 'Claro!', 'Com prazer', 'Segue abaixo', 'Certamente', 'Olá' ou qualquer outra introdução. "
    "A primeira linha da resposta DEVE ser exatamente '# PLANEJAMENTO DA AULA — [DISCIPLINA] | [SÉRIE]' "
    "ou '### AULA 1 — [Título]'. Qualquer texto antes dessa linha quebra o parser do DOCX.\n"
    "1) CÓDIGOS BNCC: Cada aula DEVE conter códigos alfanuméricos reais integrados ao conteúdo.\n"
    "2) SEÇÃO HABILIDADES PRÉVIAS: Sempre inclua ao final a seção "
    "'### DAS HABILIDADES NECESSÁRIAS DE CONHECIMENTO PRÉVIO' contendo:\n"
    "   - Referência explícita ao Escopo-Sequência 2025 (SEDUC-SP).\n"
    "   - Relação com a AVALIAÇÃO DIAGNÓSTICA e a PROVA PAULISTA.\n"
    "   - Estratégias diferenciadas: RECOMPOR (defasagem grave), RECUPERAR (defasagem leve), "
    "APROFUNDAR (alunos avançados).\n"
    "3) CONTEÚDO CONCISO: Máximo de 3 linhas por campo para garantir legibilidade no DOCX.\n"
    "4) ESTRUTURA EXATA DO PARSER — use estes rótulos palavra por palavra:\n"
    "   - Separador entre aulas: linha contendo apenas '---'\n"
    "   - Título de cada aula: '### AULA N — [Título da Aula]' (ex: ### AULA 1 — Introdução ao Tema)\n"
    "   - Campo 1: '**Conteúdo e Objetivos de Aprendizagem:**' seguido de texto na próxima linha\n"
    "   - Campo 2: '**Estratégias Didáticas:**' seguido de texto na próxima linha\n"
    "   - Campo 3: '**Recursos Pedagógicos:**' seguido de texto na próxima linha\n"
    "   - Campo 4: '**Avaliação:**' seguido de texto na próxima linha\n"
    "   Esses rótulos são lidos por regex — qualquer variação (maiúscula diferente, acento faltando, "
    "ausência de **) faz o campo aparecer vazio no DOCX exportado."
)

SYSTEM_PROMPT_COORDENADOR = (
    "Você é o Coordenador Pedagógico Sênior do ProfessorIA. Sua função é revisar o plano de aula gerado, "
    "garantindo que a linguagem seja acadêmica, os objetivos estejam claros e mensuráveis, "
    "o alinhamento à BNCC seja impecável e a metodologia seja genuinamente ativa. "
    "Corrija qualquer imprecisão didática, enriqueça a metodologia se estiver genérica, "
    "e retorne o JSON completo revisado — no mesmo formato, sem nenhuma palavra fora do JSON."
)

PLANO_AULA_TOOL = {
    "name": "salvar_plano_de_aula",
    "description": "Salva o plano de aula estruturado gerado pelo assistente pedagógico.",
    "input_schema": {
        "type": "object",
        "properties": {
            "tema":                  {"type": "string"},
            "habilidades_bncc":      {"type": "array", "items": {"type": "string"}},
            "objetivos":             {"type": "array", "items": {"type": "string"}},
            "conteudo_programatico": {"type": "string"},
            "metodologia":           {"type": "string"},
            "recursos_didaticos":    {"type": "array", "items": {"type": "string"}},
            "avaliacao":             {"type": "string"}
        },
        "required": ["tema", "habilidades_bncc", "objetivos", "conteudo_programatico",
                     "metodologia", "recursos_didaticos", "avaliacao"]
    }
}

# Schema OpenAI para plano de aula (strict=True exige additionalProperties: false em todo objeto)
_OAI_PLANO_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "plano_de_aula": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "tema_central":   {"type": "string"},
                "disciplina":     {"type": "string"},
                "ano_escolar":    {"type": "string"},
                "tempo_estimado": {"type": "string"},
                "habilidades_bncc": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "codigo":    {"type": "string"},
                            "descricao": {"type": "string"}
                        },
                        "required": ["codigo", "descricao"]
                    }
                },
                "desenvolvimento": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "etapa":                 {"type": "string"},
                            "conteudo":              {"type": "string"},
                            "estrategias_didaticas": {"type": "string"},
                            "recursos_pedagogicos":  {"type": "array", "items": {"type": "string"}}
                        },
                        "required": ["etapa", "conteudo", "estrategias_didaticas", "recursos_pedagogicos"]
                    }
                },
                "avaliacao_e_fechamento": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "metodo":    {"type": "string"},
                        "criterios": {"type": "string"}
                    },
                    "required": ["metodo", "criterios"]
                }
            },
            "required": [
                "tema_central", "disciplina", "ano_escolar", "tempo_estimado",
                "habilidades_bncc", "desenvolvimento", "avaliacao_e_fechamento"
            ]
        }
    },
    "required": ["plano_de_aula"]
}

# ─── Helpers de IA (Gemini first, Claude fallback) ────────────────────────────

def _gemini_disponivel():
    return bool(_gemini_model and GEMINI_API_KEY)

def _to_gemini_parts(content):
    """Converte conteúdo de mensagem (str ou lista multimodal) para partes do Gemini."""
    if isinstance(content, str):
        return [content]
    parts = []
    for item in content:
        if item.get('type') == 'text':
            parts.append(item['text'])
        elif item.get('type') == 'image':
            import base64 as _b64
            parts.append({
                'mime_type': item['source']['media_type'],
                'data': _b64.b64decode(item['source']['data'])
            })
    return parts

def _llm_cadeia_chat(sistema, messages, max_tokens=4000):
    """Cadeia de fallback: OpenAI (gpt-4o-mini) → Gemini → Claude.
    Retorna o texto da resposta ou lança RuntimeError se todos falharem."""
    ultimo_erro = None

    # 1. OpenAI — motor primário
    if client_openai:
        try:
            resp = client_openai.chat.completions.create(
                model='gpt-4o-mini',
                max_tokens=max_tokens,
                messages=[{'role': 'system', 'content': sistema}] + messages
            )
            logger.info('LLM: OpenAI gpt-4o-mini')
            return resp.choices[0].message.content
        except Exception as e:
            ultimo_erro = e
            logger.warning('OpenAI falhou, tentando Gemini: %s', str(e)[:150])
    else:
        logger.debug('OPENAI_API_KEY ausente, pulando OpenAI')

    # 2. Gemini — primeiro fallback
    if _gemini_disponivel():
        try:
            import google.generativeai as genai
            gm = genai.GenerativeModel(
                model_name='gemini-2.0-flash',
                system_instruction=sistema
            )
            msgs_g = []
            for m in messages:
                role = 'user' if m['role'] == 'user' else 'model'
                cont = m['content'] if isinstance(m['content'], str) else str(m['content'])
                msgs_g.append({'role': role, 'parts': [cont]})
            chat_g = gm.start_chat(history=msgs_g[:-1])
            resp_g = chat_g.send_message(msgs_g[-1]['parts'][0] if msgs_g else '')
            logger.info('LLM: Gemini gemini-2.0-flash')
            return resp_g.text
        except Exception as e:
            ultimo_erro = e
            logger.warning('Gemini falhou, tentando Claude: %s', str(e)[:150])
    else:
        logger.debug('GEMINI_API_KEY ausente, pulando Gemini')

    # 3. Claude — último recurso
    if os.environ.get('ANTHROPIC_API_KEY'):
        try:
            resposta = client.messages.create(
                model='claude-sonnet-4-6',
                max_tokens=max_tokens,
                system=sistema,
                messages=messages
            )
            logger.info('LLM: Claude claude-sonnet-4-6')
            return resposta.content[0].text
        except Exception as e:
            ultimo_erro = e
            logger.warning('Claude também falhou: %s', str(e)[:150])
    else:
        logger.debug('ANTHROPIC_API_KEY ausente, pulando Claude')

    raise RuntimeError(f'Todos os motores de IA falharam. Último erro: {ultimo_erro}')


def _llm_cadeia_simples(prompt, sistema='', max_tokens=4000):
    """Cadeia de fallback para prompt único. OpenAI → Gemini → Claude."""
    messages = [{'role': 'user', 'content': prompt}]
    return _llm_cadeia_chat(sistema or SYSTEM_PROMPT, messages, max_tokens=max_tokens)


# Aliases para compatibilidade com código legado
def chamar_ia_chat(sistema, messages):
    return _llm_cadeia_chat(sistema, messages)

def chamar_ia_simples(prompt):
    return _llm_cadeia_simples(prompt)


# ─── Banco de dados ───────────────────────────────────────────────────────────

DATABASE_URL = os.environ.get('DATABASE_URL', '').replace('postgres://', 'postgresql://', 1)

class _DbConn:
    """Wrapper que faz psycopg2 se comportar como sqlite3 no resto do código."""
    def __init__(self, conn):
        self._conn = conn

    def execute(self, sql, params=()):
        sql = sql.replace('?', '%s')
        params = tuple(
            psycopg2.Binary(p) if isinstance(p, (bytes, bytearray)) else p
            for p in params
        )
        cur = self._conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql, params)
        return cur

    def commit(self):
        self._conn.commit()

    def close(self):
        self._conn.close()

def get_db():
    if not DATABASE_URL:
        raise RuntimeError('DATABASE_URL não configurada. Adicione a variável de ambiente no Render.')
    conn = psycopg2.connect(DATABASE_URL)
    return _DbConn(conn)

def init_db():
    conn = get_db()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS usuarios (
            id         SERIAL PRIMARY KEY,
            nome       TEXT NOT NULL,
            email      TEXT UNIQUE NOT NULL,
            senha      TEXT NOT NULL,
            plano      TEXT DEFAULT '',
            ativo      INTEGER DEFAULT 0,
            valido_ate TEXT DEFAULT '',
            criado_em  TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS historico (
            id           SERIAL PRIMARY KEY,
            usuario_id   INTEGER DEFAULT 0,
            data         TEXT,
            professor    TEXT,
            escola       TEXT,
            disciplina   TEXT,
            turma        TEXT,
            num_aulas    INTEGER,
            periodo      TEXT,
            datas        TEXT,
            temas        TEXT,
            arquivo      BYTEA,
            nome_arquivo TEXT
        )
    ''')
    conn.execute('ALTER TABLE historico ADD COLUMN IF NOT EXISTS usuario_id INTEGER DEFAULT 0')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS chat_messages (
            id         SERIAL PRIMARY KEY,
            usuario_id INTEGER,
            role       TEXT,
            content    TEXT,
            criado_em  TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS planejamento_anual (
            id         SERIAL PRIMARY KEY,
            usuario_id INTEGER,
            disciplina TEXT,
            turma      TEXT,
            ano        TEXT,
            conteudo   TEXT,
            criado_em  TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS reset_tokens (
            id         SERIAL PRIMARY KEY,
            usuario_id INTEGER,
            token      TEXT,
            expira_em  TEXT,
            usado      INTEGER DEFAULT 0
        )
    ''')
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_template TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS onboarding_done INTEGER DEFAULT 0")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_nome TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS professor_nome TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS default_segment TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS logo_path TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS logo_estado_path TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_id INTEGER DEFAULT NULL")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS papel TEXT DEFAULT 'professor'")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_governo TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_secretaria TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_diretoria TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_endereco TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_fone TEXT DEFAULT ''")
    conn.execute("ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS escola_email TEXT DEFAULT ''")
    conn.execute('''
        CREATE TABLE IF NOT EXISTS questions_bank (
            id           SERIAL PRIMARY KEY,
            usuario_id   INTEGER NOT NULL,
            disciplina   TEXT,
            serie        TEXT,
            tipo         TEXT DEFAULT 'multipla_escolha',
            dificuldade  TEXT DEFAULT 'medio',
            enunciado    TEXT NOT NULL,
            alternativas TEXT,
            resposta_correta TEXT,
            bncc_codigo  TEXT,
            criado_em    TEXT
        )
    ''')
    conn.execute("ALTER TABLE questions_bank ADD COLUMN IF NOT EXISTS gabarito TEXT DEFAULT ''")
    conn.execute("ALTER TABLE questions_bank ADD COLUMN IF NOT EXISTS ano_serie TEXT DEFAULT ''")
    conn.execute("ALTER TABLE questions_bank ADD COLUMN IF NOT EXISTS habilidade_bncc TEXT DEFAULT ''")
    conn.execute('''
        CREATE TABLE IF NOT EXISTS referrals (
            id         SERIAL PRIMARY KEY,
            usuario_id INTEGER UNIQUE,
            codigo     TEXT UNIQUE,
            usos       INTEGER DEFAULT 0,
            conversoes INTEGER DEFAULT 0,
            creditos   INTEGER DEFAULT 0,
            criado_em  TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS escolas (
            id        SERIAL PRIMARY KEY,
            nome      TEXT NOT NULL,
            cnpj      TEXT,
            plano     TEXT DEFAULT 'escola',
            ativo     INTEGER DEFAULT 1,
            criado_em TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS escola_membros (
            id         SERIAL PRIMARY KEY,
            escola_id  INTEGER NOT NULL,
            usuario_id INTEGER NOT NULL,
            papel      TEXT DEFAULT 'professor',
            ativo      INTEGER DEFAULT 1,
            criado_em  TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS escola_convites (
            id        SERIAL PRIMARY KEY,
            escola_id INTEGER NOT NULL,
            email     TEXT NOT NULL,
            token     TEXT UNIQUE,
            usado     INTEGER DEFAULT 0,
            criado_em TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS lista_vip (
            id         SERIAL PRIMARY KEY,
            nome       TEXT NOT NULL,
            email      TEXT UNIQUE NOT NULL,
            whatsapp   TEXT DEFAULT '',
            criado_em  TEXT
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS leads (
            id             SERIAL PRIMARY KEY,
            nome           TEXT NOT NULL,
            contato        TEXT NOT NULL,
            tema_pesquisado TEXT DEFAULT '',
            criado_em      TEXT
        )
    ''')
    # ─── Índices de performance ────────────────────────────────────────────────
    conn.execute('CREATE INDEX IF NOT EXISTS idx_historico_usuario ON historico(usuario_id)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_chat_messages_usuario ON chat_messages(usuario_id)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_questions_bank_usuario ON questions_bank(usuario_id)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_questions_bank_disciplina ON questions_bank(disciplina)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_questions_bank_ano_serie ON questions_bank(ano_serie)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_planejamento_usuario ON planejamento_anual(usuario_id)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_reset_tokens_token ON reset_tokens(token)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_usuarios_email ON usuarios(email)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_historico_data ON historico(data)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_chat_messages_criado_em ON chat_messages(criado_em)')
    conn.commit()
    conn.close()

try:
    init_db()
except Exception as _init_err:
    logger.critical('init_db falhou — servidor iniciando sem DB: %s', _init_err)

# ─── Modelo de usuário ────────────────────────────────────────────────────────

class Usuario(UserMixin):
    def __init__(self, row):
        self.id              = row['id']
        self.nome            = row['nome']
        self.email           = row['email']
        self.plano           = row['plano']
        self.ativo           = row['ativo']
        self.valido_ate      = row['valido_ate']
        self.escola_template  = row.get('escola_template', '') or ''
        self.onboarding_done  = row.get('onboarding_done', 0) or 0
        self.escola_nome      = row.get('escola_nome', '') or ''
        self.professor_nome   = row.get('professor_nome', '') or ''
        self.default_segment  = row.get('default_segment', '') or ''
        # Valida prefixo para evitar path traversal caso o DB seja comprometido
        _logo = row.get('logo_path', '') or ''
        self.logo_path = _logo if _logo.startswith('static/logos/') else ''
        _logo_e = row.get('logo_estado_path', '') or ''
        self.logo_estado_path = _logo_e if _logo_e.startswith('static/logos/') else ''
        self.escola_id        = row.get('escola_id', None)
        self.papel            = row.get('papel', 'professor') or 'professor'
        self.escola_governo   = row.get('escola_governo', '') or ''
        self.escola_secretaria = row.get('escola_secretaria', '') or ''
        self.escola_diretoria  = row.get('escola_diretoria', '') or ''
        self.escola_endereco   = row.get('escola_endereco', '') or ''
        self.escola_fone       = row.get('escola_fone', '') or ''
        self.escola_email      = row.get('escola_email', '') or ''

    @property
    def assinatura_ativa(self):
        if not self.ativo or not self.valido_ate:
            return False
        try:
            valido = datetime.strptime(self.valido_ate, '%Y-%m-%d')
            return valido >= datetime.now()
        except Exception:
            return False

    @property
    def is_admin(self):
        return ADMIN_EMAIL and self.email == ADMIN_EMAIL

@login_manager.user_loader
def load_user(user_id):
    conn = get_db()
    row = conn.execute(
        'SELECT id, nome, email, senha, plano, ativo, valido_ate, criado_em,'
        ' escola_nome, professor_nome, logo_path, logo_estado_path,'
        ' escola_template, onboarding_done, escola_id, papel, default_segment,'
        ' escola_governo, escola_secretaria, escola_diretoria,'
        ' escola_endereco, escola_fone, escola_email'
        ' FROM usuarios WHERE id = %s', (user_id,)).fetchone()
    conn.close()
    return Usuario(row) if row else None

# ─── Helper: ativar assinatura ────────────────────────────────────────────────

def ativar_assinatura(usuario_id, plano_id):
    """Ativa ou renova assinatura de um usuário. Usado por todos os gateways."""
    if plano_id not in PLANOS:
        logger.warning('ativar_assinatura: plano inválido "%s" para usuário %s', plano_id, usuario_id)
        return False
    dias = PLANOS[plano_id]['dias']
    valido_ate = (datetime.now() + timedelta(days=dias)).strftime('%Y-%m-%d')
    conn = get_db()
    conn.execute(
        'UPDATE usuarios SET ativo = 1, plano = ?, valido_ate = ? WHERE id = ?',
        (plano_id, valido_ate, usuario_id)
    )
    conn.commit()
    conn.close()
    logger.info('Assinatura ativada: usuário %s → plano %s até %s', usuario_id, plano_id, valido_ate)
    return True

# ─── Helper: verificar assinatura ─────────────────────────────────────────────

def assinatura_required(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for('login'))
        if not current_user.assinatura_ativa and not current_user.is_admin:
            return redirect(url_for('chat'))
        return f(*args, **kwargs)
    return decorated

# ─── Helper: verificar admin ──────────────────────────────────────────────────

def admin_required(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for('login'))
        if not current_user.is_admin:
            return jsonify({'erro': 'Acesso negado'}), 403
        return f(*args, **kwargs)
    return decorated

# ─── Handler 429 (rate limit) ─────────────────────────────────────────────────

@app.errorhandler(429)
def ratelimit_handler(e):
    logger.warning('Rate limit atingido: %s — %s', _limiter_key(), request.path)
    # Rotas HTML recebem flash; rotas de API recebem JSON
    if request.path.startswith('/api/') or request.is_json:
        return jsonify({
            'erro': 'muitas_requisicoes',
            'msg': 'Muitas requisições em pouco tempo. Aguarde um momento e tente novamente.'
        }), 429
    flash('Muitas tentativas em pouco tempo. Aguarde um momento e tente novamente.', 'erro')
    return redirect(request.referrer or url_for('login'))

# ─── Email helper ─────────────────────────────────────────────────────────────

def enviar_email(to, subject, body_html):
    if not SMTP_HOST or not SMTP_USER or not SMTP_PASS:
        return False
    try:
        msg = MIMEText(body_html, 'html', 'utf-8')
        msg['Subject'] = subject
        msg['From'] = FROM_EMAIL
        msg['To'] = to
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as s:
            s.ehlo(); s.starttls(); s.login(SMTP_USER, SMTP_PASS)
            s.sendmail(FROM_EMAIL, [to], msg.as_string())
        return True
    except Exception as e:
        logger.error('Falha ao enviar email para %s: %s', to, e)
        return False

# ─── PDF (reportlab) ──────────────────────────────────────────────────────────

AZUL        = colors.HexColor('#2b4fc7')
AZUL_ESCURO = colors.HexColor('#1a3399')
AZUL_CLARO  = colors.HexColor('#eef2ff')
BRANCO      = colors.white
TEXTO       = colors.HexColor('#1a1a2e')
CINZA       = colors.HexColor('#6b7280')

def criar_pdf(dados_form, aulas_ia):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=1.8*rcm, rightMargin=1.8*rcm,
        topMargin=1.5*rcm, bottomMargin=1.5*rcm)

    escola     = dados_form.get('escola', '').strip()
    diretoria  = dados_form.get('diretoria', '').strip()
    endereco   = dados_form.get('endereco', '').strip()
    ano_letivo = dados_form.get('ano_letivo', str(datetime.now().year))

    st_centro_negrito = ParagraphStyle('cn', fontName='Helvetica-Bold', fontSize=10,
                                        alignment=TA_CENTER, textColor=TEXTO, leading=14)
    st_centro         = ParagraphStyle('c',  fontName='Helvetica', fontSize=9,
                                        alignment=TA_CENTER, textColor=TEXTO, leading=13)
    st_centro_pequeno = ParagraphStyle('cp', fontName='Helvetica', fontSize=8,
                                        alignment=TA_CENTER, textColor=CINZA, leading=12)
    st_header_tabela  = ParagraphStyle('ht', fontName='Helvetica-Bold', fontSize=8,
                                        alignment=TA_CENTER, textColor=BRANCO, leading=11)
    st_celula_titulo  = ParagraphStyle('ct', fontName='Helvetica-Bold', fontSize=8.5,
                                        alignment=TA_CENTER, textColor=AZUL, leading=12)
    st_celula         = ParagraphStyle('ce', fontName='Helvetica', fontSize=7.5,
                                        textColor=TEXTO, leading=11)
    st_sub            = ParagraphStyle('s',  fontName='Helvetica-Bold', fontSize=7.5,
                                        textColor=AZUL, leading=11)
    st_rodape         = ParagraphStyle('r',  fontName='Helvetica-Oblique', fontSize=7,
                                        alignment=TA_CENTER, textColor=CINZA, leading=10)

    story = []

    if escola or diretoria:
        story.append(Paragraph("GOVERNO DO ESTADO DE SÃO PAULO", st_centro_negrito))
        story.append(Paragraph("SECRETARIA DE ESTADO DA EDUCAÇÃO", st_centro))
        if diretoria:
            story.append(Paragraph(diretoria.upper(), st_centro))
        if escola:
            story.append(Paragraph(escola.upper(), st_centro_negrito))
        if endereco:
            story.append(Paragraph(endereco, st_centro_pequeno))
        story.append(Spacer(1, 8))

    titulo_data = [[Paragraph(f"PLANEJAMENTO DE AULA  {ano_letivo}", ParagraphStyle(
        'tit', fontName='Helvetica-Bold', fontSize=13,
        alignment=TA_CENTER, textColor=BRANCO, leading=16))]]
    t_titulo = Table(titulo_data, colWidths=[doc.width])
    t_titulo.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), AZUL),
        ('TOPPADDING',    (0,0), (-1,-1), 10),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
    ]))
    story.append(t_titulo)
    story.append(Spacer(1, 10))

    def info(label, valor):
        return Paragraph(f'<font color="#2b4fc7"><b>{label}:</b></font> {valor}',
                         ParagraphStyle('inf', fontName='Helvetica', fontSize=8.5,
                                        textColor=TEXTO, leading=13))

    w = doc.width
    t_info = Table([
        [info("Professor(a)", dados_form.get('professor', ''))],
        [info("Componente Curricular", dados_form.get('disciplina', '')),
         info("Nº de Aulas", str(dados_form.get('num_aulas', '')))],
        [info("Ano/Série/Turma", dados_form.get('turma', '')),
         info("Período", f"{dados_form.get('periodo','')}  |  {dados_form.get('datas','')}")],
    ], colWidths=[w*0.65, w*0.35])
    t_info.setStyle(TableStyle([
        ('SPAN',         (0,0), (1,0)),
        ('GRID',         (0,0), (-1,-1), 0.5, colors.HexColor('#c0c8e8')),
        ('TOPPADDING',    (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING',   (0,0), (-1,-1), 8),
        ('RIGHTPADDING',  (0,0), (-1,-1), 8),
        ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(t_info)
    story.append(Spacer(1, 10))

    col_w = [2.5*rcm, 5.2*rcm, 4.0*rcm, 2.9*rcm, 3.2*rcm]
    headers = [Paragraph(h, st_header_tabela) for h in
               ['AULA / DATA', 'CONTEÚDO E OBJETIVOS', 'ESTRATÉGIAS DIDÁTICAS', 'RECURSOS', 'AVALIAÇÃO']]
    rows = [headers]

    for i, aula in enumerate(aulas_ia):
        bg = AZUL_CLARO if i % 2 == 0 else BRANCO
        partes = aula['titulo'].split(' - ', 1)
        col0 = [Paragraph(partes[0], st_celula_titulo),
                Paragraph(partes[1] if len(partes) > 1 else '', ParagraphStyle(
                    's0', fontName='Helvetica', fontSize=7.5,
                    alignment=TA_CENTER, textColor=CINZA, leading=11))]
        col1 = [Paragraph("CONTEÚDOS", st_sub)]
        for c in aula.get('conteudos', []):
            col1.append(Paragraph(f"• {c}", st_celula))
        col1.append(Spacer(1, 4))
        col1.append(Paragraph("OBJETIVOS", st_sub))
        for o in aula.get('objetivos', []):
            col1.append(Paragraph(f"• {o}", st_celula))
        rows.append([col0, col1,
                     [Paragraph(aula.get('estrategias', ''), st_celula)],
                     [Paragraph(aula.get('recursos', ''), st_celula)],
                     [Paragraph(aula.get('avaliacao', ''), st_celula)]])

    t_main = Table(rows, colWidths=col_w, repeatRows=1)
    style = [
        ('BACKGROUND',    (0,0), (-1,0), AZUL),
        ('GRID',         (0,0), (-1,-1), 0.5, colors.HexColor('#c0c8e8')),
        ('VALIGN',       (0,0), (-1,-1), 'TOP'),
        ('TOPPADDING',    (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING',   (0,0), (-1,-1), 6),
        ('RIGHTPADDING',  (0,0), (-1,-1), 6),
        ('ALIGN',        (0,1), (0,-1), 'CENTER'),
        ('VALIGN',       (0,1), (0,-1), 'MIDDLE'),
    ]
    for i in range(len(aulas_ia)):
        bg = AZUL_CLARO if i % 2 == 0 else BRANCO
        style.append(('BACKGROUND', (0, i+1), (-1, i+1), bg))
    t_main.setStyle(TableStyle(style))
    story.append(t_main)
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}  •  Plano de Aula IA",
        st_rodape))
    doc.build(story)
    buf.seek(0)
    return buf

# ─── PDF extrator ──────────────────────────────────────────────────────────────

def extrair_pdf(url):
    try:
        import pdfplumber
        import urllib.request
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            urllib.request.urlretrieve(url, tmp.name)
            with pdfplumber.open(tmp.name) as pdf:
                texto = ''
                for page in pdf.pages[:10]:
                    t = page.extract_text()
                    if t:
                        texto += t + '\n'
            os.unlink(tmp.name)
        return texto[:3000]
    except Exception:
        return None

# ─── IA ───────────────────────────────────────────────────────────────────────

def gerar_conteudo_ia(disciplina, turma, temas, periodo, datas, aula_inicio=1, conteudo_pdf=None):
    referencia_pdf = ""
    if conteudo_pdf:
        referencia_pdf = f"\n\nMATERIAL DE REFERÊNCIA:\n{conteudo_pdf}\n\nUse esse material como base."

    prompt = f"""Você é um assistente especializado em educação brasileira.
Gere o conteúdo para um plano de aula seguindo EXATAMENTE este formato JSON.

Dados:
- Disciplina: {disciplina}
- Turma: {turma}
- Período: {periodo}
- Datas: {datas}
- Temas das aulas: {temas}
- Numeração começa na aula: {aula_inicio}{referencia_pdf}

Retorne SOMENTE um JSON válido neste formato (sem markdown, sem explicações):
{{
  "aulas": [
    {{
      "numero": {aula_inicio},
      "titulo": "Aula {aula_inicio} - [título baseado no tema]",
      "conteudos": ["conteúdo 1", "conteúdo 2", "conteúdo 3"],
      "objetivos": ["objetivo 1", "objetivo 2", "objetivo 3"],
      "estrategias": "Descrição das estratégias didáticas em 2-3 frases.",
      "recursos": "Kit Multimídia, quadro branco, [outros recursos relevantes]",
      "avaliacao": "Observar participação e desempenho dos alunos. [avaliação específica]"
    }}
  ]
}}

Gere {len(temas)} aulas, uma para cada tema. A primeira aula é número {aula_inicio}. Seja específico e pedagógico."""

    texto = chamar_ia_simples(prompt).strip()
    if texto.startswith("```"):
        texto = texto.split("```")[1]
        if texto.startswith("json"):
            texto = texto[4:]
    return json.loads(texto.strip())

# ─── DOCX ─────────────────────────────────────────────────────────────────────

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def set_cell_border(cell, color='c0c8e8'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), color)
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_run(paragraph, text, bold=False, size=9, color='1a1a2e', italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    r, g, b = hex_to_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    return run

def criar_docx(dados_form, aulas_ia):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    escola     = dados_form.get('escola', '').strip()
    diretoria  = dados_form.get('diretoria', '').strip()
    endereco   = dados_form.get('endereco', '').strip()
    ano_letivo = dados_form.get('ano_letivo', str(datetime.now().year))

    def header_line(text, bold=False, size=9):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, text, bold=bold, size=size, color='1a1a2e')

    header_line("GOVERNO DO ESTADO DE SÃO PAULO", bold=True, size=10)
    header_line("SECRETARIA DE ESTADO DA EDUCAÇÃO")
    if diretoria:
        header_line(diretoria.upper())
    if escola:
        header_line(escola.upper(), bold=True, size=10)
    if endereco:
        header_line(endereco, size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    t0 = doc.add_table(rows=1, cols=1)
    t0.style = 'Table Grid'
    cell = t0.cell(0, 0)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, f"PLANEJAMENTO DE AULA  {ano_letivo}", bold=True, size=12, color='FFFFFF')
    set_cell_bg(cell, '2b4fc7')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    t1 = doc.add_table(rows=3, cols=2)
    t1.style = 'Table Grid'

    def info_cell(cell, label, value):
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, f"{label}: ", bold=True, size=9, color='2b4fc7')
        add_run(p, value, size=9, color='333333')

    info_cell(t1.cell(0, 0), "Professor(a)", dados_form['professor'])
    t1.cell(0, 0).merge(t1.cell(0, 1))
    info_cell(t1.cell(1, 0), "Componente Curricular", dados_form['disciplina'])
    info_cell(t1.cell(1, 1), "Nº de Aulas", str(dados_form['num_aulas']))
    info_cell(t1.cell(2, 0), "Ano/Série/Turma", dados_form['turma'])
    info_cell(t1.cell(2, 1), "Período", f"{dados_form['periodo']}  |  {dados_form['datas']}")

    for row in t1.rows:
        for c in row.cells:
            set_cell_border(c)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    t2 = doc.add_table(rows=1 + len(aulas_ia), cols=5)
    t2.style = 'Table Grid'
    headers = ['AULA / DATA', 'CONTEÚDO E OBJETIVOS', 'ESTRATÉGIAS DIDÁTICAS', 'RECURSOS', 'AVALIAÇÃO']
    widths  = [Cm(2.8), Cm(5.8), Cm(4.5), Cm(3.3), Cm(3.6)]

    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = t2.cell(0, i)
        cell.width = w
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, h, bold=True, size=8, color='FFFFFF')
        set_cell_bg(cell, '2b4fc7')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, color='1a3399')

    for i, aula in enumerate(aulas_ia):
        ri = i + 1
        bg = 'f0f4ff' if i % 2 == 0 else 'FFFFFF'

        cell = t2.cell(ri, 0)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        partes = aula['titulo'].split(' - ', 1)
        add_run(p, partes[0] + '\n', bold=True, size=9, color='2b4fc7')
        if len(partes) > 1:
            add_run(p, partes[1], size=8, color='555555')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, bg)

        cell = t2.cell(ri, 1)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, "CONTEÚDOS\n", bold=True, size=8, color='2b4fc7')
        for c in aula['conteudos']:
            add_run(p, f"• {c}\n", size=8)
        add_run(p, "\nOBJETIVOS\n", bold=True, size=8, color='2b4fc7')
        for o in aula['objetivos']:
            add_run(p, f"• {o}\n", size=8)
        set_cell_bg(cell, bg)

        for col_i, key in enumerate(['estrategias', 'recursos', 'avaliacao'], start=2):
            cell = t2.cell(ri, col_i)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            add_run(p, aula[key], size=8)
            set_cell_bg(cell, bg)

        for ci in range(5):
            set_cell_border(t2.cell(ri, ci))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}  •  Plano de Aula IA",
            size=7, color='aaaaaa', italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ─── Auth ─────────────────────────────────────────────────────────────────────

@app.route('/login', methods=['GET', 'POST'])
@limiter.limit('10 per minute', methods=['POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('chat'))
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        senha = request.form.get('senha', '')
        conn  = get_db()
        row   = conn.execute(
            'SELECT id, nome, email, senha, plano, ativo, valido_ate, criado_em,'
            ' escola_nome, professor_nome, logo_path, logo_estado_path,'
            ' escola_template, onboarding_done, escola_id, papel'
            ' FROM usuarios WHERE email = ?', (email,)).fetchone()
        conn.close()
        if row and check_password_hash(row['senha'], senha):
            login_user(Usuario(row))
            return redirect(url_for('chat'))
        flash('E-mail ou senha incorretos.', 'erro')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/cadastro', methods=['GET', 'POST'])
@limiter.limit('5 per minute', methods=['POST'])
def cadastro():
    if current_user.is_authenticated:
        return redirect(url_for('chat'))
    if request.method == 'POST':
        nome  = request.form.get('nome', '').strip()[:120]
        email = request.form.get('email', '').strip().lower()[:254]
        senha = request.form.get('senha', '')[:128]
        aceito_termos = request.form.get('aceito_termos')
        if not nome or not email or not senha:
            flash('Preencha todos os campos.', 'erro')
            return render_template('cadastro.html')
        if not aceito_termos:
            flash('Você precisa aceitar os Termos de Uso e a Política de Privacidade para criar uma conta.', 'erro')
            return render_template('cadastro.html')
        if len(senha) < 6:
            flash('A senha deve ter pelo menos 6 caracteres.')
            return render_template('cadastro.html')
        conn = get_db()
        existe = conn.execute('SELECT id FROM usuarios WHERE email = ?', (email,)).fetchone()
        if existe:
            conn.close()
            flash('Este e-mail já está cadastrado.')
            return render_template('cadastro.html')
        try:
            conn.execute(
                'INSERT INTO usuarios (nome, email, senha, criado_em) VALUES (?, ?, ?, ?)',
                (nome, email, generate_password_hash(senha), datetime.now().strftime('%Y-%m-%d'))
            )
            conn.commit()
            row = conn.execute('SELECT * FROM usuarios WHERE email = ?', (email,)).fetchone()
            conn.close()
            login_user(Usuario(row))
            _capi_event('CompleteRegistration', user_data={'email': email, 'name': nome})
            return redirect(url_for('chat'))
        except Exception as e:
            conn.close()
            flash('Erro ao criar conta. Tente novamente.')
            return render_template('cadastro.html')
    return render_template('cadastro.html')

# ─── Recuperação de senha ─────────────────────────────────────────────────────

@app.route('/esqueci-senha', methods=['GET', 'POST'])
@limiter.limit('5 per minute', methods=['POST'])
def esqueci_senha():
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        conn = get_db()
        row = conn.execute("SELECT id FROM usuarios WHERE email = ?", (email,)).fetchone()
        if row:
            token = secrets.token_urlsafe(32)
            expira = (datetime.now() + timedelta(hours=2)).strftime('%Y-%m-%d %H:%M:%S')
            conn.execute("INSERT INTO reset_tokens (usuario_id, token, expira_em, usado) VALUES (?, ?, ?, 0)",
                        (row['id'], token, expira))
            conn.commit()
            link = f"{SITE_URL}/redefinir-senha/{token}"
            enviado = enviar_email(email, 'ProfessorIA — Redefinição de senha',
                f'<p>Clique para redefinir sua senha (válido por 2h):</p><p><a href="{link}">{link}</a></p>')
            if enviado:
                flash('Email enviado! Verifique sua caixa de entrada.', 'ok')
            else:
                logger.warning('SMTP não configurado — link de reset gerado para %s (não exposto ao usuário)', email)
                flash('Se esse email estiver cadastrado, você receberá as instruções.', 'ok')
        else:
            flash('Se esse email estiver cadastrado, você receberá as instruções.', 'ok')
        conn.close()
        return redirect(url_for('esqueci_senha'))
    return render_template('esqueci_senha.html')

@app.route('/redefinir-senha/<token>', methods=['GET', 'POST'])
@limiter.limit('5 per 10 minutes', methods=['POST'])
def redefinir_senha(token):
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT * FROM reset_tokens WHERE token = ? AND usado = 0", (token,)).fetchone()
        # Mensagem unificada para token inválido/expirado/já usado — evita enumeração de tokens
        if not row:
            flash('Link inválido, expirado ou já utilizado. Solicite um novo.', 'erro')
            return redirect(url_for('esqueci_senha'))
        if datetime.strptime(row['expira_em'], '%Y-%m-%d %H:%M:%S') < datetime.now():
            flash('Link inválido, expirado ou já utilizado. Solicite um novo.', 'erro')
            return redirect(url_for('esqueci_senha'))
        if request.method == 'POST':
            senha = request.form.get('senha', '')
            confirma = request.form.get('confirma', '')
            if len(senha) < 6:
                flash('Senha deve ter pelo menos 6 caracteres.', 'erro')
                return render_template('redefinir_senha.html', token=token)
            if senha != confirma:
                flash('As senhas não coincidem.', 'erro')
                return render_template('redefinir_senha.html', token=token)
            conn.execute("UPDATE usuarios SET senha = ? WHERE id = ?",
                        (generate_password_hash(senha), row['usuario_id']))
            conn.execute("UPDATE reset_tokens SET usado = 1 WHERE id = ?", (row['id'],))
            conn.commit()
            flash('Senha atualizada com sucesso!', 'ok')
            return redirect(url_for('login'))
        return render_template('redefinir_senha.html', token=token)
    finally:
        conn.close()

# ─── Planos e Pagamento ───────────────────────────────────────────────────────

@app.route('/planos')
@login_required
def planos():
    return render_template('planos.html', planos=PLANOS,
                           assinatura_ativa=current_user.assinatura_ativa,
                           valido_ate=current_user.valido_ate,
                           plano_atual=current_user.plano)

# ─── Stripe ───────────────────────────────────────────────────────────────────

@app.route('/stripe/checkout/<plano_id>')
@login_required
def stripe_checkout(plano_id):
    if plano_id not in PLANOS:
        flash(f'Plano inválido: {plano_id}', 'erro')
        return redirect(url_for('chat'))
    if not STRIPE_SECRET_KEY:
        flash('Chave Stripe não configurada no servidor (STRIPE_SECRET_KEY).', 'erro')
        return redirect(url_for('chat'))
    price_id = STRIPE_PRICES.get(plano_id, '')
    if not price_id:
        flash(f'Price ID não configurado para o plano "{plano_id}".', 'erro')
        return redirect(url_for('chat'))
    try:
        stripe_lib.api_key = STRIPE_SECRET_KEY
        session = stripe_lib.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{'price': price_id, 'quantity': 1}],
            mode='subscription',
            customer_email=current_user.email,
            metadata={'usuario_id': str(current_user.id), 'plano_id': plano_id},
            success_url=f"{SITE_URL}/stripe/sucesso?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{SITE_URL}/planos",
            allow_promotion_codes=True,
        )
        _capi_event('InitiateCheckout',
                    user_data={'email': current_user.email, 'name': current_user.nome},
                    custom_data={'content_name': plano_id,
                                 'value': PLANOS[plano_id]['preco'], 'currency': 'BRL'})
        return redirect(session.url, code=303)
    except Exception as e:
        logger.error('Stripe checkout erro: %s', e)
        flash(f'Erro ao iniciar pagamento. Tente novamente.', 'erro')
        return redirect(url_for('chat'))


@app.route('/stripe/sucesso')
@login_required
def stripe_sucesso():
    session_id = request.args.get('session_id', '')
    if session_id and STRIPE_SECRET_KEY:
        try:
            stripe_lib.api_key = STRIPE_SECRET_KEY
            session = stripe_lib.checkout.Session.retrieve(session_id)
            if session.payment_status in ('paid', 'no_payment_required'):
                plano_id = session.metadata.get('plano_id', 'basic')
                ativar_assinatura(current_user.id, plano_id)
                conn = get_db()
                row = conn.execute('SELECT * FROM usuarios WHERE id=?', (current_user.id,)).fetchone()
                conn.close()
                if row:
                    login_user(Usuario(row))
                preco = PLANOS.get(plano_id, {}).get('preco', 0)
                _capi_event('Purchase',
                            user_data={'email': current_user.email, 'name': current_user.nome},
                            custom_data={'content_name': plano_id,
                                         'value': preco, 'currency': 'BRL'})
        except Exception as e:
            logger.warning('Stripe sucesso: erro ao verificar sessão %s — %s', session_id, e)
    return render_template('pagamento_status.html',
                           status='sucesso',
                           titulo='PAGAMENTO APROVADO',
                           mensagem='Sua assinatura está ativa. Bom trabalho!')


@app.route('/stripe/webhook', methods=['POST'])
def stripe_webhook():
    stripe_lib.api_key = STRIPE_SECRET_KEY
    payload = request.get_data()
    sig     = request.headers.get('Stripe-Signature', '')
    try:
        event = stripe_lib.Webhook.construct_event(payload, sig, STRIPE_WEBHOOK_SECRET)
    except stripe_lib.errors.SignatureVerificationError as e:
        logger.warning('Stripe webhook: assinatura inválida — %s', e)
        return '', 400
    except Exception as e:
        logger.error('Stripe webhook: erro ao construir evento — %s', e)
        return '', 400

    etype = event['type']
    logger.info('Stripe evento recebido: %s', etype)

    if etype == 'checkout.session.completed':
        session  = event['data']['object']
        plano_id = session.get('metadata', {}).get('plano_id', 'basic')
        uid      = session.get('metadata', {}).get('usuario_id')
        if uid and plano_id in PLANOS:
            ativar_assinatura(int(uid), plano_id)

    elif etype == 'invoice.payment_succeeded':
        # Renovação mensal/anual — mantém a assinatura ativa no banco
        invoice = event['data']['object']
        sub_id  = invoice.get('subscription')
        if sub_id:
            try:
                sub      = stripe_lib.Subscription.retrieve(sub_id)
                customer = stripe_lib.Customer.retrieve(sub.customer)
                email    = getattr(customer, 'email', '') or ''
                if email:
                    conn = get_db()
                    row  = conn.execute('SELECT id, plano FROM usuarios WHERE email=?', (email,)).fetchone()
                    conn.close()
                    if row:
                        ativar_assinatura(row['id'], row['plano'] or 'basic')
                        logger.info('Stripe renovação: usuário %s renovado via invoice', row['id'])
            except Exception as e:
                logger.error('Stripe invoice.payment_succeeded: erro — %s', e)

    elif etype in ('customer.subscription.deleted', 'customer.subscription.paused'):
        sub         = event['data']['object']
        customer_id = sub.get('customer', '')
        if customer_id:
            try:
                customer = stripe_lib.Customer.retrieve(customer_id)
                email    = getattr(customer, 'email', '') or ''
                if email:
                    conn = get_db()
                    conn.execute("UPDATE usuarios SET plano='', ativo=0, valido_ate='' WHERE email=?", (email,))
                    conn.commit()
                    conn.close()
                    logger.info('Stripe: assinatura cancelada/pausada para %s', email)
            except Exception as e:
                logger.error('Stripe subscription cancel: erro — %s', e)

    return '', 200


# ─── Admin ────────────────────────────────────────────────────────────────────

@app.route('/admin')
@admin_required
def admin():
    conn  = get_db()
    users = conn.execute('SELECT * FROM usuarios ORDER BY id DESC LIMIT 200').fetchall()
    conn.close()
    return render_template('admin.html', users=users)

@app.route('/admin/ativar/<int:uid>', methods=['POST'])
@admin_required
def admin_ativar(uid):
    dias = int(request.form.get('dias', 30))
    plano = request.form.get('plano', 'professor')
    valido_ate = (datetime.now() + timedelta(days=dias)).strftime('%Y-%m-%d')
    conn = get_db()
    conn.execute('UPDATE usuarios SET ativo = 1, plano = ?, valido_ate = ? WHERE id = ?',
                 (plano, valido_ate, uid))
    conn.commit()
    conn.close()
    return redirect(url_for('admin'))

@app.route('/admin/desativar/<int:uid>', methods=['POST'])
@admin_required
def admin_desativar(uid):
    conn = get_db()
    conn.execute('UPDATE usuarios SET ativo = 0 WHERE id = ?', (uid,))
    conn.commit()
    conn.close()
    return redirect(url_for('admin'))

@app.route('/admin/update', methods=['POST'])
@admin_required
def admin_update():
    data   = request.json or {}
    uid    = data.get('uid')
    action = data.get('action')
    if not uid:
        return jsonify({'erro': 'uid obrigatório'}), 400
    conn = get_db()
    if action == 'ativar':
        dias = int(data.get('dias', 30))
        plano = data.get('plano', 'professor')
        if dias == 0:
            valido_ate = '2099-12-31'
        else:
            valido_ate = (datetime.now() + timedelta(days=dias)).strftime('%Y-%m-%d')
        conn.execute('UPDATE usuarios SET ativo=1, plano=?, valido_ate=? WHERE id=?',
                     (plano, valido_ate, uid))
    elif action == 'desativar':
        conn.execute('UPDATE usuarios SET ativo=0 WHERE id=?', (uid,))
    elif action == 'set_plan':
        conn.execute('UPDATE usuarios SET plano=? WHERE id=?', (data.get('plano','professor'), uid))
    else:
        conn.close()
        return jsonify({'erro': 'Ação inválida'}), 400
    conn.commit()
    row = conn.execute('SELECT id, ativo, plano, valido_ate FROM usuarios WHERE id=?', (uid,)).fetchone()
    conn.close()
    return jsonify({'ok': True, 'ativo': row['ativo'], 'plano': row['plano'],
                    'valido_ate': row['valido_ate'] or ''})

# ─── Conta / Perfil ───────────────────────────────────────────────────────────

@app.route('/perfil')
@login_required
def conta():
    conn = get_db()
    rows = conn.execute(
        'SELECT num_aulas, disciplina FROM historico WHERE usuario_id = ?',
        (current_user.id,)
    ).fetchall()
    conn.close()
    stats = {
        'total':       len(rows),
        'aulas':       sum(int(r['num_aulas'] or 0) for r in rows),
        'disciplinas': len(set(r['disciplina'] for r in rows if r['disciplina']))
    }
    return render_template('conta.html', stats=stats)

@app.route('/perfil/senha', methods=['POST'])
@login_required
def conta_senha():
    senha_atual = request.form.get('senha_atual', '')
    senha_nova  = request.form.get('senha_nova', '')
    senha_conf  = request.form.get('senha_conf', '')

    conn = get_db()
    row  = conn.execute('SELECT senha FROM usuarios WHERE id = ?', (current_user.id,)).fetchone()

    if not check_password_hash(row['senha'], senha_atual):
        conn.close()
        flash('Senha atual incorreta.', 'erro')
        return redirect(url_for('conta'))

    if senha_nova != senha_conf:
        conn.close()
        flash('As senhas não coincidem.', 'erro')
        return redirect(url_for('conta'))

    if len(senha_nova) < 6:
        conn.close()
        flash('A nova senha deve ter pelo menos 6 caracteres.', 'erro')
        return redirect(url_for('conta'))

    conn.execute('UPDATE usuarios SET senha = ? WHERE id = ?',
                 (generate_password_hash(senha_nova), current_user.id))
    conn.commit()
    conn.close()
    flash('Senha atualizada com sucesso!', 'ok')
    return redirect(url_for('conta'))


@app.route('/conta/escola', methods=['POST'])
@login_required
def conta_escola():
    """Salva os dados da escola enviados pelo formulário HTML da página Minha Conta."""
    f = request.form
    escola   = f.get('escola_nome', '').strip()[:200]
    prof     = f.get('professor_nome', current_user.professor_nome or '').strip()[:200]
    gov      = f.get('escola_governo', '').strip()[:200]
    sec      = f.get('escola_secretaria', '').strip()[:200]
    dire     = f.get('escola_diretoria', '').strip()[:200]
    ender    = f.get('escola_endereco', '').strip()[:300]
    fone     = f.get('escola_fone', '').strip()[:50]
    email_e  = f.get('escola_email', '').strip()[:200]
    conn = get_db()
    conn.execute(
        "UPDATE usuarios SET escola_nome=?, professor_nome=?,"
        " escola_governo=?, escola_secretaria=?, escola_diretoria=?,"
        " escola_endereco=?, escola_fone=?, escola_email=?"
        " WHERE id=?",
        (escola, prof, gov, sec, dire, ender, fone, email_e, current_user.id)
    )
    conn.commit()
    conn.close()
    flash('Dados da escola salvos com sucesso!', 'ok')
    return redirect(url_for('conta'))


@app.route('/conta/senha', methods=['POST'])
@login_required
def conta_senha_alias():
    """Alias de /perfil/senha para compatibilidade com o formulário da página Minha Conta."""
    return conta_senha()


@app.route('/api/profile', methods=['GET', 'PUT'])
@login_required
def api_profile():
    """Perfil Global do professor — fricção zero.
    GET  → { display_name, school_name }
    PUT  → { display_name?, school_name? }  (campos opcionais, atualiza só o que vier)
    """
    if request.method == 'GET':
        return jsonify({
            'display_name': current_user.professor_nome or '',
            'school_name':  current_user.escola_nome or '',
        })

    data         = request.get_json(force=True) or {}
    display_name = data.get('display_name')
    school_name  = data.get('school_name')

    # Atualiza apenas os campos enviados
    updates, params = [], []
    if display_name is not None:
        updates.append('professor_nome = ?')
        params.append(str(display_name).strip()[:200])
    if school_name is not None:
        updates.append('escola_nome = ?')
        params.append(str(school_name).strip()[:200])

    if not updates:
        return jsonify({'erro': 'Nenhum campo para atualizar'}), 400

    params.append(current_user.id)
    conn = get_db()
    conn.execute(f"UPDATE usuarios SET {', '.join(updates)} WHERE id = ?", params)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ─── Rotas principais ─────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/vip')
def vip():
    return render_template('index.html', abrir_modal=True)

@app.route('/teste3d')
def teste3d():
    return render_template('landing3d.html')

@app.route('/historico')
@login_required
@assinatura_required
def historico_page():
    return render_template('historico.html')

@app.route('/api/historico')
@login_required
@assinatura_required
def api_historico():
    conn = get_db()
    rows = conn.execute(
        '''SELECT id, data, professor, escola, disciplina, turma,
                  num_aulas, periodo, datas, temas, nome_arquivo
           FROM historico WHERE usuario_id = ? ORDER BY id DESC LIMIT 100''',
        (current_user.id,)
    ).fetchall()
    conn.close()
    result = []
    for r in rows:
        try:
            temas_parsed = json.loads(r['temas']) if r['temas'] else []
        except (json.JSONDecodeError, TypeError):
            temas_parsed = []
        result.append({
            'id': r['id'], 'data': r['data'], 'professor': r['professor'],
            'escola': r['escola'], 'disciplina': r['disciplina'],
            'turma': r['turma'], 'num_aulas': r['num_aulas'],
            'periodo': r['periodo'], 'datas': r['datas'],
            'temas': temas_parsed,
            'nome_arquivo': r['nome_arquivo']
        })
    return jsonify(result)

@app.route('/download/<int:item_id>')
@login_required
@assinatura_required
def download_historico(item_id):
    conn = get_db()
    row  = conn.execute(
        'SELECT arquivo, nome_arquivo FROM historico WHERE id = ? AND usuario_id = ?',
        (item_id, current_user.id)
    ).fetchone()
    conn.close()
    if not row:
        return 'Não encontrado', 404
    buf = io.BytesIO(row['arquivo'])
    return send_file(buf, as_attachment=True, download_name=row['nome_arquivo'],
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/deletar/<int:item_id>', methods=['DELETE'])
@login_required
@assinatura_required
def deletar_historico(item_id):
    conn = get_db()
    conn.execute('DELETE FROM historico WHERE id = ? AND usuario_id = ?',
                 (item_id, current_user.id))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/gerar', methods=['POST'])
@login_required
@assinatura_required
@limiter.limit('10 per minute')
def gerar():
    dados = {
        'professor':  (request.form.get('professor') or current_user.professor_nome or '')[:200],
        'escola':     (request.form.get('escola') or current_user.escola_nome or '')[:200],
        'diretoria':  request.form.get('diretoria', '')[:200],
        'endereco':   request.form.get('endereco', '')[:300],
        'ano_letivo': request.form.get('ano_letivo', str(datetime.now().year))[:4],
        'disciplina': request.form.get('disciplina', '')[:100],
        'turma':      request.form.get('turma', '')[:50],
        'num_aulas':  request.form.get('num_aulas', '1')[:3],
        'aula_inicio':request.form.get('aula_inicio', '1')[:3],
        'periodo':    request.form.get('periodo', 'quinzenal')[:20],
        'datas':      request.form.get('datas', '')[:500],
    }
    temas    = [t[:200] for t in request.form.getlist('temas[]')]
    urls_pdf = [u.strip() for u in request.form.getlist('urls_pdf[]') if u.strip()]
    formato  = request.form.get('formato', 'docx')

    conteudo_pdf = None
    if urls_pdf:
        partes = [f"--- PDF {i+1} ---\n{t}" for i, u in enumerate(urls_pdf)
                  if (t := extrair_pdf(u))]
        if partes:
            conteudo_pdf = "\n\n".join(partes)

    try:
        conteudo = gerar_conteudo_ia(dados['disciplina'], dados['turma'], temas,
                                     dados['periodo'], dados['datas'],
                                     int(dados.get('aula_inicio', 1)), conteudo_pdf)
    except Exception as e:
        flash(f'Erro ao gerar conteúdo: {str(e)[:200]}', 'erro')
        return redirect(url_for('chat'))

    base_nome = f"Plano_{dados['disciplina'].replace(' ', '_')}_{dados['turma'].replace(' ', '_')}"

    try:
        if formato == 'pdf':
            buf        = criar_pdf(dados, conteudo['aulas'])
            file_bytes = buf.read()
            nome       = base_nome + '.pdf'
            mimetype   = 'application/pdf'
        else:
            buf        = criar_docx(dados, conteudo['aulas'])
            file_bytes = buf.read()
            nome       = base_nome + '.docx'
            mimetype   = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    except Exception as e:
        flash(f'Erro ao gerar arquivo: {str(e)[:200]}', 'erro')
        return redirect(url_for('chat'))

    conn = get_db()
    conn.execute(
        '''INSERT INTO historico
           (usuario_id, data, professor, escola, disciplina, turma,
            num_aulas, periodo, datas, temas, arquivo, nome_arquivo)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
        (current_user.id, datetime.now().strftime('%d/%m/%Y %H:%M'),
         dados['professor'], dados['escola'], dados['disciplina'],
         dados['turma'], dados['num_aulas'], dados['periodo'],
         dados['datas'], json.dumps(temas, ensure_ascii=False), file_bytes, nome)
    )
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(file_bytes), as_attachment=True,
                     download_name=nome, mimetype=mimetype)


# ─── API Plano de Aula — Structured Output ────────────────────────────────────

@app.route('/api/gerar-plano', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_gerar_plano():
    """Gera um plano de aula estruturado via JSON Schema.
    Entrada: { tema, ano, disciplina }
    Saída:   { plano_de_aula: { ... } }
    """
    if not current_user.assinatura_ativa and not current_user.is_admin:
        geracoes = get_geracoes_mes(current_user.id)
        if geracoes >= LIMITE_GRATIS:
            return jsonify({
                'erro': 'limite_atingido',
                'geracoes': geracoes,
                'cta': '/planos',
                'mensagem': 'Você atingiu o limite do plano grátis. Faça upgrade para continuar gerando materiais!'
            }), 403

    data = request.get_json(force=True) or {}
    tema       = str(data.get('tema', '')).strip()[:300]
    ano        = str(data.get('ano', '')).strip()[:50]
    disciplina = str(data.get('disciplina', '')).strip()[:100]

    if not tema or not ano or not disciplina:
        return jsonify({'erro': 'Campos obrigatórios: tema, ano, disciplina'}), 400

    user_prompt = (
        f"Gere um plano de aula completo para:\n"
        f"- Tema: {tema}\n"
        f"- Ano/Série: {ano}\n"
        f"- Disciplina: {disciplina}\n\n"
        "Use habilidades BNCC reais e metodologias ativas. "
        "Inclua pelo menos 2 etapas de desenvolvimento (Introdução e Prática)."
    )

    plano_json = None
    erro_motores = []

    # ── Tenta Claude (tool_use para schema garantido) ──
    try:
        if os.environ.get('ANTHROPIC_API_KEY'):
            resp = client.messages.create(
                model='claude-sonnet-4-6',
                max_tokens=4000,
                system=SYSTEM_PROMPT_PLANO,
                tools=[PLANO_AULA_TOOL],
                tool_choice={"type": "tool", "name": "salvar_plano_de_aula"},
                messages=[{"role": "user", "content": user_prompt}]
            )
            for block in resp.content:
                if block.type == 'tool_use' and block.name == 'salvar_plano_de_aula':
                    plano_json = block.input
                    break
            if plano_json is None:
                # Claude respondeu mas sem tool_use — tenta extrair JSON do texto
                for block in resp.content:
                    if hasattr(block, 'text'):
                        txt = block.text.strip()
                        m = re.search(r'\{[\s\S]+\}', txt)
                        if m:
                            try:
                                plano_json = json.loads(m.group())
                            except Exception:
                                pass
                        break
                if plano_json is None:
                    erro_motores.append('Claude: resposta sem bloco tool_use')
                    logger.warning('api_gerar_plano — Claude sem tool_use, stop_reason=%s', resp.stop_reason)
    except Exception as e:
        erro_motores.append(f'Claude: {e}')
        logger.warning('api_gerar_plano — Claude falhou: %s', e)

    # ── Fallback: Gemini com response_schema ──
    if plano_json is None and _gemini_disponivel():
        try:
            import google.generativeai as genai

            gemini_schema = {
                'type': 'object',
                'properties': {
                    'tema':                  {'type': 'string'},
                    'habilidades_bncc':      {'type': 'array', 'items': {'type': 'string'}},
                    'objetivos':             {'type': 'array', 'items': {'type': 'string'}},
                    'conteudo_programatico': {'type': 'string'},
                    'metodologia':           {'type': 'string'},
                    'recursos_didaticos':    {'type': 'array', 'items': {'type': 'string'}},
                    'avaliacao':             {'type': 'string'},
                }
            }
            gm = genai.GenerativeModel(
                model_name='gemini-2.0-flash',
                system_instruction=SYSTEM_PROMPT_PLANO,
                generation_config=genai.GenerationConfig(
                    response_mime_type='application/json',
                    response_schema=gemini_schema
                )
            )
            resp_g = gm.generate_content(user_prompt)
            plano_json = json.loads(resp_g.text)
        except Exception as e:
            erro_motores.append(f'Gemini: {e}')
            logger.warning('api_gerar_plano — Gemini falhou: %s', e)

    # ── Fallback: OpenAI com json_object ──
    if plano_json is None:
        try:
            oai_key = os.environ.get('OPENAI_API_KEY')
            if oai_key:
                import openai as _oai
                oai_client = _oai.OpenAI(api_key=oai_key)
                resp_o = oai_client.chat.completions.create(
                    model='gpt-4o-mini',
                    messages=[
                        {'role': 'system',  'content': SYSTEM_PROMPT_PLANO},
                        {'role': 'user',    'content': user_prompt}
                    ],
                    response_format={'type': 'json_object'},
                    max_tokens=4000
                )
                plano_json = json.loads(resp_o.choices[0].message.content)
        except Exception as e:
            erro_motores.append(f'OpenAI: {e}')
            logger.warning('api_gerar_plano — OpenAI falhou: %s', e)

    if plano_json is None:
        return jsonify({
            'erro': 'Todos os motores de IA falharam. Verifique as chaves de API.',
            'detalhes': erro_motores
        }), 503

    # ── Coordenador Pedagógico Sênior: revisão e refinamento do plano ──────────
    try:
        plano_bruto_str = json.dumps(plano_json, ensure_ascii=False)
        plano_refinado_str = _llm_cadeia_simples(
            f"Revise e melhore este plano de aula, retornando o JSON completo revisado:\n{plano_bruto_str}",
            sistema=SYSTEM_PROMPT_COORDENADOR,
            max_tokens=4000
        )
        # Extrai o JSON da resposta do Coordenador
        m = re.search(r'\{[\s\S]+\}', plano_refinado_str)
        if m:
            plano_refinado = json.loads(m.group())
            plano_json = plano_refinado
            logger.info('api_gerar_plano — Coordenador revisou o plano com sucesso.')
    except Exception as e:
        # Se o Coordenador falhar, mantém o plano original sem quebrar
        logger.warning('api_gerar_plano — Coordenador falhou (plano original mantido): %s', e)

    # Garante que disciplina e ano estejam no JSON para o PDF
    plano_json.setdefault('disciplina', disciplina)
    plano_json.setdefault('ano_escolar', ano)

    # Contabiliza geração
    conn = get_db()
    conn.execute(
        "INSERT INTO chat_messages (usuario_id, role, content, criado_em) VALUES (?, ?, ?, ?)",
        (current_user.id, 'assistant',
         f'[plano estruturado] {tema} — {disciplina} {ano}',
         datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    )
    conn.commit()
    conn.close()

    return jsonify(plano_json)


# ─── Helpers de plano ─────────────────────────────────────────────────────────

def get_geracoes_mes(usuario_id):
    mes_atual = datetime.now().strftime('%Y-%m')
    conn = get_db()
    row = conn.execute(
        "SELECT COUNT(*) as total FROM chat_messages WHERE usuario_id = ? AND role = 'assistant' AND criado_em LIKE ?",
        (usuario_id, f'{mes_atual}%')
    ).fetchone()
    conn.close()
    return row['total'] if row else 0

# ─── Chat ──────────────────────────────────────────────────────────────────────

@app.route('/api/salvar-template', methods=['POST'])
@login_required
def salvar_template():
    data = request.json or {}
    escola_nome     = data.get('escola_nome', '').strip()[:200]
    default_segment = data.get('default_segment', '').strip()[:100]
    # Mantém compatibilidade com campo legado
    template = data.get('template', '').strip()[:5000]
    conn = get_db()
    conn.execute(
        "UPDATE usuarios SET escola_template = ?, escola_nome = ?, default_segment = ?, onboarding_done = 1 WHERE id = ?",
        (template, escola_nome, default_segment, current_user.id)
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/chat')
@login_required
def chat():
    geracoes = get_geracoes_mes(current_user.id)
    tem_plano = current_user.assinatura_ativa or current_user.is_admin
    limite_atingido = not tem_plano and geracoes >= LIMITE_GRATIS
    return render_template('chat.html',
                           geracoes=geracoes,
                           limite=LIMITE_GRATIS,
                           limite_atingido=limite_atingido,
                           tem_plano=tem_plano,
                           plano=current_user.plano,
                           onboarding_done=current_user.onboarding_done)


@app.route('/api/processar-arquivo', methods=['POST'])
@login_required
def processar_arquivo():
    """Extrai texto de PDFs, DOCX e TXT enviados pelo professor."""
    if 'arquivo' not in request.files:
        return jsonify({'erro': 'Nenhum arquivo enviado'}), 400

    f    = request.files['arquivo']
    nome = f.filename or 'arquivo'
    mime = f.content_type or ''
    dados = f.read()

    MAX_SIZE = 10 * 1024 * 1024  # 10 MB
    if len(dados) > MAX_SIZE:
        return jsonify({'erro': 'Arquivo muito grande (máximo 10 MB)'}), 400

    nome_lower = nome.lower()

    # PDF
    if mime == 'application/pdf' or nome_lower.endswith('.pdf'):
        try:
            import pdfplumber, io as _io
            texto = ''
            with pdfplumber.open(_io.BytesIO(dados)) as pdf:
                for pg in pdf.pages[:30]:
                    t = pg.extract_text()
                    if t:
                        texto += t + '\n'
            if not texto.strip():
                return jsonify({'erro': 'Não foi possível extrair texto deste PDF.'}), 400
            return jsonify({'tipo': 'documento', 'texto': texto[:20000], 'nome': nome})
        except Exception as e:
            return jsonify({'erro': f'Erro ao ler PDF: {str(e)}'}), 500

    # DOCX
    if (mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            or nome_lower.endswith('.docx')):
        try:
            from docx import Document as _Doc
            import io as _io
            doc = _Doc(_io.BytesIO(dados))
            texto = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            return jsonify({'tipo': 'documento', 'texto': texto[:20000], 'nome': nome})
        except Exception as e:
            return jsonify({'erro': f'Erro ao ler DOCX: {str(e)}'}), 500

    # TXT / CSV
    if mime.startswith('text/') or nome_lower.endswith(('.txt', '.csv', '.md')):
        try:
            texto = dados.decode('utf-8', errors='ignore')
            return jsonify({'tipo': 'documento', 'texto': texto[:20000], 'nome': nome})
        except Exception as e:
            return jsonify({'erro': f'Erro ao ler arquivo: {str(e)}'}), 500

    return jsonify({'erro': 'Tipo não suportado via upload. Imagens são enviadas diretamente — use JPG, PNG ou WEBP.'}), 400


@app.route('/api/chat', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_chat():
    if not current_user.assinatura_ativa and not current_user.is_admin:
        geracoes = get_geracoes_mes(current_user.id)
        if geracoes >= LIMITE_GRATIS:
            return jsonify({
                'erro': 'limite_atingido',
                'geracoes': geracoes,
                'cta': '/planos',
                'mensagem': 'Você atingiu o limite do plano grátis. Faça upgrade para continuar gerando materiais!'
            }), 403

    data = request.json or {}
    messages = data.get('messages', [])
    anexo   = data.get('anexo')   # { tipo, base64, mime, nome } ou { tipo, texto, nome }
    if not messages:
        return jsonify({'erro': 'Mensagem vazia'}), 400
    if len(messages) > 100:
        return jsonify({'erro': 'Histórico muito longo. Inicie uma nova conversa.'}), 400

    # Extrai conteúdo de texto da última mensagem para salvar no DB
    last_content = messages[-1].get('content', '')
    if isinstance(last_content, list):
        text_parts = [p['text'] for p in last_content if p.get('type') == 'text']
        db_content = ' '.join(text_parts)
        if anexo:
            db_content += f' [arquivo: {anexo.get("nome", "")}]'
    else:
        db_content = str(last_content) if last_content else ''
    db_content = db_content[:4000]  # cap para armazenamento no banco

    conn = get_db()
    conn.execute(
        "INSERT INTO chat_messages (usuario_id, role, content, criado_em) VALUES (?, ?, ?, ?)",
        (current_user.id, 'user', db_content,
         datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    )
    conn.commit()
    conn.close()

    # Prepara sistema e mensagens antes de entrar no generator
    sistema = SYSTEM_PROMPT
    if current_user.escola_template:
        sistema += f"\n\nO professor usa o seguinte esqueleto/modelo padrão de plano de aula da sua escola. SEMPRE que gerar planos de aula, use EXATAMENTE essa estrutura como base:\n\n{current_user.escola_template}"

    if anexo:
        tipo_anexo = anexo.get('tipo')
        if tipo_anexo == 'documento':
            nome_doc = anexo.get('nome', 'documento')
            texto_doc = anexo.get('texto', '')
            sistema += f"\n\n=== DOCUMENTO ENVIADO PELO PROFESSOR: {nome_doc} ===\n{texto_doc}\n=== FIM DO DOCUMENTO ===\n\nUse este documento como referência e base estrutural para criar o material solicitado."
        elif tipo_anexo == 'image':
            b64    = anexo.get('base64', '')
            mime_t = anexo.get('mime', 'image/jpeg')
            texto_msg = messages[-1].get('content', '')
            if not isinstance(texto_msg, list):
                texto_msg = texto_msg or 'Use esta imagem como base para criar o material.'
            messages = messages[:-1] + [{
                'role': 'user',
                'content': [
                    {'type': 'image', 'source': {'type': 'base64', 'media_type': mime_t, 'data': b64}},
                    {'type': 'text',  'text': texto_msg if isinstance(texto_msg, str) else ' '.join(p.get('text','') for p in texto_msg if p.get('type')=='text')}
                ]
            }]

    usuario_id = current_user.id

    @stream_with_context
    def generate():
        chunks = []
        try:
            # ── 1. OpenAI streaming — motor primário ─────────────────────────
            openai_ok = False
            openai_key = os.environ.get('OPENAI_API_KEY', '').strip()
            if openai_key:
                try:
                    import requests as req_lib
                    openai_msgs = [{'role': 'system', 'content': sistema}] + [
                        {'role': m['role'],
                         'content': m['content'] if isinstance(m['content'], str) else str(m['content'])}
                        for m in messages
                    ]
                    ro = req_lib.post(
                        'https://api.openai.com/v1/chat/completions',
                        json={'model': 'gpt-4o-mini', 'max_tokens': 8000,
                              'messages': openai_msgs, 'stream': True},
                        headers={'Authorization': f'Bearer {openai_key}',
                                 'content-type': 'application/json'},
                        stream=True, timeout=120
                    )
                    if ro.status_code == 200:
                        for line in ro.iter_lines():
                            if not line:
                                continue
                            line = line.decode('utf-8')
                            if line.startswith('data: '):
                                line = line[6:]
                            if line == '[DONE]':
                                break
                            try:
                                parsed = json.loads(line)
                                if 'error' in parsed:
                                    raise RuntimeError(parsed['error'].get('message', 'Erro OpenAI'))
                                delta = parsed['choices'][0]['delta'].get('content', '')
                                if delta:
                                    chunks.append(delta)
                                    yield f"data: {json.dumps({'chunk': delta}, ensure_ascii=False)}\n\n"
                            except (json.JSONDecodeError, KeyError):
                                pass
                        openai_ok = bool(chunks)
                        if openai_ok:
                            logger.info('Streaming: OpenAI gpt-4o-mini')
                    else:
                        logger.warning('OpenAI streaming %s, tentando Gemini', ro.status_code)
                except Exception as oe:
                    logger.warning('OpenAI streaming falhou, tentando Gemini: %s', oe)
                    chunks = []
            else:
                logger.debug('OPENAI_API_KEY ausente, pulando OpenAI no streaming')

            # ── 2. Gemini streaming — primeiro fallback ───────────────────────
            gemini_ok = False
            if not openai_ok and _gemini_disponivel():
                try:
                    import google.generativeai as genai
                    historico_g = []
                    for m in messages[:-1]:
                        role = 'user' if m['role'] == 'user' else 'model'
                        historico_g.append({'role': role, 'parts': _to_gemini_parts(m['content'])})
                    gm = genai.GenerativeModel(model_name='gemini-2.0-flash', system_instruction=sistema)
                    chat_g = gm.start_chat(history=historico_g)
                    resp_g = chat_g.send_message(
                        _to_gemini_parts(messages[-1]['content']),
                        stream=True, request_options={'timeout': 60}
                    )
                    for chunk in resp_g:
                        text = getattr(chunk, 'text', '') or ''
                        if text:
                            chunks.append(text)
                            yield f"data: {json.dumps({'chunk': text}, ensure_ascii=False)}\n\n"
                    gemini_ok = bool(chunks)
                    if gemini_ok:
                        logger.info('Streaming: Gemini gemini-2.0-flash')
                    else:
                        logger.warning('Gemini retornou vazio, tentando Claude')
                except Exception as ge:
                    logger.warning('Gemini streaming falhou, tentando Claude: %s', ge)
                    chunks = []

            # ── 3. Claude streaming — último recurso ─────────────────────────
            if not openai_ok and not gemini_ok:
                try:
                    with client.messages.stream(
                        model='claude-sonnet-4-6',
                        max_tokens=8000,
                        system=sistema,
                        messages=messages
                    ) as stream:
                        for text in stream.text_stream:
                            chunks.append(text)
                            yield f"data: {json.dumps({'chunk': text}, ensure_ascii=False)}\n\n"
                    logger.info('Streaming: Claude claude-sonnet-4-6')
                except Exception as ce:
                    logger.error('Claude streaming também falhou: %s', ce)
                    if not chunks:
                        yield f"data: {json.dumps({'erro': 'Todos os motores de IA estão indisponíveis. Tente novamente em instantes.'}, ensure_ascii=False)}\n\n"
                        return

            yield "data: [DONE]\n\n"
            resposta = ''.join(chunks)
            conn2 = get_db()
            conn2.execute(
                "INSERT INTO chat_messages (usuario_id, role, content, criado_em) VALUES (?, ?, ?, ?)",
                (usuario_id, 'assistant', resposta, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            )
            conn2.commit(); conn2.close()

        except Exception as e:
            logger.error('Erro no streaming de IA: %s', traceback.format_exc())
            yield f"data: {json.dumps({'erro': str(e)}, ensure_ascii=False)}\n\n"

    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive',
        }
    )


@app.route('/api/transcribe', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_transcribe():
    """Transcreve áudio do usuário usando OpenAI Whisper."""
    if 'audio' not in request.files:
        return jsonify({'erro': 'Nenhum arquivo de áudio enviado'}), 400

    audio_file = request.files['audio']
    api_key = os.environ.get('OPENAI_API_KEY', '').strip()
    if not api_key:
        return jsonify({'erro': 'OPENAI_API_KEY não configurada no servidor'}), 500

    try:
        import requests as req_lib
        fname = audio_file.filename or 'audio.webm'
        ctype = audio_file.content_type or 'audio/webm'
        files = {'file': (fname, audio_file.read(), ctype)}
        data  = {'model': 'whisper-1', 'language': 'pt'}
        r = req_lib.post(
            'https://api.openai.com/v1/audio/transcriptions',
            headers={'Authorization': f'Bearer {api_key}'},
            files=files, data=data, timeout=30
        )
        if r.status_code != 200:
            return jsonify({'erro': f'Whisper {r.status_code}: {r.text[:200]}'}), 500
        texto = r.json().get('text', '').strip()
        return jsonify({'texto': texto})
    except Exception as e:
        return jsonify({'erro': str(e)}), 500


@app.route('/api/tts', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_tts():
    """Converte texto em áudio usando OpenAI TTS."""
    data = request.json or {}
    texto = data.get('texto', '').strip()[:4000]
    voz   = data.get('voz', 'nova')
    VOZES_VALIDAS = {'alloy', 'echo', 'fable', 'onyx', 'nova', 'shimmer'}
    if voz not in VOZES_VALIDAS:
        voz = 'nova'
    if not texto:
        return jsonify({'erro': 'Texto vazio'}), 400

    api_key = os.environ.get('OPENAI_API_KEY', '').strip()
    if not api_key:
        return jsonify({'erro': 'OPENAI_API_KEY não configurada no servidor'}), 500

    try:
        import requests as req_lib
        r = req_lib.post(
            'https://api.openai.com/v1/audio/speech',
            headers={'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'},
            json={'model': 'tts-1', 'input': texto, 'voice': voz},
            timeout=60
        )
        if r.status_code != 200:
            return jsonify({'erro': f'TTS {r.status_code}: {r.text[:200]}'}), 500
        from flask import Response
        return Response(r.content, mimetype='audio/mpeg',
                        headers={'Content-Disposition': 'inline; filename=resposta.mp3'})
    except Exception as e:
        return jsonify({'erro': str(e)}), 500


# ─── Gerador de DOCX com design ProfessorIA ────────────────────────────────────
# Inspirado nas fichas pedagógicas dos exemplos: cabeçalho de marca, campos do
# aluno, seções com caixas, grid monospace, rodapé — tudo black & white p/ impressão.

def _pr(paragraph, text, bold=False, size=10, color='0a0a0a',
        italic=False, font='Arial', underline=False):
    """Adiciona um run estilizado a um parágrafo."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size)
    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return run

def _pr_fmt(paragraph, text, size=10, font='Arial'):
    """Adiciona texto com suporte a **bold**, *italic* e `code` inline."""
    import re
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            _pr(paragraph, part[2:-2], bold=True, size=size, font=font)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            _pr(paragraph, part[1:-1], italic=True, size=size, font=font)
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            _pr(paragraph, part[1:-1], font='Courier New', size=size - 1, color='333333')
        elif part:
            _pr(paragraph, part, size=size, font=font)

def _pia_no_borders(table):
    """Remove todas as bordas de uma tabela."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                el.set(qn('w:sz'), '0')
                el.set(qn('w:color'), 'auto')
                tcBorders.append(el)
            tcPr.append(tcBorders)

def _pia_hrule(doc, thick=True, color='0a0a0a'):
    """Linha horizontal fina usando borda inferior de parágrafo."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '16' if thick else '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

def _pia_section_box(doc, title):
    """Caixa de seção: fundo preto, texto branco — como nos exemplos."""
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.25)
    _pr(p, '  ' + title.upper(), bold=True, size=10, color='FFFFFF')
    set_cell_bg(cell, '0a0a0a')
    # bordas pretas
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), '0a0a0a')
        tcBorders.append(el)
    tcPr.append(tcBorders)
    ep = doc.add_paragraph()
    ep.paragraph_format.space_after = Pt(3)

def _is_letter_grid(code):
    """Returns True if code block is a caça-palavras letter grid."""
    import re
    lines = [l for l in code.strip().split('\n') if l.strip()]
    if len(lines) < 5:
        return False
    letter_lines = 0
    for line in lines:
        tokens = line.split()
        if not tokens:
            continue
        if all(re.match(r'^\d+$', t) for t in tokens):
            continue  # skip pure number rows
        single = sum(1 for t in tokens if len(t) == 1 and t.isalpha())
        if single >= 8:
            letter_lines += 1
    return letter_lines >= 5


def _pia_caca_palavras_table(doc, code):
    """Renders a caça-palavras letter grid as a proper bordered Word table."""
    import re
    lines = [l for l in code.strip().split('\n') if l.strip()]

    grid = []
    for line in lines:
        tokens = line.split()
        if not tokens:
            continue
        if all(re.match(r'^\d+$', t) for t in tokens):
            continue  # skip pure number rows
        # Collect all single alpha tokens — works for both formats (with or without row prefix)
        row = [t.upper() for t in tokens if len(t) == 1 and t.isalpha()]
        if len(row) >= 8:
            grid.append(row)

    if not grid:
        _pia_code_block(doc, code)
        return

    n_rows = len(grid)
    n_cols = max(len(r) for r in grid)

    lbl = doc.add_paragraph()
    lbl.paragraph_format.space_before = Pt(6)
    lbl.paragraph_format.space_after  = Pt(3)
    _pr(lbl, 'GRADE DE LETRAS', bold=True, size=8, color='333333')

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    cell_w = Cm(15.5 / n_cols)

    for ri, row_letters in enumerate(grid):
        for ci in range(n_cols):
            cell = tbl.cell(ri, ci)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            letter = row_letters[ci] if ci < len(row_letters) else ' '
            _pr(p, letter, bold=True, size=10, color='0a0a0a')
            cell.width = cell_w
            tc    = cell._tc
            tcPr  = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), 'bbbbbb')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_after = Pt(8)


def _pia_code_block(doc, code):
    """Bloco monospace para grades de caça-palavras, cruzadinhas, mapas mentais."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    lines = code.split('\n')
    for idx, ln in enumerate(lines):
        if idx > 0:
            p.add_run().add_break()
        _pr(p, ln, font='Courier New', size=8, color='0a0a0a')
    p.paragraph_format.space_after = Pt(4)
    set_cell_bg(cell, 'f5f5f2')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'aaaaaa')
        tcBorders.append(el)
    tcPr.append(tcBorders)
    ep = doc.add_paragraph()
    ep.paragraph_format.space_after = Pt(4)

def _pia_md_table(doc, lines):
    """Renderiza tabela Markdown como tabela Word estilizada."""
    import re
    rows = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^\|[-\s|:]+\|$', stripped):
            continue  # linha separadora do markdown
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        if any(cells):
            rows.append(cells)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=max_cols)
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(max_cols):
            cell = t.cell(ri, ci)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            text = row[ci] if ci < len(row) else ''
            if ri == 0:
                set_cell_bg(cell, 'e8e8e5')
                _pr(p, text, bold=True, size=8, color='0a0a0a')
            else:
                if ri % 2 == 0:
                    set_cell_bg(cell, 'f8f8f5')
                _pr_fmt(p, text, size=9)
    ep = doc.add_paragraph()
    ep.paragraph_format.space_after = Pt(4)

# Palette: 6 vivid colors for mind map categories (matching emoji color order: red, blue, yellow, green, purple, orange)
_MM_PALETTE = [
    ('e53e3e', 'fff5f5'),  # red
    ('2b6cb0', 'ebf8ff'),  # blue
    ('d69e2e', 'fffff0'),  # yellow
    ('276749', 'f0fff4'),  # green
    ('6b46c1', 'faf5ff'),  # purple
    ('c05621', 'fffaf0'),  # orange
]


def _limpar_codigo_sujo(texto: str) -> str:
    """Remove blocos de código Mermaid/JSON vazados que a IA pode emitir indevidamente.
    Usado como filtro de entrada antes de parsear mapas mentais e planos de aula.
    Preserva o conteúdo do bloco para formatos que precisam do código raw (Mermaid).
    Se o bloco inteiro é Mermaid mindmap, devolve o bloco tal qual (para _extrair_mermaid).
    Para outros blocos de código (json, python, etc.) remove as cercas mas mantém o texto."""
    # Se o texto já é um bloco Mermaid puro (único bloco, sem texto antes/depois), retorna intacto
    puro = re.fullmatch(r'\s*```mermaid[\s\S]+?```\s*', texto, re.IGNORECASE)
    if puro:
        return texto.strip()
    # Remove cercas de código json / genéricas — mantém o conteúdo interno
    texto = re.sub(r'```(?:json|javascript|js|python|py|text|txt)?\s*\n([\s\S]+?)\n```',
                   r'\1', texto, flags=re.IGNORECASE)
    # Remove blocos Mermaid que vazaram dentro de uma resposta maior (não é o bloco principal)
    texto = re.sub(r'```mermaid[\s\S]+?```', '', texto, flags=re.IGNORECASE)
    # Remove cercas de código sozinhas (``` sem linguagem no final do texto)
    texto = re.sub(r'```\s*\n?', '', texto)
    return texto.strip()


def _extrair_mermaid(texto):
    """Extrai o código Mermaid de um bloco ```mermaid ... ```.
    Retorna a string do código ou None se não encontrado."""
    m = re.search(r'```mermaid\s*\n([\s\S]+?)\n```', texto, re.IGNORECASE)
    return m.group(1).strip() if m else None


def _parse_mermaid_mindmap(texto):
    """
    Parseia um bloco Mermaid mindmap e retorna (titulo, categorias)
    no mesmo formato de _parse_mapa_mental.
    Formato esperado:
        mindmap
          root((TITULO))
            Categoria1
              - item1
              - item2
            Categoria2
              item1
    """
    codigo = _extrair_mermaid(texto)
    if not codigo:
        return 'MAPA MENTAL', []

    lines = codigo.split('\n')
    titulo = 'MAPA MENTAL'
    # Extrai título do root((...))
    for line in lines:
        m = re.search(r'root\s*\(\((.+?)\)\)', line, re.IGNORECASE)
        if m:
            titulo = m.group(1).strip().upper()
            break

    # Determina indentação do root para calcular nível relativo
    root_indent = None
    for line in lines:
        if re.search(r'root\s*\(\(', line, re.IGNORECASE):
            root_indent = len(line) - len(line.lstrip())
            break
    if root_indent is None:
        root_indent = 2

    categoria_indent = root_indent + 2  # nível imediatamente abaixo do root
    item_indent = categoria_indent + 2  # nível dos itens

    categorias = []
    cat_atual = None
    ci = 0

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.lower().startswith('mindmap') or re.search(r'root\s*\(\(', line, re.IGNORECASE):
            continue
        indent = len(line) - len(line.lstrip())
        # Remove sintaxe Mermaid: ::, (), [], {}, parênteses especiais
        clean = re.sub(r'\[{1,2}.*?\]{1,2}|\({1,2}.*?\){1,2}|\{{1,2}.*?\}{1,2}|:::\w+', '', stripped)
        clean = re.sub(r'^[-•*]\s*', '', clean).strip()
        clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean).strip()
        if not clean:
            continue

        if indent <= categoria_indent:
            # Nova categoria
            if cat_atual and cat_atual['itens']:
                categorias.append(cat_atual)
            cat_atual = {'titulo': clean.upper(), 'cor_idx': ci, 'itens': []}
            ci += 1
        else:
            # Item da categoria atual
            if cat_atual is not None:
                cat_atual['itens'].append(clean)

    if cat_atual and cat_atual['itens']:
        categorias.append(cat_atual)

    return titulo, categorias


def _mermaid_para_png(codigo_mermaid):
    """Converte código Mermaid em bytes PNG usando a API pública mermaid.ink.
    Retorna bytes da imagem ou None em caso de falha."""
    import requests as req_lib
    try:
        # mermaid.ink aceita JSON encodado em base64 URL-safe
        config = json.dumps({'code': codigo_mermaid, 'mermaid': {'theme': 'default'}})
        encoded = base64.urlsafe_b64encode(config.encode()).decode().rstrip('=')
        url = f'https://mermaid.ink/img/{encoded}?type=png&width=1200&height=900'
        r = req_lib.get(url, timeout=30)
        if r.status_code == 200 and r.headers.get('content-type', '').startswith('image/'):
            return r.content
        logger.warning('mermaid.ink retornou status %s', r.status_code)
    except Exception as e:
        logger.warning('_mermaid_para_png falhou: %s', e)
    return None


def _detect_doc_type(texto):
    """Returns 'plano_aula', 'mapa_mental', or 'outro'."""
    t = texto.lower()
    # Detecta Mermaid mindmap
    if _extrair_mermaid(texto) and 'mindmap' in t:
        return 'mapa_mental'
    # Formato legado ## 🧠
    if ('🧠 tema central' in t or '## 🧠' in t or
            (('### 🔴' in t or '### 🔵' in t) and '## 🧠' in t)):
        return 'mapa_mental'
    # Sinais fortes — qualquer UM já indica plano de aula
    strong = ['### aula', '# planejamento da aula', '# planejamento de aula',
              'planejamento da aula —', 'planejamento de aula —']
    if any(s in t for s in strong):
        return 'plano_aula'
    # Sinais fracos — precisa de 2+
    weak = ['**conteúdo e objetivos', '**estratégias didáticas',
            '**recursos pedagógicos', '**avaliação:', 'componente curricular',
            'habilidades bncc', 'metodologia ativa']
    return 'plano_aula' if sum(1 for s in weak if s in t) >= 2 else 'outro'


def _parse_mapa_mental(texto):
    """
    Parses structured mind map text into (titulo, categorias).
    Returns (str, list of {'titulo': str, 'cor_idx': int, 'itens': [str]})
    """
    import re
    texto = _limpar_codigo_sujo(texto)
    titulo = 'MAPA MENTAL'
    m = re.search(r'##\s+🧠\s+TEMA\s+CENTRAL\s*:\s*(.+)', texto, re.IGNORECASE)
    if m:
        titulo = re.sub(r'[*_`]', '', m.group(1)).strip().upper()

    categorias = []
    # Split by ### sections
    parts = re.split(r'\n###\s+', texto)
    color_emojis = {'🔴': 0, '🔵': 1, '🟡': 2, '🟢': 3, '🟣': 4, '🟠': 5}
    ci = 0
    for part in parts[1:]:  # skip content before first ###
        lines = part.strip().split('\n')
        if not lines:
            continue
        cat_title = re.sub(r'[*_`]', '', lines[0]).strip()
        # Determine color index from emoji
        cor_idx = ci % len(_MM_PALETTE)
        for emoji, idx in color_emojis.items():
            if emoji in cat_title:
                cor_idx = idx
                break
        # Clean emoji from title
        cat_clean = re.sub(r'[\U0001F7E0-\U0001F7E6🧠🔴🔵🟡🟢🟣🟠]', '', cat_title).strip(' —:-')
        cat_clean = re.sub(r'\s+', ' ', cat_clean).strip()
        # Extract bullet items
        itens = []
        for line in lines[1:]:
            line = line.strip()
            if line.startswith('-') or line.startswith('•') or line.startswith('*'):
                item = re.sub(r'^[-•*]\s*', '', line).strip()
                item = re.sub(r'\*\*([^*]+)\*\*', r'\1', item)
                if item:
                    itens.append(item)
        if cat_clean and itens:
            categorias.append({'titulo': cat_clean.upper(), 'cor_idx': cor_idx, 'itens': itens})
        ci += 1
    return titulo, categorias


def _hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def gerar_mapa_mental_docx(texto, meta=None):
    """Gera DOCX visual de mapa mental no estilo infográfico Descomplica.
    Suporta formato legado (## 🧠) e Mermaid mindmap."""
    import re
    if meta is None:
        meta = {}
    escola    = meta.get('escola', '').strip()
    professor = meta.get('professor', '').strip()
    disciplina = meta.get('disciplina', '').strip()

    # Detecta qual formato e faz o parse correto
    if _extrair_mermaid(texto):
        titulo, categorias = _parse_mermaid_mindmap(texto)
    else:
        titulo, categorias = _parse_mapa_mental(texto)

    doc = Document()
    for sec in doc.sections:
        sec.page_height   = Cm(29.7)
        sec.page_width    = Cm(21.0)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    # ── HEADER ──────────────────────────────────────────────────────────
    if escola or professor:
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after  = Pt(2)
        parts = []
        if escola: parts.append(escola)
        if professor: parts.append(f'Prof(a). {professor}')
        if disciplina: parts.append(disciplina)
        run = hp.add_run('  ·  '.join(parts))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── CENTRAL TITLE BOX ───────────────────────────────────────────────
    title_tbl = doc.add_table(rows=1, cols=1)
    _pia_no_borders(title_tbl)
    tc = title_tbl.cell(0, 0)
    # Dark indigo background for title
    set_cell_bg(tc, '312e81')
    # Add colored top stripe feel via border
    tcPr = tc._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), '4338ca')
        tcBorders.append(el)
    tcPr.append(tcBorders)

    tp = tc.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(10)
    tp.paragraph_format.space_after  = Pt(10)
    tr = tp.add_run(titulo)
    tr.font.bold  = True
    tr.font.size  = Pt(18)
    tr.font.name  = 'Arial'
    tr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── CATEGORY GRID: 2 columns ─────────────────────────────────────────
    # Fill in pairs; if odd, last row has 1 cell spanning
    n = len(categorias)
    pairs = [(categorias[i], categorias[i+1] if i+1 < n else None)
             for i in range(0, n, 2)]

    for left_cat, right_cat in pairs:
        row_tbl = doc.add_table(rows=1, cols=2)
        _pia_no_borders(row_tbl)
        row_tbl.columns[0].width = Cm(8.5)
        row_tbl.columns[1].width = Cm(8.5)

        for col_i, cat in enumerate([left_cat, right_cat]):
            cell = row_tbl.cell(0, col_i)
            if cat is None:
                continue
            fg, bg = _MM_PALETTE[cat['cor_idx'] % len(_MM_PALETTE)]

            # Background of card
            set_cell_bg(cell, bg)

            # Card borders
            tc2  = cell._tc
            tcP2 = tc2.get_or_add_tcPr()
            tcB2 = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '8')
                el.set(qn('w:color'), fg)
                tcB2.append(el)
            tcP2.append(tcB2)

            # Cell left margin via inner table padding
            tcMar = OxmlElement('w:tcMar')
            for m_side, val in [('top','80'),('left','120'),('bottom','80'),('right','120')]:
                m_el = OxmlElement(f'w:{m_side}')
                m_el.set(qn('w:w'), val)
                m_el.set(qn('w:type'), 'dxa')
                tcMar.append(m_el)
            tcP2.append(tcMar)

            # Category title
            title_p = cell.paragraphs[0]
            title_p.paragraph_format.space_before = Pt(4)
            title_p.paragraph_format.space_after  = Pt(4)
            t_run = title_p.add_run(cat['titulo'])
            t_run.font.bold  = True
            t_run.font.size  = Pt(10)
            t_run.font.name  = 'Arial'
            t_run.font.color.rgb = _hex_to_rgb(fg)

            # Items
            for item in cat['itens']:
                item_p = cell.add_paragraph()
                item_p.paragraph_format.space_before = Pt(1)
                item_p.paragraph_format.space_after  = Pt(1)
                bullet = item_p.add_run('▸ ')
                bullet.font.size  = Pt(8)
                bullet.font.color.rgb = _hex_to_rgb(fg)
                i_run = item_p.add_run(item)
                i_run.font.size  = Pt(8.5)
                i_run.font.name  = 'Arial'
                i_run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)

            item_p = cell.add_paragraph()
            item_p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── FOOTER ──────────────────────────────────────────────────────────
    _pia_hrule(doc, thick=False, color='aaaaaa')
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(0)
    parts_f = ['Gerado por ProfessorIA™', datetime.now().strftime('%d/%m/%Y')]
    if escola: parts_f.append(escola)
    _pr(pf, '  ·  '.join(parts_f), size=7, color='888880')

    return doc


_PREAMBLE_PLANO_RE = re.compile(
    r'^[\s\S]*?(?=(?:#\s+PLANEJAMENTO|###\s+AULA\s+\d))',
    re.IGNORECASE
)
_PREAMBLE_FRASES = re.compile(
    r'^\s*(?:aqui\s+est[aá]|claro\s*[!,]|com\s+prazer|segue\s+(?:abaixo|o|um)|'
    r'certamente|ol[aá]\s*[!,]|vou\s+gerar|abaixo\s+(?:est[aá]|segue)|'
    r'segue\s+o\s+plano|preparei\s+o)[^\n]*\n+',
    re.IGNORECASE
)


def _limpar_preamble_plano(texto: str) -> str:
    """Remove frases de cortesia e preambles que a IA insere antes do conteúdo estruturado.
    O parser depende de ### AULA N ou # PLANEJAMENTO na primeira linha útil."""
    # Remove linhas de cortesia no início
    texto = _PREAMBLE_FRASES.sub('', texto)
    # Avança até o primeiro marcador estrutural esperado
    m = re.search(r'(?:^|\n)(#\s+PLANEJAMENTO|###\s+AULA\s+\d)', texto, re.IGNORECASE)
    if m:
        texto = texto[m.start():].lstrip('\n')
    return texto.strip()


def _parse_plano_aula(texto):
    """
    Extrai metadados e seções de aula do texto estruturado gerado pela IA.
    Retorna (meta_extra, aulas) onde aulas é lista de dicts.
    """
    import re
    texto = _limpar_preamble_plano(texto)
    meta_extra = {}

    for pattern, key in [
        (r'(?:nº\s+de\s+aulas|número\s+de\s+aulas)[^*\n]*?\*\*\s*[:\s]+([^\n|]+)', 'num_aulas'),
        (r'período[^*\n]*?\*\*\s*[:\s]+([^\n|]+)', 'periodo'),
        (r'data[^*\n]*?\*\*\s*[:\s]+([^\n|]+)', 'data_range'),
        (r'(?:ano|série)[^*\n]*?turma[^*\n]*?\*\*\s*[:\s]+([^\n|]+)', 'serie_turma'),
    ]:
        m = re.search(pattern, texto, re.IGNORECASE)
        if m:
            meta_extra[key] = re.sub(r'\*+', '', m.group(1)).strip(' |,')

    # Split by ### AULA N sections
    section_re = re.compile(r'(?:^|\n)#{2,3}\s+(AULA\s+\d+[^\n]*)', re.IGNORECASE)
    matches = list(section_re.finditer(texto))
    aulas = []

    def extract_field(corpo, keys):
        for k in keys:
            m = re.search(
                r'\*\*' + re.escape(k) + r'[^*]*?\*\*\s*[:\n]+([\s\S]*?)(?=\n\s*\*\*[A-ZÀ-Ú]|\n#{2,3}|\n---|\Z)',
                corpo, re.IGNORECASE
            )
            if m:
                return m.group(1).strip()
        return ''

    for idx, match in enumerate(matches):
        titulo = re.sub(r'[#*]+', '', match.group(1)).strip()
        start  = match.end()
        end    = matches[idx + 1].start() if idx + 1 < len(matches) else len(texto)
        corpo  = texto[start:end]
        aulas.append({
            'titulo':      titulo,
            'conteudo':    extract_field(corpo, ['conteúdo e objetivos de aprendizagem', 'conteúdo e objetivos', 'objetivos de aprendizagem', 'conteúdo']),
            'estrategias': extract_field(corpo, ['estratégias didáticas', 'estratégias', 'metodologia']),
            'recursos':    extract_field(corpo, ['recursos pedagógicos', 'recursos']),
            'avaliacao':   extract_field(corpo, ['avaliação', 'verificar se']),
        })

    # Fallback: sem seções ### AULA N — trata o texto inteiro como uma única aula
    if not aulas:
        aulas.append({
            'titulo':      'Aula',
            'conteudo':    extract_field(texto, ['conteúdo e objetivos de aprendizagem', 'conteúdo e objetivos', 'objetivos de aprendizagem', 'conteúdo', 'objetivos']),
            'estrategias': extract_field(texto, ['estratégias didáticas', 'estratégias', 'metodologia']),
            'recursos':    extract_field(texto, ['recursos pedagógicos', 'recursos didáticos', 'recursos']),
            'avaliacao':   extract_field(texto, ['avaliação', 'verificar se']),
        })
        # Se nem os campos estruturados existem, usa o texto completo como conteúdo
        if not any(aulas[0].values()):
            aulas[0]['conteudo'] = texto.strip()

    return meta_extra, aulas


def _set_cell_borders_plano(cell, color='cccccc'):
    """Adiciona bordas simples a uma célula da tabela de plano de aula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_bg_plano(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _gerar_pdf_mermaid(codigo_mermaid, titulo='Mapa Mental'):
    """Gera PDF A4 landscape com imagem do Mermaid (via mermaid.ink)."""
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.utils import ImageReader
    import io

    png_bytes = _mermaid_para_png(codigo_mermaid)

    W, H = landscape(A4)
    buf  = io.BytesIO()
    c    = rl_canvas.Canvas(buf, pagesize=(W, H))

    # Fundo branco
    c.setFillColor(white)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # Barra superior
    c.setFillColor(HexColor('#1E3A5F'))
    c.rect(0, H - 36, W, 36, fill=1, stroke=0)
    c.setFont('Helvetica-Bold', 13)
    c.setFillColor(white)
    c.drawCentredString(W / 2, H - 23, titulo.upper())

    # Rodapé
    c.setFont('Helvetica', 7)
    c.setFillColor(HexColor('#6b7280'))
    c.drawCentredString(W / 2, 14, f'Gerado por ProfessorIA™  ·  {datetime.now().strftime("%d/%m/%Y")}')

    if png_bytes:
        img_buf = io.BytesIO(png_bytes)
        img     = ImageReader(img_buf)
        iw, ih  = img.getSize()
        # Escala proporcional dentro da área útil
        margin  = 24
        max_w   = W - margin * 2
        max_h   = H - 36 - 28 - margin  # topo + rodapé + margem
        scale   = min(max_w / iw, max_h / ih)
        dw, dh  = iw * scale, ih * scale
        dx      = (W - dw) / 2
        dy      = 28 + margin + (max_h - dh) / 2
        c.drawImage(img, dx, dy, width=dw, height=dh, mask='auto')
    else:
        # Fallback: mensagem de erro
        c.setFont('Helvetica', 10)
        c.setFillColor(HexColor('#374151'))
        c.drawCentredString(W / 2, H / 2, 'Não foi possível renderizar o mapa mental. Tente novamente.')

    c.save()
    return buf.getvalue()


def gerar_mapa_mental_pdf(texto, meta=None):
    """Gera PDF visual de mapa mental.
    Suporta Mermaid (```mermaid mindmap```) e formato legado ## 🧠."""
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.colors import HexColor, white
    import io, math

    if meta is None:
        meta = {}

    # ── Mermaid mindmap — caminho novo ────────────────────────────────────────
    codigo_mermaid = _extrair_mermaid(texto)
    if codigo_mermaid and 'mindmap' in codigo_mermaid.lower():
        titulo_m = re.search(r'root\(\((.+?)\)\)', codigo_mermaid)
        titulo   = titulo_m.group(1).strip() if titulo_m else 'Mapa Mental'
        return _gerar_pdf_mermaid(codigo_mermaid, titulo=titulo)

    # ── Formato legado ## 🧠 — caminho existente ─────────────────────────────

    titulo, categorias = _parse_mapa_mental(texto)
    W, H = landscape(A4)   # 841.89 x 595.28 pts
    cx, cy = W / 2, H / 2  # 420.94, 297.64

    C_NAVY  = HexColor('#1E3A5F')
    C_BLUE  = HexColor('#1E40AF')
    C_LBLUE = HexColor('#BFDBFE')
    C_GRAY  = HexColor('#374151')
    C_LGRAY = HexColor('#F0F7FF')
    C_LINE  = HexColor('#93C5FD')

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=(W, H))

    # Fundo branco
    c.setFillColor(white)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # Blocos de cor nos cantos
    c.setFillColor(C_LGRAY)
    for rx, ry in [(0, H * 0.72), (W * 0.78, H * 0.72), (0, 0), (W * 0.78, 0)]:
        c.rect(rx, ry, W * 0.22, H * 0.28, fill=1, stroke=0)

    # Borda superior
    c.setFillColor(C_NAVY)
    c.rect(0, H - 7, W, 7, fill=1, stroke=0)
    c.setFillColor(C_BLUE)
    c.rect(0, H - 9, W, 2, fill=1, stroke=0)

    ORX, ORY   = 128, 60
    BOX_W      = 188
    HEADER_H   = 26
    ITEM_H     = 20
    MAX_ITEMS  = 6

    # Posições [bottom-left x, y] — suporta até 6 categorias
    POS = [
        (22,             H * 0.62),   # topo-esquerda
        (W - BOX_W - 22, H * 0.62),   # topo-direita
        (22,             H * 0.32),   # meio-esquerda
        (W - BOX_W - 22, H * 0.32),   # meio-direita
        (22,             H * 0.04),   # baixo-esquerda
        (W - BOX_W - 22, H * 0.04),   # baixo-direita
    ]

    n = min(len(categorias), 6)

    # ── Linhas de conexão (desenhadas antes das caixas) ──
    for i in range(n):
        cat = categorias[i]
        bx, by = POS[i]
        n_items = min(len(cat['itens']), MAX_ITEMS)
        box_h   = HEADER_H + n_items * ITEM_H + 6
        bcx     = bx + BOX_W / 2
        bcy     = by + box_h / 2
        angle   = math.atan2(bcy - cy, bcx - cx)
        ox      = cx + math.cos(angle) * ORX
        oy      = cy + math.sin(angle) * ORY
        mid_y   = (oy + bcy) / 2

        c.saveState()
        c.setStrokeColor(C_LINE)
        c.setLineWidth(2.2)
        p = c.beginPath()
        p.moveTo(ox, oy)
        p.curveTo(ox, mid_y, bcx, mid_y, bcx, bcy)
        c.drawPath(p, stroke=1, fill=0)
        c.setFillColor(C_LINE)
        c.circle(bcx, bcy, 4, fill=1, stroke=0)
        c.restoreState()

    # ── Caixas de categoria ──
    for i in range(n):
        cat    = categorias[i]
        bx, by = POS[i]
        items  = cat['itens'][:MAX_ITEMS]
        items_h = len(items) * ITEM_H + 6

        # Fundo dos itens
        c.saveState()
        c.setFillColor(HexColor('#F8FAFF'))
        c.setStrokeColor(C_LBLUE)
        c.setLineWidth(0.8)
        c.roundRect(bx, by, BOX_W, items_h, 5, fill=1, stroke=1)
        c.restoreState()

        # Cabeçalho
        c.saveState()
        c.setFillColor(C_NAVY)
        c.roundRect(bx, by + items_h, BOX_W, HEADER_H, 6, fill=1, stroke=0)
        c.setFillColor(white)
        title_str = cat['titulo']
        fs = 8.5
        c.setFont('Helvetica-Bold', fs)
        if c.stringWidth(title_str, 'Helvetica-Bold', fs) > BOX_W - 10:
            fs = 7
            c.setFont('Helvetica-Bold', fs)
        tw = c.stringWidth(title_str, 'Helvetica-Bold', fs)
        c.drawString(bx + (BOX_W - tw) / 2, by + items_h + 9, title_str)
        c.restoreState()

        # Itens com quebra de linha
        iy = by + items_h - 4
        for item in items:
            c.saveState()
            c.setFillColor(C_BLUE)
            c.setFont('Helvetica-Bold', 9)
            c.drawString(bx + 7, iy - ITEM_H + 6, '›')
            c.setFillColor(C_GRAY)
            c.setFont('Helvetica', 7.5)
            max_w = BOX_W - 22
            words  = item.split()
            l1, l2 = '', ''
            for w in words:
                test = l1 + (' ' if l1 else '') + w
                if c.stringWidth(test, 'Helvetica', 7.5) <= max_w:
                    l1 = test
                else:
                    test2 = l2 + (' ' if l2 else '') + w
                    if c.stringWidth(test2, 'Helvetica', 7.5) <= max_w:
                        l2 = test2
                    else:
                        l2 = (l2[:-1] + '…') if l2 else w
                        break
            base = iy - ITEM_H + 7 + (4 if not l2 else 0)
            if l1:
                c.drawString(bx + 19, base, l1)
            if l2:
                c.drawString(bx + 19, base - 9, l2)
            c.restoreState()
            iy -= ITEM_H

    # ── Oval central — sombra ──
    c.saveState()
    c.setFillColor(HexColor('#C8D5E8'))
    c.ellipse(cx - ORX + 5, cy - ORY - 5, cx + ORX + 5, cy + ORY - 5, fill=1, stroke=0)
    c.restoreState()

    # Oval principal
    c.saveState()
    c.setFillColor(C_NAVY)
    c.setStrokeColor(C_BLUE)
    c.setLineWidth(3)
    c.ellipse(cx - ORX, cy - ORY, cx + ORX, cy + ORY, fill=1, stroke=1)
    c.restoreState()

    # Texto do título no oval
    c.saveState()
    c.setFillColor(white)
    words = titulo.split() or ['MAPA MENTAL']
    fs = 16
    c.setFont('Helvetica-Bold', fs)
    tw = c.stringWidth(titulo, 'Helvetica-Bold', fs)
    if tw <= ORX * 2 - 24:
        c.drawString(cx - tw / 2, cy - fs * 0.35, titulo)
    else:
        mid = len(words) // 2
        l1, l2 = ' '.join(words[:mid]), ' '.join(words[mid:])
        fs = 13
        c.setFont('Helvetica-Bold', fs)
        while max(c.stringWidth(l1, 'Helvetica-Bold', fs),
                  c.stringWidth(l2, 'Helvetica-Bold', fs)) > ORX * 2 - 20 and fs > 8:
            fs -= 1
            c.setFont('Helvetica-Bold', fs)
        c.drawString(cx - c.stringWidth(l1, 'Helvetica-Bold', fs) / 2, cy + 5, l1)
        c.drawString(cx - c.stringWidth(l2, 'Helvetica-Bold', fs) / 2, cy - fs - 3, l2)
    c.restoreState()

    # ── Cabeçalho escola/professor ──
    parts = []
    if meta.get('escola'): parts.append(meta['escola'].strip())
    if meta.get('professor'): parts.append(f"Prof(a). {meta['professor'].strip()}")
    if parts:
        c.saveState()
        c.setFillColor(HexColor('#6B7280'))
        c.setFont('Helvetica', 7)
        c.drawString(22, H - 20, '  ·  '.join(parts))
        c.restoreState()

    # ── Marca ProfessorIA ──
    c.saveState()
    c.setFillColor(C_NAVY)
    c.setFont('Helvetica-Bold', 8.5)
    brand = 'ProfessorIA™'
    c.drawString(W - c.stringWidth(brand, 'Helvetica-Bold', 8.5) - 18, 16, brand)
    c.restoreState()

    c.save()
    buf.seek(0)
    return buf.read()


def gerar_plano_aula_docx(texto, meta=None, logo_estado_path=None):
    """
    Gera DOCX no formato oficial da Secretaria de Educação Estadual.
    Estrutura: cabeçalho gov + tabela 5 colunas (Aula | Conteúdo | Estratégias | Recursos | Avaliação)

    CONTRATO DE FLUXO:
    - Chamada por: api_chat_download → gerar_docx_pia (quando _detect_doc_type == 'plano_aula')
    - Entrada esperada: texto Markdown com ### AULA N gerado pelo SYSTEM_PROMPT (chat)
    - NÃO usar com JSON do /api/gerar-plano — esse fluxo vai para gerar_plano_pdf()
    """
    import re
    if meta is None:
        meta = {}

    # Proteção: se receber JSON em vez de Markdown, converte para texto legível
    if texto.strip().startswith('{'):
        try:
            d = json.loads(texto)
            linhas = [f"# PLANEJAMENTO DA AULA — {d.get('tema', '')}"]
            linhas.append(f"\n**Habilidades BNCC:** {', '.join(d.get('habilidades_bncc', []))}")
            linhas.append(f"\n**Objetivos:** {' '.join(d.get('objetivos', []))}")
            linhas.append(f"\n### AULA 1 — {d.get('tema', '')}")
            linhas.append(f"\n**Conteúdo e Objetivos de Aprendizagem:**\n{d.get('conteudo_programatico', '')}")
            linhas.append(f"\n**Estratégias Didáticas:**\n{d.get('metodologia', '')}")
            linhas.append(f"\n**Recursos Pedagógicos:**\n{', '.join(d.get('recursos_didaticos', []))}")
            linhas.append(f"\n**Avaliação:**\n{d.get('avaliacao', '')}")
            texto = '\n'.join(linhas)
        except Exception:
            pass

    escola     = meta.get('escola', '').strip()
    professor  = meta.get('professor', '').strip()
    disciplina = meta.get('disciplina', '').strip()
    estado     = meta.get('estado', '').strip()
    # Campos completos do cabecalho escolar
    escola_governo    = meta.get('escola_governo', '').strip()
    escola_secretaria = meta.get('escola_secretaria', '').strip() or 'SECRETARIA DE ESTADO DA EDUCAÇÃO'
    escola_diretoria  = meta.get('escola_diretoria', '').strip()
    escola_endereco   = meta.get('escola_endereco', '').strip()
    escola_fone       = meta.get('escola_fone', '').strip()
    escola_email      = meta.get('escola_email', '').strip()

    meta_extra, aulas = _parse_plano_aula(texto)
    num_aulas  = meta_extra.get('num_aulas', '3 semanais')
    periodo    = meta_extra.get('periodo', 'quinzenal')
    data_range = meta_extra.get('data_range', '')
    serie_turma = meta_extra.get('serie_turma', meta.get('serie', '').strip())

    doc = Document()
    for section in doc.sections:
        section.page_height   = Cm(29.7)
        section.page_width    = Cm(21.0)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)

    # ── CABEÇALHO OFICIAL ──────────────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=3)
    _pia_no_borders(hdr)
    hdr.columns[0].width = Cm(3.0)
    hdr.columns[1].width = Cm(13.5)
    hdr.columns[2].width = Cm(1.5)

    # Coluna esquerda: brasão / logo do governo estadual
    logo_cell = hdr.cell(0, 0)
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(0)
    if logo_estado_path:
        try:
            run = lp.add_run()
            run.add_picture(logo_estado_path, height=Cm(2.6))
        except Exception:
            _pr(lp, '[BRASÃO]', size=7, color='aaaaaa', italic=True)
    else:
        _pr(lp, '🏛', size=20, color='4338ca')

    # Coluna central: info governo + escola
    mid = hdr.cell(0, 1)
    mp = mid.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(0)
    mp.paragraph_format.space_after  = Pt(1)

    # Governo (usa campo personalizado ou fallback)
    if escola_governo:
        estado_txt = escola_governo.upper()
    elif estado:
        estado_txt = f'GOVERNO DO ESTADO DE {estado.upper()}'
    else:
        estado_txt = 'GOVERNO DO ESTADO'
    _pr(mp, estado_txt, bold=True, size=9, color='0a0a0a')

    # Secretaria
    mp2 = mid.add_paragraph()
    mp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp2.paragraph_format.space_before = Pt(0)
    mp2.paragraph_format.space_after  = Pt(0)
    _pr(mp2, escola_secretaria.upper(), bold=True, size=8, color='222222')

    # Diretoria de Ensino
    if escola_diretoria:
        mp_dir = mid.add_paragraph()
        mp_dir.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp_dir.paragraph_format.space_before = Pt(0)
        mp_dir.paragraph_format.space_after  = Pt(0)
        _pr(mp_dir, escola_diretoria.upper(), bold=False, size=8, color='222222')

    # Nome da escola (negrito, maior)
    if escola:
        mp3 = mid.add_paragraph()
        mp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp3.paragraph_format.space_before = Pt(2)
        mp3.paragraph_format.space_after  = Pt(0)
        _pr(mp3, escola.upper(), bold=True, size=10, color='0a0a0a')

    # Endereco da escola
    if escola_endereco:
        mp_end = mid.add_paragraph()
        mp_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp_end.paragraph_format.space_before = Pt(0)
        mp_end.paragraph_format.space_after  = Pt(0)
        _pr(mp_end, escola_endereco, bold=False, size=7, color='444444')

    # Fone e email na mesma linha
    fone_email_parts = []
    if escola_fone:
        fone_email_parts.append(f'Fone {escola_fone}')
    if escola_email:
        fone_email_parts.append(f'Email: {escola_email}')
    if fone_email_parts:
        mp_fe = mid.add_paragraph()
        mp_fe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp_fe.paragraph_format.space_before = Pt(0)
        mp_fe.paragraph_format.space_after  = Pt(0)
        _pr(mp_fe, '  |  '.join(fone_email_parts), bold=False, size=7, color='444444')

    # Coluna direita: watermark ProfessorIA
    rp = hdr.cell(0, 2).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    _pr(rp, 'ProfessorIA™', size=6, color='aaaaaa', italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── TÍTULO E METADADOS ──────────────────────────────────────────────
    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_p.paragraph_format.space_before = Pt(2)
    titulo_p.paragraph_format.space_after  = Pt(4)
    _pr(titulo_p, f'PLANEJAMENTO DA AULA  {datetime.now().year}', bold=True, size=11, color='0a0a0a')

    # Linha de metadados 1: Professor | Componente | Nº aulas
    meta1 = doc.add_table(rows=1, cols=3)
    _pia_no_borders(meta1)
    meta1.columns[0].width = Cm(5.5)
    meta1.columns[1].width = Cm(7.0)
    meta1.columns[2].width = Cm(5.5)

    def _meta_field(cell, label, value):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        _pr(p, f'{label}', bold=True, size=8, color='333333')
        run = p.add_run(value or '________________________________')
        run.font.size = Pt(8)
        run.font.bold = False

    _meta_field(meta1.cell(0, 0), 'Professor(a): ', professor)
    _meta_field(meta1.cell(0, 1), 'Componente Curricular: ', disciplina)
    _meta_field(meta1.cell(0, 2), 'Nº de aulas: ', num_aulas)

    # Linha de metadados 2: Série/Turma | Período | Data
    _pia_hrule(doc, thick=False, color='cccccc')
    meta2 = doc.add_table(rows=1, cols=3)
    _pia_no_borders(meta2)
    meta2.columns[0].width = Cm(5.5)
    meta2.columns[1].width = Cm(7.0)
    meta2.columns[2].width = Cm(5.5)

    _meta_field(meta2.cell(0, 0), 'Ano/Série/Turma: ', serie_turma)
    _meta_field(meta2.cell(0, 1), 'Período do plano: ', periodo)
    _meta_field(meta2.cell(0, 2), 'Data: ', data_range)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_after = Pt(4)

    # ── TABELA PRINCIPAL 5 COLUNAS ──────────────────────────────────────
    COL_HEADERS = [
        'AULA/DATA',
        'CONTEÚDO E OBJETIVOS DE APRENDIZAGEM',
        'ESTRATÉGIAS DIDÁTICAS',
        'RECURSOS PEDAGÓGICOS',
        'AVALIAÇÃO\nVerificar se o objetivo foi cumprido',
    ]
    COL_WIDTHS = [Cm(2.5), Cm(4.5), Cm(4.0), Cm(3.0), Cm(4.0)]

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'

    # Larguras
    for ci, w in enumerate(COL_WIDTHS):
        for row in tbl.rows:
            row.cells[ci].width = w

    # Linha de cabeçalho
    hrow = tbl.rows[0]
    for ci, hdr_txt in enumerate(COL_HEADERS):
        cell = hrow.cells[ci]
        _set_cell_bg_plano(cell, 'e8eaf6')  # azul lavanda claro
        _set_cell_borders_plano(cell, '9fa8da')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        for line in hdr_txt.split('\n'):
            run = p.add_run(line + ('\n' if '\n' in hdr_txt and line == hdr_txt.split('\n')[0] else ''))
            run.font.bold = True
            run.font.size = Pt(7)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)

    # Linhas de conteúdo
    if not aulas:
        # Fallback: sem parser, coloca texto bruto na coluna conteúdo
        aulas = [{'titulo': 'Aula', 'conteudo': texto, 'estrategias': '', 'recursos': '', 'avaliacao': ''}]

    for aula in aulas:
        row = tbl.add_row()
        for ci, w in enumerate(COL_WIDTHS):
            row.cells[ci].width = w

        fields = [
            aula.get('titulo', ''),
            aula.get('conteudo', ''),
            aula.get('estrategias', ''),
            aula.get('recursos', ''),
            aula.get('avaliacao', ''),
        ]
        for ci, txt in enumerate(fields):
            cell = row.cells[ci]
            _set_cell_borders_plano(cell, 'bbbbbb')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            # Clean markdown bold from field text
            txt_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', txt or '').strip()
            run = p.add_run(txt_clean)
            run.font.size = Pt(8)
            run.font.name = 'Arial'
            if ci == 0:
                run.font.bold = True

    # ── RODAPÉ ──────────────────────────────────────────────────────────
    ep2 = doc.add_paragraph()
    ep2.paragraph_format.space_after = Pt(2)
    _pia_hrule(doc, thick=False, color='aaaaaa')
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(0)
    rodape_parts = ['Gerado por ProfessorIA™', datetime.now().strftime('%d/%m/%Y')]
    if escola:
        rodape_parts.append(escola)
    _pr(pf, '  ·  '.join(rodape_parts), size=7, color='888880')

    return doc


def gerar_docx_pia(texto, meta=None, logo_path=None):
    """
    Gera DOCX com design ProfessorIA™ inspirado nas fichas pedagógicas dos exemplos.
    Black & white — imprime bem em qualquer impressora.
    meta dict: escola, professor, disciplina, bimestre, serie
    """
    import re
    if meta is None:
        meta = {}

    # Roteamento por tipo de documento
    doc_type = _detect_doc_type(texto)
    if doc_type == 'plano_aula':
        logo_estado_abs = meta.get('logo_estado_path')
        return gerar_plano_aula_docx(texto, meta=meta, logo_estado_path=logo_estado_abs)
    if doc_type == 'mapa_mental':
        return gerar_mapa_mental_docx(texto, meta=meta)

    escola    = meta.get('escola', '').strip()
    professor = meta.get('professor', '').strip()
    disciplina = meta.get('disciplina', '').strip()
    bimestre   = meta.get('bimestre', '').strip()
    serie      = meta.get('serie', '').strip()

    doc = Document()

    # Margens A4
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    # ── CABEÇALHO ESCOLAR ──────────────────────────────────────────────────
    # Tabela: [LOGO | DADOS DA ESCOLA | ProfessorIA™ watermark]
    hdr = doc.add_table(rows=1, cols=3)
    _pia_no_borders(hdr)
    hdr.columns[0].width = Cm(3.2)
    hdr.columns[1].width = Cm(12.0)
    hdr.columns[2].width = Cm(2.6)

    # Coluna esquerda: logo ou placeholder
    logo_cell = hdr.cell(0, 0)
    logo_cell._tc.get_or_add_tcPr()
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(0)
    if logo_path:
        try:
            run = lp.add_run()
            run.add_picture(logo_path, width=Cm(2.8))
        except Exception:
            _pr(lp, '[LOGO]', size=7, color='aaaaaa', italic=True)
    else:
        # Placeholder box para logo
        t2 = logo_cell.add_table(rows=1, cols=1)
        t2_c = t2.cell(0, 0)
        t2_c._tc.get_or_add_tcPr()
        t2_p = t2_c.paragraphs[0]
        t2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t2_p.paragraph_format.space_before = Pt(14)
        t2_p.paragraph_format.space_after  = Pt(14)
        _pr(t2_p, 'LOGO', size=7, color='aaaaaa', italic=True)
        tc = t2_c._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:color'), 'cccccc')
            tcBorders.append(el)
        tcPr.append(tcBorders)

    # Coluna central: nome da escola, disciplina/bimestre, professor
    mid_cell = hdr.cell(0, 1)
    mp = mid_cell.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(0)
    mp.paragraph_format.space_after  = Pt(2)
    if escola:
        _pr(mp, escola.upper(), bold=True, size=13, color='0a0a0a')
    else:
        _pr(mp, 'ESCOLA / INSTITUIÇÃO DE ENSINO', bold=True, size=11, color='555555')

    # Linha 2: "Avaliação de DISCIPLINA — Xº Bimestre"
    mp2 = mid_cell.add_paragraph()
    mp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp2.paragraph_format.space_before = Pt(1)
    mp2.paragraph_format.space_after  = Pt(1)
    partes = []
    if disciplina:
        partes.append(f'Avaliação de {disciplina}')
    if bimestre:
        partes.append(f'{bimestre}º Bimestre')
    if serie:
        partes.append(serie)
    if partes:
        _pr(mp2, '  ·  '.join(partes), bold=False, size=10, color='222222')
    else:
        _pr(mp2, 'Avaliação', bold=False, size=10, color='555555')

    # Linha 3: professor
    mp3 = mid_cell.add_paragraph()
    mp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp3.paragraph_format.space_before = Pt(1)
    mp3.paragraph_format.space_after  = Pt(0)
    if professor:
        _pr(mp3, f'Prof(a). {professor}', italic=True, size=9, color='444444')
    else:
        _pr(mp3, 'Prof(a). ______________________________', italic=False, size=9, color='888880')

    # Coluna direita: marca ProfessorIA™ (pequena, discreta)
    rc = hdr.cell(0, 2)
    pr_cell = rc.paragraphs[0]
    pr_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr_cell.paragraph_format.space_before = Pt(0)
    pr_cell.paragraph_format.space_after  = Pt(0)
    _pr(pr_cell, 'Professor', bold=True, size=7, color='bbbbbb')
    _pr(pr_cell, 'IA', bold=True, size=7, color='bbbbbb')
    ps2 = rc.add_paragraph()
    ps2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ps2.paragraph_format.space_before = Pt(0)
    ps2.paragraph_format.space_after  = Pt(0)
    _pr(ps2, '™', size=6, color='cccccc')

    # Linha separadora grossa
    _pia_hrule(doc, thick=True)

    # ── CAMPOS DO ALUNO ────────────────────────────────────────────────────
    # Formato: Nome:___ Data:__/__/__ Ano:
    pn = doc.add_paragraph()
    pn.paragraph_format.space_before = Pt(4)
    pn.paragraph_format.space_after  = Pt(4)
    _pr(pn, 'Nome:', bold=True, size=10, color='0a0a0a')
    _pr(pn, '_' * 44, size=10, color='555550')
    _pr(pn, '   Data:', bold=True, size=10, color='0a0a0a')
    _pr(pn, '__/__/__', size=10, color='555550')
    _pr(pn, '   Ano:', bold=True, size=10, color='0a0a0a')
    _pr(pn, '_______________', size=10, color='555550')

    _pia_hrule(doc, thick=False, color='555555')

    # ── CONTEÚDO MARKDOWN ─────────────────────────────────────────────────
    lines = texto.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Bloco de código (``` ... ```)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            if _is_letter_grid(code_text):
                _pia_caca_palavras_table(doc, code_text)
            else:
                _pia_code_block(doc, code_text)
            i += 1
            continue

        # Tabela markdown
        if line.strip().startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            _pia_md_table(doc, tbl_lines)
            continue

        # H1 → título grande centrado
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(10)
            _pr(p, line[2:].strip().upper(), bold=True, size=20, color='0a0a0a')

        # H2 → caixa preta de seção
        elif line.startswith('## '):
            _pia_section_box(doc, line[3:].strip())

        # H3 → sub-seção negrito com linha fina
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            _pr(p, line[4:].strip().upper(), bold=True, size=10, color='0a0a0a')
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
            bot.set(qn('w:space'), '1');   bot.set(qn('w:color'), 'aaaaaa')
            pBdr.append(bot); pPr.append(pBdr)

        # Lista com marcador
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            _pr_fmt(p, line[2:].strip())

        # Lista numerada
        elif re.match(r'^\d+[\.\)\:]\s', line):
            m = re.match(r'^\d+[\.\)\:]\s+(.*)', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _pr_fmt(p, m.group(1) if m else line)

        # Separador ou linha vazia
        elif line.strip() == '' or re.match(r'^-{3,}$', line.strip()):
            ep = doc.add_paragraph()
            ep.paragraph_format.space_after = Pt(3)

        # Texto normal
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            _pr_fmt(p, line)

        i += 1

    # ── RODAPÉ ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _pia_hrule(doc, thick=False, color='aaaaaa')
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(0)
    rodape_parts = ['Gerado por ProfessorIA™', datetime.now().strftime('%d/%m/%Y')]
    if escola:
        rodape_parts.append(escola)
    _pr(pf, '  ·  '.join(rodape_parts), size=7, color='888880')

    return doc


@app.route('/api/chat-download', methods=['POST'])
@login_required
def api_chat_download():
    """Converte o texto de uma mensagem do chat em DOCX com design ProfessorIA."""
    import os, traceback
    data = request.json or {}
    texto = data.get('texto', '').strip()
    if not texto:
        return jsonify({'erro': 'Texto vazio'}), 400

    meta = {
        'escola':    data.get('escola', current_user.escola_nome).strip(),
        'professor': data.get('professor', current_user.professor_nome).strip(),
        'disciplina': data.get('disciplina', '').strip(),
        'bimestre':   data.get('bimestre', '').strip(),
        'serie':      data.get('serie', '').strip(),
        'estado':     data.get('estado', '').strip(),
        # Campos completos da escola (preenchidos na pagina Minha Conta)
        'escola_governo':    data.get('escola_governo', getattr(current_user, 'escola_governo', '')).strip(),
        'escola_secretaria': data.get('escola_secretaria', getattr(current_user, 'escola_secretaria', '')).strip(),
        'escola_diretoria':  data.get('escola_diretoria', getattr(current_user, 'escola_diretoria', '')).strip(),
        'escola_endereco':   data.get('escola_endereco', getattr(current_user, 'escola_endereco', '')).strip(),
        'escola_fone':       data.get('escola_fone', getattr(current_user, 'escola_fone', '')).strip(),
        'escola_email':      data.get('escola_email', getattr(current_user, 'escola_email', '')).strip(),
    }

    logo_abs = None
    if current_user.logo_path:
        candidate = os.path.join(os.path.dirname(__file__), current_user.logo_path)
        if os.path.isfile(candidate):
            logo_abs = candidate

    # Logo do governo estadual (brasão)
    logo_estado_abs = None
    if current_user.logo_estado_path:
        candidate_e = os.path.join(os.path.dirname(__file__), current_user.logo_estado_path)
        if os.path.isfile(candidate_e):
            logo_estado_abs = candidate_e
    meta['logo_estado_path'] = logo_estado_abs

    try:
        doc_type = _detect_doc_type(texto)
        if doc_type == 'mapa_mental':
            doc = gerar_mapa_mental_docx(texto, meta=meta)
            download_name = 'mapa-mental-ProfessorIA.docx'
        elif doc_type == 'plano_aula':
            doc = gerar_plano_aula_docx(texto, meta=meta, logo_estado_path=logo_estado_abs)
            download_name = 'plano-de-aula-ProfessorIA.docx'
        else:
            doc = gerar_docx_pia(texto, meta=meta, logo_path=logo_abs)
            download_name = 'material-ProfessorIA.docx'
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf, as_attachment=True,
            download_name=download_name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error('Erro no chat-download: %s', traceback.format_exc())
        return jsonify({'erro': f'Falha ao gerar arquivo: {str(e)[:200]}'}), 500


@app.route('/api/config-escola', methods=['GET', 'POST'])
@login_required
def api_config_escola():
    """Salva ou retorna configurações da escola do professor."""
    if request.method == 'GET':
        return jsonify({
            'escola_nome':      current_user.escola_nome,
            'professor_nome':   current_user.professor_nome,
            'logo_path':        current_user.logo_path,
            'logo_estado_path': current_user.logo_estado_path,
            'escola_governo':    getattr(current_user, 'escola_governo', ''),
            'escola_secretaria': getattr(current_user, 'escola_secretaria', ''),
            'escola_diretoria':  getattr(current_user, 'escola_diretoria', ''),
            'escola_endereco':   getattr(current_user, 'escola_endereco', ''),
            'escola_fone':       getattr(current_user, 'escola_fone', ''),
            'escola_email':      getattr(current_user, 'escola_email', ''),
        })
    data = request.json or {}
    escola   = data.get('escola_nome', '').strip()[:200]
    prof     = data.get('professor_nome', '').strip()[:200]
    # Campos completos do cabecalho escolar
    gov      = data.get('escola_governo', '').strip()[:200]
    sec      = data.get('escola_secretaria', '').strip()[:200]
    dire     = data.get('escola_diretoria', '').strip()[:200]
    ender    = data.get('escola_endereco', '').strip()[:300]
    fone     = data.get('escola_fone', '').strip()[:50]
    email_e  = data.get('escola_email', '').strip()[:200]
    conn = get_db()
    conn.execute(
        "UPDATE usuarios SET escola_nome=?, professor_nome=?,"
        " escola_governo=?, escola_secretaria=?, escola_diretoria=?,"
        " escola_endereco=?, escola_fone=?, escola_email=?"
        " WHERE id=?",
        (escola, prof, gov, sec, dire, ender, fone, email_e, current_user.id)
    )
    conn.commit(); conn.close()
    return jsonify({'ok': True})


_LOGO_EXTS    = {'.png', '.jpg', '.jpeg', '.webp'}
_LOGO_MAX_BYTES = 2 * 1024 * 1024  # 2 MB
_LOGO_CAMPOS_PERMITIDOS = {'logo_path', 'logo_estado_path'}

def _salvar_logo(f, prefixo, campo_db):
    """Valida, salva e atualiza o logo/brasão do usuário. Retorna (rel_path, erro_msg)."""
    import uuid
    if campo_db not in _LOGO_CAMPOS_PERMITIDOS:
        return None, 'Campo inválido'
    if not f or not f.filename:
        return None, 'Nenhum arquivo enviado'
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in _LOGO_EXTS:
        return None, 'Formato não suportado. Use PNG, JPG ou WEBP.'
    data = f.read()
    if len(data) > _LOGO_MAX_BYTES:
        return None, 'Arquivo muito grande. Máximo 2 MB.'
    fname    = f'{prefixo}_{current_user.id}_{uuid.uuid4().hex[:8]}{ext}'
    save_dir = os.path.join(os.path.dirname(__file__), 'static', 'logos')
    os.makedirs(save_dir, exist_ok=True)
    fpath = os.path.join(save_dir, fname)
    with open(fpath, 'wb') as fp:
        fp.write(data)
    rel = f'static/logos/{fname}'
    conn = get_db()
    conn.execute(f'UPDATE usuarios SET {campo_db}=? WHERE id=?', (rel, current_user.id))
    conn.commit(); conn.close()
    return rel, None


@app.route('/api/upload-logo', methods=['POST'])
@login_required
def api_upload_logo():
    rel, erro = _salvar_logo(request.files.get('logo'), 'logo', 'logo_path')
    if erro:
        return jsonify({'erro': erro}), 400
    return jsonify({'ok': True, 'logo_path': rel})


@app.route('/api/upload-logo-estado', methods=['POST'])
@login_required
def api_upload_logo_estado():
    rel, erro = _salvar_logo(request.files.get('logo'), 'brasao', 'logo_estado_path')
    if erro:
        return jsonify({'erro': erro}), 400
    return jsonify({'ok': True, 'logo_estado_path': rel})


# ─── DOCX via templates oficiais (docxtpl) ───────────────────────────────────

PLANO_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), 'static', 'templates', 'plano_de_aula.docx'
)
ENSINO_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), 'static', 'templates', 'plano_de_ensino_tpl.docx'
)

# Ferramenta Claude — gera JSON com as chaves exatas do template PLANO DE AULA.docx
PLANO_AULA_DOCX_TOOL = {
    "name": "gerar_plano_aula_docx",
    "description": (
        "Gera os dados para preencher o template oficial PLANO DE AULA.docx. "
        "Retorne as chaves exatas que o template usa: disciplina, num_aulas, turma e aulas[]."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "disciplina": {"type": "string", "description": "Nome da disciplina"},
            "num_aulas":  {"type": "string", "description": "Quantidade e duração (ex: '2 aulas de 50 min')"},
            "turma":      {"type": "string", "description": "Ano, série e turma (ex: '8º Ano B')"},
            "aulas": {
                "type": "array",
                "description": "Cada linha da tabela de desenvolvimento",
                "items": {
                    "type": "object",
                    "properties": {
                        "data":        {"type": "string", "description": "Data da aula (pode ser vazio)"},
                        "conteudo":    {"type": "string", "description": "Conteúdo/tema com habilidades BNCC"},
                        "estrategias": {"type": "string", "description": "Estratégias didáticas e metodologia"},
                        "recursos":    {"type": "string", "description": "Recursos pedagógicos"},
                        "avaliacao":   {"type": "string", "description": "Método de avaliação"},
                        "verificacao": {"type": "string", "description": "Critérios de verificação da aprendizagem"},
                    },
                    "required": ["data", "conteudo", "estrategias", "recursos", "avaliacao", "verificacao"]
                }
            }
        },
        "required": ["disciplina", "num_aulas", "turma", "aulas"]
    }
}

# Ferramenta Claude — gera JSON com as chaves exatas do template PLANO DE ENSINO.docx
PLANO_ENSINO_DOCX_TOOL = {
    "name": "gerar_plano_ensino_docx",
    "description": (
        "Gera os dados para preencher o template oficial PLANO DE ENSINO.docx. "
        "Retorne as chaves exatas que o template usa."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "disciplina":  {"type": "string"},
            "num_aulas":   {"type": "string", "description": "Nº de aulas semanais (ex: '2')"},
            "turma":       {"type": "string", "description": "Ano/Série/Turma (ex: '8º Ano B')"},
            "objetivos":   {"type": "string", "description": "Objetivos gerais da disciplina no ano"},
            "obj_b1":      {"type": "string", "description": "Objetos de conhecimento — 1º bimestre"},
            "obj_b2":      {"type": "string", "description": "Objetos de conhecimento — 2º bimestre"},
            "obj_b3":      {"type": "string", "description": "Objetos de conhecimento — 3º bimestre"},
            "obj_b4":      {"type": "string", "description": "Objetos de conhecimento — 4º bimestre"},
            "comp_b1":     {"type": "string", "description": "Competências e habilidades BNCC — 1º bimestre"},
            "comp_b2":     {"type": "string", "description": "Competências e habilidades BNCC — 2º bimestre"},
            "comp_b3":     {"type": "string", "description": "Competências e habilidades BNCC — 3º bimestre"},
            "comp_b4":     {"type": "string", "description": "Competências e habilidades BNCC — 4º bimestre"},
            "metodologias":                 {"type": "string"},
            "competencias_socioemocionais": {"type": "string"},
            "programas_projetos":           {"type": "string"},
            "materiais_apoio":              {"type": "string"},
            "avaliacao":                    {"type": "string"},
            "recuperacao":                  {"type": "string"},
            "diagnostico":                  {"type": "string"},
            "referencias":                  {"type": "string"},
        },
        "required": [
            "disciplina", "num_aulas", "turma", "objetivos",
            "obj_b1", "obj_b2", "obj_b3", "obj_b4",
            "comp_b1", "comp_b2", "comp_b3", "comp_b4",
            "metodologias", "competencias_socioemocionais",
            "programas_projetos", "materiais_apoio",
            "avaliacao", "recuperacao", "diagnostico", "referencias"
        ]
    }
}


def _renderizar_docx_tpl(template_path: str, context: dict) -> bytes:
    """Renderiza qualquer template docxtpl com o contexto fornecido."""
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()


def _gerar_plano_aula_docx_interno(tema: str, ano: str, disciplina: str) -> dict:
    """Chama a IA para gerar o JSON com as chaves exatas do template PLANO DE AULA.docx."""
    user_prompt = (
        f"Gere um plano de aula completo para:\n"
        f"- Tema: {tema}\n- Ano/Série: {ano}\n- Disciplina: {disciplina}\n\n"
        "Use habilidades BNCC reais. Inclua pelo menos 3 aulas com conteúdo detalhado, "
        "estratégias ativas e recursos variados."
    )
    if os.environ.get('ANTHROPIC_API_KEY'):
        try:
            resp = client.messages.create(
                model='claude-sonnet-4-6', max_tokens=4000,
                system=SYSTEM_PROMPT_PLANO,
                tools=[PLANO_AULA_DOCX_TOOL],
                tool_choice={"type": "tool", "name": "gerar_plano_aula_docx"},
                messages=[{"role": "user", "content": user_prompt}]
            )
            for block in resp.content:
                if block.type == 'tool_use' and block.name == 'gerar_plano_aula_docx':
                    return block.input
        except Exception as e:
            logger.warning('_gerar_plano_aula_docx_interno Claude falhou: %s', e)

    # Fallback: parseia JSON do texto
    texto = _llm_cadeia_simples(user_prompt, sistema=SYSTEM_PROMPT_PLANO)
    m = re.search(r'\{[\s\S]+\}', texto)
    if m:
        return json.loads(m.group())
    raise ValueError('Não foi possível parsear JSON do plano de aula')


def _gerar_plano_ensino_docx_interno(disciplina: str, ano: str) -> dict:
    """Chama a IA para gerar o JSON com as chaves exatas do template PLANO DE ENSINO.docx."""
    user_prompt = (
        f"Gere um plano de ensino anual completo para:\n"
        f"- Disciplina: {disciplina}\n- Ano/Série: {ano}\n\n"
        "Use habilidades BNCC reais e distribua os objetos de conhecimento nos 4 bimestres."
    )
    if os.environ.get('ANTHROPIC_API_KEY'):
        try:
            resp = client.messages.create(
                model='claude-sonnet-4-6', max_tokens=4000,
                system=SYSTEM_PROMPT_PLANO,
                tools=[PLANO_ENSINO_DOCX_TOOL],
                tool_choice={"type": "tool", "name": "gerar_plano_ensino_docx"},
                messages=[{"role": "user", "content": user_prompt}]
            )
            for block in resp.content:
                if block.type == 'tool_use' and block.name == 'gerar_plano_ensino_docx':
                    return block.input
        except Exception as e:
            logger.warning('_gerar_plano_ensino_docx_interno Claude falhou: %s', e)

    texto = _llm_cadeia_simples(user_prompt, sistema=SYSTEM_PROMPT_PLANO)
    m = re.search(r'\{[\s\S]+\}', texto)
    if m:
        return json.loads(m.group())
    raise ValueError('Não foi possível parsear JSON do plano de ensino')


@app.route('/api/plano-aula/docx', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_plano_aula_docx():
    """Preenche o template oficial PLANO DE AULA.docx via docxtpl.

    Aceita dois formatos:
      A) { "plano_docx": { disciplina, num_aulas, turma, aulas[] } }  — JSON já no formato do template
      B) { "tema", "ano", "disciplina" }  — gera o plano via IA com as chaves exatas do template
    """
    if not current_user.assinatura_ativa and not current_user.is_admin:
        if get_geracoes_mes(current_user.id) >= LIMITE_GRATIS:
            return jsonify({'erro': 'limite_atingido', 'cta': '/planos'}), 403

    data = request.get_json(force=True) or {}
    context = data.get('plano_docx')

    if not context:
        tema       = str(data.get('tema', '')).strip()[:300]
        ano        = str(data.get('ano', '')).strip()[:50]
        disciplina = str(data.get('disciplina', '')).strip()[:100]
        if not tema or not ano or not disciplina:
            return jsonify({'erro': 'Forneça plano_docx ou tema + ano + disciplina'}), 400
        try:
            context = _gerar_plano_aula_docx_interno(tema, ano, disciplina)
        except Exception as e:
            return jsonify({'erro': f'Falha ao gerar plano: {str(e)[:200]}'}), 500

    if not context:
        return jsonify({'erro': 'Dados do plano não encontrados'}), 400
    if not os.path.exists(PLANO_TEMPLATE_PATH):
        return jsonify({'erro': 'Template DOCX não encontrado no servidor.'}), 500

    # Adiciona campos gerados pelo servidor
    context['professor'] = current_user.professor_nome or ''
    context['data_plano'] = datetime.now().strftime('%d/%m/%Y')

    try:
        docx_bytes = _renderizar_docx_tpl(PLANO_TEMPLATE_PATH, context)
        slug = str(context.get('disciplina', 'plano'))[:40].replace(' ', '_')
        return send_file(
            io.BytesIO(docx_bytes),
            as_attachment=True,
            download_name=f'PlanoDeAula_{slug}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error('api_plano_aula_docx erro: %s', traceback.format_exc())
        return jsonify({'erro': f'Erro ao preencher template: {str(e)[:200]}'}), 500


@app.route('/api/plano-ensino/docx', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_plano_ensino_docx():
    """Preenche o template oficial PLANO DE ENSINO.docx via docxtpl.

    Aceita dois formatos:
      A) { "plano_docx": { disciplina, num_aulas, turma, objetivos, obj_b1..b4, comp_b1..b4, ... } }
      B) { "disciplina", "ano" }  — gera via IA com as chaves exatas do template
    """
    if not current_user.assinatura_ativa and not current_user.is_admin:
        return jsonify({'erro': 'Plano necessário', 'cta': '/planos'}), 403

    data = request.get_json(force=True) or {}
    context = data.get('plano_docx')

    if not context:
        disciplina = str(data.get('disciplina', '')).strip()[:100]
        ano        = str(data.get('ano', '')).strip()[:50]
        if not disciplina or not ano:
            return jsonify({'erro': 'Forneça plano_docx ou disciplina + ano'}), 400
        try:
            context = _gerar_plano_ensino_docx_interno(disciplina, ano)
        except Exception as e:
            return jsonify({'erro': f'Falha ao gerar plano de ensino: {str(e)[:200]}'}), 500

    if not context:
        return jsonify({'erro': 'Dados do plano de ensino não encontrados'}), 400
    if not os.path.exists(ENSINO_TEMPLATE_PATH):
        return jsonify({'erro': 'Template PLANO DE ENSINO não encontrado no servidor.'}), 500

    context['professor'] = current_user.professor_nome or ''
    context['data_plano'] = datetime.now().strftime('%d/%m/%Y')

    try:
        docx_bytes = _renderizar_docx_tpl(ENSINO_TEMPLATE_PATH, context)
        slug = str(context.get('disciplina', 'ensino'))[:40].replace(' ', '_')
        return send_file(
            io.BytesIO(docx_bytes),
            as_attachment=True,
            download_name=f'PlanoDeEnsino_{slug}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error('api_plano_ensino_docx erro: %s', traceback.format_exc())
        return jsonify({'erro': f'Erro ao preencher template: {str(e)[:200]}'}), 500


# ─── Máquina de Leads — Degustação ───────────────────────────────────────────

@app.route('/teste-gratis')
def teste_gratis():
    return render_template('teste_gratis.html')


@app.route('/api/degustacao', methods=['POST'])
@limiter.limit('20 per minute')
def api_degustacao():
    """Captura lead e retorna 1 aula de degustação via OpenAI gpt-4o-mini."""
    data = request.get_json(force=True) or {}
    nome    = str(data.get('nome', '')).strip()[:100]
    contato = str(data.get('contato', '')).strip()[:200]
    tema    = str(data.get('tema', '')).strip()[:300]
    serie   = str(data.get('serie', '')).strip()[:50]

    if not nome or not contato or not tema:
        return jsonify({'erro': 'nome, contato e tema são obrigatórios'}), 400

    # Salva o lead no banco
    try:
        conn = get_db()
        conn.execute(
            'INSERT INTO leads (nome, contato, tema_pesquisado, criado_em) VALUES (?, ?, ?, ?)',
            (nome, contato, tema, datetime.now().isoformat())
        )
        conn.commit()
        conn.close()
    except Exception as e:
        logger.warning('api_degustacao — erro ao salvar lead: %s', e)

    # Gera 1 aula via OpenAI gpt-4o-mini (rápido e barato)
    if not client_openai:
        return jsonify({'erro': 'Serviço temporariamente indisponível. Tente novamente em breve.'}), 503

    prompt = (
        f"Você é um assistente pedagógico especialista na BNCC brasileira.\n"
        f"Gere APENAS UMA linha de plano de aula para o tema '{tema}' no {serie or '8º ano'}.\n"
        "Retorne EXCLUSIVAMENTE um JSON válido com as chaves:\n"
        '{"conteudo": "...", "estrategias": "...", "recursos": "..."}\n'
        "conteudo: título + habilidade BNCC real (ex: EF08HI03)\n"
        "estrategias: metodologia ativa em 1 frase\n"
        "recursos: lista de recursos em 1 frase\n"
        "Nenhum texto adicional fora do JSON."
    )
    try:
        resp = client_openai.chat.completions.create(
            model='gpt-4o-mini',
            messages=[{'role': 'user', 'content': prompt}],
            max_tokens=300,
            temperature=0.7,
            response_format={'type': 'json_object'},
        )
        aula = json.loads(resp.choices[0].message.content)
        return jsonify({'ok': True, 'aula': aula, 'tema': tema})
    except Exception as e:
        logger.error('api_degustacao — erro IA: %s', e)
        return jsonify({'erro': 'Erro ao gerar prévia. Tente novamente.'}), 500


# ─── Gerador de Provas ────────────────────────────────────────────────────────

@app.route('/prova')
@login_required
def prova_page():
    geracoes = get_geracoes_mes(current_user.id)
    tem_plano = current_user.assinatura_ativa or current_user.is_admin
    limite_atingido = not tem_plano and geracoes >= LIMITE_GRATIS
    return render_template('prova.html',
                           geracoes=geracoes,
                           limite=LIMITE_GRATIS,
                           limite_atingido=limite_atingido,
                           tem_plano=tem_plano)


# ─── Planejamento Anual ────────────────────────────────────────────────────────

@app.route('/planejamento')
@login_required
def planejamento():
    if not current_user.assinatura_ativa and not current_user.is_admin:
        return redirect(url_for('chat'))
    return render_template('planejamento.html')


@app.route('/api/planejamento', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_planejamento():
    if not current_user.assinatura_ativa and not current_user.is_admin:
        return jsonify({'erro': 'Plano necessário'}), 403

    data = request.get_json(silent=True) or {}
    disciplina    = data.get('disciplina', '').strip()
    turma         = data.get('turma', '').strip()
    if not disciplina or not turma:
        return jsonify({'erro': 'Disciplina e turma são obrigatórios'}), 400
    ano           = data.get('ano', str(datetime.now().year))
    aulas_semana  = int(data.get('aulas_semana', 2))
    inicio        = data.get('inicio', f'01/02/{ano}')
    fim           = data.get('fim',   f'30/11/{ano}')

    prompt = f"""Crie um planejamento anual completo para professor brasileiro.

Dados:
- Disciplina: {disciplina}
- Turma/Série: {turma}
- Ano letivo: {ano}
- Aulas por semana: {aulas_semana}
- Período: {inicio} até {fim}

Gere um planejamento bimestral detalhado seguindo a BNCC com:
- Divisão por bimestre (4 bimestres)
- Conteúdos de cada bimestre
- Habilidades trabalhadas (códigos BNCC quando aplicável)
- Sugestão de avaliações

Formato: texto estruturado e claro, pronto para entregar à coordenação."""

    try:
        conteudo = chamar_ia_simples(prompt)
    except Exception as e:
        return jsonify({'erro': f'Erro ao gerar planejamento: {str(e)[:200]}'}), 500

    conn = get_db()
    conn.execute(
        "INSERT INTO planejamento_anual (usuario_id, disciplina, turma, ano, conteudo, criado_em) VALUES (?, ?, ?, ?, ?, ?)",
        (current_user.id, disciplina, turma, ano, conteudo,
         datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    )
    conn.commit()
    conn.close()

    return jsonify({'conteudo': conteudo})


# ─── Termos e Privacidade ─────────────────────────────────────────────────────

@app.route('/termos')
def termos():
    return render_template('termos.html')

@app.route('/privacidade')
def privacidade():
    return render_template('privacidade.html')

# ─── Error handlers ────────────────────────────────────────────────────────────

@app.errorhandler(404)
def nao_encontrado(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def erro_interno(e):
    return render_template('500.html'), 500

# ═══════════════════════════════════════════════════════════════════════════════
# BANCO DE QUESTÕES
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/banco-questoes')
@login_required
def banco_questoes():
    return render_template('banco_questoes.html')

@app.route('/api/questoes', methods=['GET'])
@login_required
def api_listar_questoes():
    conn = get_db()
    disciplina = request.args.get('disciplina', '')
    ano_serie  = request.args.get('ano_serie', '')
    busca      = request.args.get('busca', '')
    sql = "SELECT * FROM questions_bank WHERE usuario_id = ?"
    params = [current_user.id]
    if disciplina:
        sql += " AND disciplina = ?"; params.append(disciplina)
    if ano_serie:
        sql += " AND (ano_serie = ? OR serie = ?)"; params += [ano_serie, ano_serie]
    if busca:
        sql += " AND (enunciado ILIKE ? OR habilidade_bncc ILIKE ?)"
        params += [f'%{busca}%', f'%{busca}%']
    sql += " ORDER BY id DESC LIMIT 200"
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/questoes', methods=['POST'])
@app.route('/api/questions/save', methods=['POST']) # Alias solicitado no Super Prompt
@login_required
def api_salvar_questao():
    d = request.get_json(force=True)
    if not d.get('enunciado'):
        return jsonify({'erro': 'Enunciado obrigatório'}), 400
    
    # Mapeamento de campos para suportar o novo Super Prompt
    enunciado = d.get('enunciado', '')
    alternativas = d.get('alternativas', [])
    gabarito = d.get('gabarito', d.get('resposta_correta', ''))
    bncc = d.get('bncc_skill', d.get('habilidade_bncc', ''))
    disciplina = d.get('disciplina', '')
    serie = d.get('serie', d.get('ano_serie', ''))
    tipo = d.get('tipo', 'multipla_escolha')

    conn = get_db()
    conn.execute(
        """INSERT INTO questions_bank
           (usuario_id, enunciado, alternativas, gabarito, ano_serie, disciplina, habilidade_bncc, tipo, criado_em)
           VALUES (?,?,?,?,?,?,?,?,?)""",
        (current_user.id,
         enunciado,
         json.dumps(alternativas, ensure_ascii=False),
         gabarito,
         serie,
         disciplina,
         bncc,
         tipo,
         datetime.now().strftime('%d/%m/%Y %H:%M'))
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/questoes/<int:qid>', methods=['DELETE'])
@login_required
def api_deletar_questao(qid):
    conn = get_db()
    conn.execute("DELETE FROM questions_bank WHERE id = ? AND usuario_id = ?", (qid, current_user.id))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/questoes/extrair-do-chat', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_extrair_questoes():
    """Extrai questões estruturadas de um texto gerado pela IA e salva no banco."""
    d = request.get_json(force=True)
    texto = d.get('texto', '')
    disciplina = d.get('disciplina', '')
    ano_serie  = d.get('ano_serie', '')
    if not texto:
        return jsonify({'erro': 'Texto vazio'}), 400
    prompt = f"""Analise o texto abaixo e extraia TODAS as questões de múltipla escolha ou discursivas.
Para cada questão, retorne um JSON array com objetos contendo:
- enunciado: texto da questão
- alternativas: array de strings ["A) ...", "B) ...", ...] (vazio se discursiva)
- gabarito: letra ou resposta correta
- tipo: "multipla_escolha" ou "discursiva"
- habilidade_bncc: código BNCC se mencionado (ex: EF05MA01), senão ""

Retorne APENAS o JSON array, sem texto adicional.

TEXTO:
{texto[:4000]}"""
    try:
        raw = chamar_ia_simples(prompt).strip()
        # Remove markdown code fences if present
        raw = re.sub(r'^```[a-z]*\n?', '', raw)
        raw = re.sub(r'\n?```$', '', raw)
        questoes = json.loads(raw)
        conn = get_db()
        salvos = 0
        for q in questoes:
            if q.get('enunciado'):
                conn.execute(
                    """INSERT INTO questions_bank
                       (usuario_id, enunciado, alternativas, gabarito, ano_serie, disciplina, habilidade_bncc, tipo, criado_em)
                       VALUES (?,?,?,?,?,?,?,?,?)""",
                    (current_user.id,
                     q.get('enunciado',''),
                     json.dumps(q.get('alternativas',[]), ensure_ascii=False),
                     q.get('gabarito',''),
                     ano_serie,
                     disciplina,
                     q.get('habilidade_bncc',''),
                     q.get('tipo','multipla_escolha'),
                     datetime.now().strftime('%d/%m/%Y %H:%M'))
                )
                salvos += 1
        conn.commit()
        conn.close()
        return jsonify({'ok': True, 'salvos': salvos})
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

# ═══════════════════════════════════════════════════════════════════════════════
# DASHBOARD DE ANALYTICS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/dashboard')
@login_required
def dashboard():
    conn = get_db()
    total_materiais = conn.execute(
        "SELECT COUNT(*) as c FROM historico WHERE usuario_id = ?", (current_user.id,)
    ).fetchone()['c']
    total_questoes = conn.execute(
        "SELECT COUNT(*) as c FROM questions_bank WHERE usuario_id = ?", (current_user.id,)
    ).fetchone()['c']
    total_chat = conn.execute(
        "SELECT COUNT(*) as c FROM chat_messages WHERE usuario_id = ? AND role = 'user'", (current_user.id,)
    ).fetchone()['c']
    por_disciplina = conn.execute(
        """SELECT disciplina, COUNT(*) as total FROM historico
           WHERE usuario_id = ? AND disciplina != ''
           GROUP BY disciplina ORDER BY total DESC LIMIT 5""",
        (current_user.id,)
    ).fetchall()
    por_mes = conn.execute(
        """SELECT SUBSTRING(data, 4, 7) as mes, COUNT(*) as total
           FROM historico WHERE usuario_id = ? AND data != ''
           GROUP BY mes ORDER BY mes DESC LIMIT 6""",
        (current_user.id,)
    ).fetchall()
    ultimas = conn.execute(
        """SELECT id, data, disciplina, turma, num_aulas, nome_arquivo
           FROM historico WHERE usuario_id = ?
           ORDER BY id DESC LIMIT 5""",
        (current_user.id,)
    ).fetchall()
    conn.close()
    horas_economizadas = round(total_materiais * 2.5, 1)
    return render_template('dashboard.html',
        total_materiais=total_materiais,
        total_questoes=total_questoes,
        total_chat=total_chat,
        horas_economizadas=horas_economizadas,
        por_disciplina=[dict(r) for r in por_disciplina],
        por_mes=[dict(r) for r in reversed(list(por_mes))],
        ultimas=[dict(r) for r in ultimas]
    )

@app.route('/api/dashboard-stats')
@login_required
def api_dashboard_stats():
    dias = int(request.args.get('dias', 30))
    conn = get_db()
    total_materiais = conn.execute(
        "SELECT COUNT(*) as c FROM historico WHERE usuario_id = ?", (current_user.id,)
    ).fetchone()['c']
    total_aulas = conn.execute(
        "SELECT COALESCE(SUM(num_aulas),0) as s FROM historico WHERE usuario_id = ?", (current_user.id,)
    ).fetchone()['s']
    total_questoes = conn.execute(
        "SELECT COUNT(*) as c FROM questions_bank WHERE usuario_id = ?", (current_user.id,)
    ).fetchone()['c']
    por_semana_rows = conn.execute(
        """SELECT TO_CHAR(TO_DATE(data,'DD/MM/YYYY'),'IYYY-IW') as semana, COUNT(*) as total
           FROM historico WHERE usuario_id = ? AND data != ''
           GROUP BY semana ORDER BY semana DESC LIMIT 8""",
        (current_user.id,)
    ).fetchall()
    semanas = [r['semana'] or '' for r in reversed(list(por_semana_rows))]
    por_semana_vals = [r['total'] for r in reversed(list(por_semana_rows))]
    # Pad to 8
    while len(semanas) < 8:
        semanas.insert(0, '')
        por_semana_vals.insert(0, 0)
    # Tipo distribution
    por_tipo = {'plano_aula': total_materiais, 'questao': int(total_questoes)}
    # Recent
    recentes = conn.execute(
        """SELECT id, data, disciplina, turma, num_aulas FROM historico
           WHERE usuario_id = ? ORDER BY id DESC LIMIT 8""",
        (current_user.id,)
    ).fetchall()
    conn.close()
    return jsonify({
        'total_materiais': total_materiais,
        'total_aulas': int(total_aulas or 0),
        'total_questoes': int(total_questoes),
        'delta_materiais': 0,
        'delta_aulas': 0,
        'semanas': [s[-2:] if s else 'S?' for s in semanas],
        'por_semana': por_semana_vals,
        'por_tipo': por_tipo,
        'recentes': [{'titulo': f"{r['disciplina']} — {r['turma']}", 'data': r['data'], 'tipo': 'plano_aula'} for r in recentes]
    })

# ═══════════════════════════════════════════════════════════════════════════════
# PROGRAMA DE INDICAÇÃO (REFERRAL)
# ═══════════════════════════════════════════════════════════════════════════════

def _get_or_create_referral(usuario_id):
    conn = get_db()
    row = conn.execute(
        "SELECT * FROM referrals WHERE usuario_id = ?", (usuario_id,)
    ).fetchone()
    if not row:
        codigo = secrets.token_urlsafe(8).upper()[:10]
        conn.execute(
            "INSERT INTO referrals (usuario_id, codigo, usos, creditos, criado_em) VALUES (?,?,0,0,?)",
            (usuario_id, codigo, datetime.now().strftime('%d/%m/%Y'))
        )
        conn.commit()
        row = conn.execute(
            "SELECT * FROM referrals WHERE usuario_id = ?", (usuario_id,)
        ).fetchone()
    conn.close()
    return dict(row)

@app.route('/indicar')
@login_required
def indicar():
    ref = _get_or_create_referral(current_user.id)
    link = f"{SITE_URL}/cadastro?ref={ref['codigo']}"
    return render_template('indicar.html', ref=ref, link=link)

@app.route('/api/referral/stats')
@login_required
def api_referral_stats():
    ref = _get_or_create_referral(current_user.id)
    link = f"{SITE_URL}/cadastro?ref={ref['codigo']}"
    return jsonify({'codigo': ref['codigo'], 'usos': ref['usos'], 'creditos': ref['creditos'], 'link': link})

# ═══════════════════════════════════════════════════════════════════════════════
# B2B ESCOLA
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/escola')
@login_required
def escola_painel():
    conn = get_db()
    # Verifica se o usuário é gestor de alguma escola
    if current_user.escola_id and current_user.papel == 'gestor':
        membros = conn.execute(
            """SELECT u.nome, u.email, em.papel, em.ativo,
                      (SELECT COUNT(*) FROM historico h WHERE h.usuario_id = u.id) as materiais
               FROM escola_membros em
               LEFT JOIN usuarios u ON u.id = em.usuario_id
               WHERE em.escola_id = ? ORDER BY em.criado_em DESC""",
            (current_user.escola_id,)
        ).fetchall()
        stats = conn.execute(
            """SELECT COUNT(DISTINCT em.usuario_id) as professores,
                      COUNT(h.id) as materiais_total
               FROM escola_membros em
               LEFT JOIN historico h ON h.usuario_id = em.usuario_id
               WHERE em.escola_id = ?""",
            (current_user.escola_id,)
        ).fetchone()
        conn.close()
        return render_template('escola_painel.html',
            membros=[dict(m) for m in membros],
            stats=dict(stats) if stats else {},
            escola_nome=current_user.escola_nome
        )
    conn.close()
    return render_template('escola_sem_acesso.html')

@app.route('/api/escola/convidar', methods=['POST'])
@login_required
def api_escola_convidar():
    if not (current_user.escola_id and current_user.papel == 'gestor'):
        return jsonify({'erro': 'Sem permissão'}), 403
    d = request.get_json(force=True)
    email = d.get('email', '').strip().lower()
    if not email:
        return jsonify({'erro': 'Email obrigatório'}), 400
    token = secrets.token_urlsafe(24)
    conn = get_db()
    conn.execute(
        "INSERT INTO escola_convites (escola_id, email, token, usado, criado_em) VALUES (?,?,?,0,?)",
        (current_user.escola_id, email, token, datetime.now().strftime('%d/%m/%Y'))
    )
    conn.commit()
    conn.close()
    link = f"{SITE_URL}/cadastro?convite={token}"
    enviar_email(email, f"Convite para {current_user.escola_nome} no ProfessorIA",
        f"""<p>Você foi convidado para fazer parte da escola <strong>{current_user.escola_nome}</strong> no ProfessorIA.</p>
        <p><a href="{link}" style="background:#4338ca;color:#fff;padding:12px 24px;border-radius:8px;text-decoration:none;font-weight:600;">Aceitar convite</a></p>
        <p style="color:#666;font-size:.85rem;">Ou acesse: {link}</p>""")
    return jsonify({'ok': True, 'link': link})

@app.route('/api/escola/relatorio')
@login_required
def api_escola_relatorio():
    if not (current_user.escola_id and current_user.papel == 'gestor'):
        return jsonify({'erro': 'Sem permissão'}), 403
    conn = get_db()
    rows = conn.execute(
        """SELECT u.nome, u.email,
                  COUNT(h.id) as materiais,
                  COUNT(DISTINCT h.disciplina) as disciplinas,
                  MAX(h.data) as ultimo_uso
           FROM escola_membros em
           JOIN usuarios u ON u.id = em.usuario_id
           LEFT JOIN historico h ON h.usuario_id = u.id
           WHERE em.escola_id = ?
           GROUP BY u.id, u.nome, u.email
           ORDER BY materiais DESC""",
        (current_user.escola_id,)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

# ═══════════════════════════════════════════════════════════════════════════════
# GERAÇÃO DE IMAGENS EDUCACIONAIS
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/api/gerar-imagem', methods=['POST'])
@login_required
def api_gerar_imagem():
    """Gera uma imagem educacional usando a API de imagens da Anthropic/DALL-E."""
    d = request.get_json(force=True)
    descricao = d.get('descricao', '').strip()
    if not descricao:
        return jsonify({'erro': 'Descrição obrigatória'}), 400
    # Usa a API de imagens do OpenAI (DALL-E 3) se disponível, senão retorna placeholder
    openai_key = os.environ.get('OPENAI_API_KEY', '')
    if not openai_key:
        return jsonify({
            'erro': 'API de imagens não configurada. Adicione OPENAI_API_KEY nas variáveis de ambiente.',
            'placeholder': True
        }), 503
    try:
        import requests as req
        prompt_educacional = (
            f"Educational illustration for Brazilian teachers, clean and professional style, "
            f"suitable for classroom use: {descricao}. "
            f"Flat design, colorful but not distracting, white background, high quality."
        )
        resp = req.post(
            "https://api.openai.com/v1/images/generations",
            headers={"Authorization": f"Bearer {openai_key}", "Content-Type": "application/json"},
            json={"model": "dall-e-3", "prompt": prompt_educacional, "n": 1, "size": "1024x1024", "quality": "standard"},
            timeout=60
        )
        data = resp.json()
        if resp.status_code == 200:
            url = data['data'][0]['url']
            return jsonify({'ok': True, 'url': url})
        else:
            return jsonify({'erro': data.get('error', {}).get('message', 'Erro ao gerar imagem')}), 400
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

# ═══════════════════════════════════════════════════════════════════════════════
# ONBOARDING MELHORADO — 3 passos
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/api/onboarding/completar', methods=['POST'])
@login_required
def api_onboarding_completar():
    d = request.get_json(force=True)
    disciplina = d.get('disciplina', '')[:100]
    serie = d.get('serie', '')[:50]
    template = d.get('template', '')[:5000]
    conn = get_db()
    conn.execute(
        """UPDATE usuarios SET onboarding_done = 1, escola_template = ? WHERE id = ?""",
        (template, current_user.id)
    )
    conn.commit()
    conn.close()
    from urllib.parse import urlencode
    return jsonify({'ok': True, 'redirect': '/chat?' + urlencode({'disciplina': disciplina, 'serie': serie})})


# ═══════════════════════════════════════════════════════════════════════════════
# GERAÇÃO GRÁTIS (lead capture + 1 plano completo)
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/gratis')
def gratis():
    return render_template('gratis.html')

@app.route('/api/gerar-gratis', methods=['POST'])
@limiter.limit('5 per hour')
def api_gerar_gratis():
    """Captura lead (nome + email + whatsapp) e gera 1 plano de aula completo."""
    data      = request.get_json(force=True) or {}
    nome      = str(data.get('nome', '')).strip()[:100]
    email     = str(data.get('email', '')).strip().lower()[:254]
    whatsapp  = str(data.get('whatsapp', '')).strip()[:20]
    tema      = str(data.get('tema', '')).strip()[:300]
    serie     = str(data.get('serie', '')).strip()[:50]
    disciplina = str(data.get('disciplina', 'Não informada')).strip()[:100]

    if not nome or not (email or whatsapp) or not tema:
        return jsonify({'erro': 'Nome, contato (email ou WhatsApp) e tema são obrigatórios'}), 400

    # Salva o lead
    try:
        conn = get_db()
        conn.execute(
            'INSERT INTO lista_vip (nome, email, whatsapp, criado_em) VALUES (%s, %s, %s, %s)',
            (nome, email or '', whatsapp, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        )
        conn.commit()
        conn.close()
        logger.info('Novo lead grátis: %s <%s> %s', nome, email, whatsapp)
        _capi_event('Lead', user_data={'email': email, 'phone': whatsapp, 'name': nome},
                    custom_data={'content_name': 'geracao_gratis'})
    except Exception as e:
        logger.warning('Erro ao salvar lead grátis: %s', e)

    # Gera o plano completo via cadeia de fallback
    prompt = (
        f"Gere um plano de aula completo e detalhado em português brasileiro para:\n"
        f"Tema: {tema}\n"
        f"Série/Ano: {serie or 'não informado'}\n"
        f"Disciplina: {disciplina}\n\n"
        "O plano deve conter:\n"
        "1. Título da aula\n"
        "2. Habilidades BNCC (códigos reais)\n"
        "3. Objetivos de aprendizagem\n"
        "4. Desenvolvimento da aula (passo a passo)\n"
        "5. Recursos didáticos\n"
        "6. Avaliação\n\n"
        "Seja específico, prático e alinhado à BNCC. Responda em português brasileiro."
    )
    try:
        resultado = _llm_cadeia_simples(prompt, sistema=SYSTEM_PROMPT, max_tokens=2000)
        return jsonify({'ok': True, 'plano': resultado, 'nome': nome})
    except Exception as e:
        logger.error('api_gerar_gratis — erro IA: %s', e)
        return jsonify({'erro': 'Serviço temporariamente indisponível. Tente novamente em instantes.'}), 503


# ═══════════════════════════════════════════════════════════════════════════════
# LISTA VIP
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/lista-vip', methods=['POST'])
def lista_vip():
    data     = request.get_json(silent=True) or {}
    nome     = data.get('nome', '').strip()[:200]
    email    = data.get('email', '').strip().lower()[:254]
    whatsapp = data.get('whatsapp', '').strip()[:20]
    if not nome or not email:
        return jsonify({'erro': 'Nome e e-mail são obrigatórios'}), 400
    try:
        conn = get_db()
        conn.execute(
            'INSERT INTO lista_vip (nome, email, whatsapp, criado_em) VALUES (%s, %s, %s, %s)',
            (nome, email, whatsapp, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        )
        conn.commit()
        conn.close()
        logger.info('Novo lead VIP: %s <%s>', nome, email)
        _capi_event('Lead', user_data={'email': email, 'phone': whatsapp, 'name': nome},
                    custom_data={'content_name': 'lista_vip'})
        return jsonify({'ok': True})
    except Exception as e:
        if 'unique' in str(e).lower():
            return jsonify({'erro': 'Este e-mail já está na Lista VIP!'}), 409
        logger.error('Erro ao salvar lead VIP: %s', e)
        return jsonify({'erro': 'Erro interno. Tente novamente.'}), 500


@app.route('/admin/leads')
def admin_leads():
    if request.args.get('senha', '') != LEADS_PASS:
        return Response('Acesso negado. Use ?senha=SUA_SENHA', status=401,
                        mimetype='text/plain; charset=utf-8')
    conn   = get_db()
    leads  = conn.execute(
        'SELECT id, nome, email, whatsapp, criado_em FROM lista_vip ORDER BY id DESC'
    ).fetchall()
    conn.close()

    if request.args.get('exportar') == 'csv':
        buf = io.StringIO()
        buf.write('ID,Nome,Email,WhatsApp,Cadastrado em\n')
        for l in leads:
            buf.write(f"{l['id']},{l['nome']},{l['email']},{l['whatsapp'] or ''},{l['criado_em']}\n")
        return Response(
            buf.getvalue(),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=leads_vip.csv'}
        )

    senha = request.args.get('senha', '')
    rows  = ''.join(
        f"<tr><td>{l['id']}</td><td>{l['nome']}</td><td>{l['email']}</td>"
        f"<td>{l['whatsapp'] or '—'}</td><td>{l['criado_em']}</td></tr>"
        for l in leads
    )
    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head><meta charset="utf-8"><title>Leads VIP — ProfessorIA</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:system-ui,sans-serif;background:#f1f5f9;padding:32px;color:#1e293b}}
h1{{color:#1e40af;margin-bottom:4px}}
.sub{{color:#64748b;font-size:.9rem;margin-bottom:20px}}
.actions{{display:flex;gap:10px;margin-bottom:20px}}
.btn{{padding:9px 20px;border-radius:8px;font-weight:600;font-size:.875rem;text-decoration:none;cursor:pointer;border:none}}
.btn-primary{{background:#1e40af;color:#fff}}
.btn-outline{{background:#fff;color:#1e40af;border:1.5px solid #1e40af}}
table{{width:100%;border-collapse:collapse;background:#fff;border-radius:12px;overflow:hidden;box-shadow:0 1px 6px rgba(0,0,0,.08)}}
th{{background:#1e40af;color:#fff;padding:12px 16px;text-align:left;font-size:.8rem;letter-spacing:.5px;text-transform:uppercase}}
td{{padding:11px 16px;border-bottom:1px solid #f1f5f9;font-size:.875rem}}
tr:last-child td{{border-bottom:none}}
tr:hover td{{background:#f8fafc}}
</style></head>
<body>
<h1>Lista VIP — ProfessorIA</h1>
<p class="sub">{len(leads)} lead(s) cadastrado(s)</p>
<div class="actions">
  <a class="btn btn-primary" href="/admin/leads?senha={senha}&exportar=csv">Exportar CSV</a>
  <a class="btn btn-outline" href="/admin/leads?senha={senha}">Atualizar</a>
</div>
<table>
  <thead><tr><th>#</th><th>Nome</th><th>E-mail</th><th>WhatsApp</th><th>Cadastrado em</th></tr></thead>
  <tbody>{rows if rows else '<tr><td colspan="5" style="text-align:center;color:#94a3b8;padding:32px">Nenhum lead ainda.</td></tr>'}</tbody>
</table>
</body></html>"""
    return html


# ─── Geração de Imagens — DALL-E 3 ──────────────────────────────────────────

@app.route('/api/generate/image', methods=['POST'])
@login_required
@limiter.limit('5 per minute')
def api_generate_image():
    """Gera uma ilustração via DALL-E 3.

    Entrada: { prompt, size? }
      - prompt: descrição da imagem (obrigatório)
      - size:   '1024x1024' | '1792x1024' | '1024x1792'  (padrão: 1024x1024)

    Saída: { url } com URL temporária da imagem (válida por 1h)
    """
    if not client_openai:
        return jsonify({'erro': 'OPENAI_API_KEY não configurada. Adicione a chave no painel do Render.'}), 503

    data   = request.get_json(force=True) or {}
    prompt = str(data.get('prompt', '')).strip()[:4000]
    size   = data.get('size', '1024x1024')

    if not prompt:
        return jsonify({'erro': 'Campo obrigatório: prompt'}), 400
    if size not in ('1024x1024', '1792x1024', '1024x1792'):
        size = '1024x1024'

    # Injeta estilo pedagógico automaticamente
    prompt_final = f"{prompt}. {IMAGE_STYLE_MODIFIER}"

    try:
        resp = client_openai.images.generate(
            model='dall-e-3',
            prompt=prompt_final,
            size=size,
            quality='standard',
            n=1
        )
        url = resp.data[0].url
        logger.info('DALL-E 3 gerou imagem para usuario %s', current_user.id)
        return jsonify({'url': url, 'prompt_revisado': resp.data[0].revised_prompt or prompt})
    except Exception as e:
        err = str(e)
        logger.error('DALL-E 3 erro: %s', err[:300])
        if 'billing' in err.lower() or 'insufficient' in err.lower():
            return jsonify({'erro': 'Créditos OpenAI esgotados. Adicione créditos em platform.openai.com'}), 402
        if 'content_policy' in err.lower() or 'safety' in err.lower():
            return jsonify({'erro': 'O prompt foi bloqueado pela política de conteúdo. Tente uma descrição diferente.'}), 422
        return jsonify({'erro': f'Erro ao gerar imagem: {err[:200]}'}), 500


@app.route('/api/generate/mapa-mental', methods=['POST'])
@login_required
@limiter.limit('5 per minute')
def api_generate_mapa_mental():
    """Gera uma ilustração visual de mapa mental via DALL-E 3.
    Entrada: { tema }
    Saída:   { url }
    """
    if not client_openai:
        return jsonify({'erro': 'OPENAI_API_KEY não configurada.'}), 503

    data = request.get_json(force=True) or {}
    tema = str(data.get('tema', '')).strip()[:300]
    if not tema:
        return jsonify({'erro': 'Campo obrigatório: tema'}), 400

    # Garante que o tema está em PT-BR no prompt enviado ao DALL-E
    # (evita que a IA alucie títulos em inglês na imagem gerada)
    prompt_final = (
        f"Mapa mental educacional em PORTUGUÊS DO BRASIL sobre '{tema}'. "
        f"O título central e todos os rótulos dos ramos DEVEM estar em português, "
        f"nunca em inglês. Estilo aquarela digital moderna, "
        "fundo 100% branco sólido, composição radial com tema central e ramos conectados "
        "a subtópicos ilustrados com ícones acadêmicos. "
        "Design premium, cores pastéis acadêmicas, sem texto longo dentro da imagem. "
        "No canto inferior direito: logotipo minimalista 'ProfessorIA™' em azul acadêmico. "
        + IMAGE_STYLE_MODIFIER
    )

    try:
        resp = client_openai.images.generate(
            model='dall-e-3',
            prompt=prompt_final[:4000],
            size='1792x1024',
            quality='standard',
            n=1
        )
        url = resp.data[0].url
        logger.info('Mapa mental visual gerado para usuario %s: %s', current_user.id, tema[:50])
        return jsonify({'url': url})
    except Exception as e:
        err = str(e)
        logger.error('DALL-E mapa-mental erro: %s', err[:300])
        return jsonify({'erro': f'Erro ao gerar mapa mental: {err[:200]}'}), 500


@app.route('/api/prova/docx', methods=['POST'])
@login_required
def api_prova_docx():
    """Gera DOCX de prova a partir dos dados estruturados.
    Entrada: { prova_dados: { titulo_prova, objetivos_bncc, questoes_multipla_escolha, questoes_dissertativas, gabarito_geral } }
    Saída:   arquivo .docx
    """
    data = request.get_json(force=True) or {}
    prova = data.get('prova_dados') or {}

    titulo     = str(prova.get('titulo_prova', 'Prova ProfessorIA'))[:200]
    objetivos  = prova.get('objetivos_bncc', [])
    mc         = prova.get('questoes_multipla_escolha', [])
    disc       = prova.get('questoes_dissertativas', [])
    gabarito   = str(prova.get('gabarito_geral', ''))

    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    def heading(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
        run.font.color.rgb = RGBColor(0x1E, 0x5A, 0x63)
        return p

    def body(text):
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(11) if p.runs else None
        return p

    heading(titulo)
    doc.add_paragraph(f"Nome: ___________________________   Turma: _______   Data: ___/___/______")
    doc.add_paragraph()

    if objetivos:
        heading('Habilidades BNCC', 2)
        for obj in objetivos:
            body(f"• {obj}")
        doc.add_paragraph()

    if mc:
        heading('Questões de Múltipla Escolha', 2)
        letras = ['A', 'B', 'C', 'D', 'E']
        for i, q in enumerate(mc, 1):
            body(f"{i}. {q.get('pergunta', '')}")
            for j, alt in enumerate(q.get('alternativas', [])):
                body(f"   {letras[j] if j < len(letras) else j+1}) {alt}")
            doc.add_paragraph()

    if disc:
        heading('Questões Dissertativas', 2)
        offset = len(mc)
        for i, q in enumerate(disc, 1):
            body(f"{offset + i}. {q.get('pergunta', '')}")
            for _ in range(4):
                body("_______________________________________________")
            doc.add_paragraph()

    if gabarito:
        heading('Gabarito (uso do professor)', 2)
        body(gabarito)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"Prova_{titulo[:40].replace(' ','_')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ─── Download PDF do Plano de Aula ───────────────────────────────────────────

@app.route('/api/plano-de-aula/pdf', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_plano_pdf():
    """Converte o JSON estruturado do plano em PDF e retorna para download.

    Body (JSON):
        plano_de_aula: objeto completo retornado pelo /api/gerar-plano
                       (aceita com ou sem o wrapper { "plano_de_aula": {...} })

    O display_name e school_name são injetados automaticamente do perfil do usuário.
    """
    data = request.get_json(force=True) or {}

    # Schema plano: JSON plano direto (flat)
    plano_json = data
    if not plano_json.get('tema'):
        return jsonify({'erro': 'JSON do plano de aula não encontrado no body'}), 400

    display_name = current_user.professor_nome or ''
    # Aceita escola temporária enviada pelo front-end (sem sobrescrever o perfil global)
    escola_override = plano_json.pop('_escola_override', None)
    school_name = escola_override or current_user.escola_nome or ''

    try:
        pdf_bytes = gerar_plano_pdf(plano_json, display_name=display_name, school_name=school_name)
    except Exception as e:
        logger.error('Erro ao gerar PDF do plano: %s', e)
        return jsonify({'erro': f'Falha ao gerar PDF: {str(e)[:200]}'}), 500

    tema = plano_json.get('tema', 'plano')
    nome_arquivo = f"PlanoDeAula_{tema[:40].replace(' ', '_')}.pdf"

    return send_file(
        io.BytesIO(pdf_bytes),
        as_attachment=True,
        download_name=nome_arquivo,
        mimetype='application/pdf'
    )


# ─── Gerador de Prova Estruturada ─────────────────────────────────────────────

SYSTEM_PROMPT_PROVA = (
    "Você é o ProfessorIA, um elaborador de exames de elite especializado no currículo "
    "educacional brasileiro. Sua missão é criar provas rigorosas, claras e pedagogicamente sólidas. "
    "REGRAS ABSOLUTAS: "
    "1) O retorno deve ser EXCLUSIVAMENTE um objeto JSON válido. "
    "2) É ESTRITAMENTE PROIBIDO usar tabelas Markdown, blocos de texto explicativo ou qualquer "
    "formatação visual. Sua resposta DEVE começar com { e terminar com }, contendo EXCLUSIVAMENTE "
    "o objeto JSON puro. Nenhuma palavra a mais, nenhuma tabela. "
    "3) Questões de múltipla escolha devem ter exatamente 4 alternativas (A, B, C, D). "
    "4) Os gabaritos devem ser precisos e pedagogicamente justificáveis. "
    "5) As questões discursivas devem ter gabarito esperado claro e objetivo."
)

PROVA_TOOL = {
    "name": "salvar_prova",
    "description": "Salva a prova estruturada gerada pelo elaborador de exames.",
    "input_schema": {
        "type": "object",
        "properties": {
            "prova": {
                "type": "object",
                "properties": {
                    "tema": {"type": "string"},
                    "questoes_verdadeiro_falso": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "afirmacao": {"type": "string"},
                                "resposta":  {"type": "string", "enum": ["V", "F"]}
                            },
                            "required": ["afirmacao", "resposta"]
                        }
                    },
                    "questoes_multipla_escolha": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "pergunta":      {"type": "string"},
                                "alternativas":  {
                                    "type": "array",
                                    "items": {"type": "string"},
                                    "minItems": 4,
                                    "maxItems": 4
                                },
                                "correta": {"type": "string", "enum": ["A", "B", "C", "D"]}
                            },
                            "required": ["pergunta", "alternativas", "correta"]
                        }
                    },
                    "questoes_discursivas": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "pergunta":          {"type": "string"},
                                "gabarito_esperado": {"type": "string"}
                            },
                            "required": ["pergunta", "gabarito_esperado"]
                        }
                    }
                },
                "required": [
                    "tema",
                    "questoes_verdadeiro_falso",
                    "questoes_multipla_escolha",
                    "questoes_discursivas"
                ]
            }
        },
        "required": ["prova"]
    }
}

# Schema OpenAI para prova (strict=True exige additionalProperties: false; sem minItems/maxItems)
_OAI_PROVA_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "prova": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "tema": {"type": "string"},
                "questoes_verdadeiro_falso": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "afirmacao": {"type": "string"},
                            "resposta":  {"type": "string", "enum": ["V", "F"]}
                        },
                        "required": ["afirmacao", "resposta"]
                    }
                },
                "questoes_multipla_escolha": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "pergunta":     {"type": "string"},
                            "alternativas": {"type": "array", "items": {"type": "string"}},
                            "correta":      {"type": "string", "enum": ["A", "B", "C", "D"]}
                        },
                        "required": ["pergunta", "alternativas", "correta"]
                    }
                },
                "questoes_discursivas": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "pergunta":          {"type": "string"},
                            "gabarito_esperado": {"type": "string"}
                        },
                        "required": ["pergunta", "gabarito_esperado"]
                    }
                }
            },
            "required": ["tema", "questoes_verdadeiro_falso", "questoes_multipla_escolha", "questoes_discursivas"]
        }
    },
    "required": ["prova"]
}


@app.route('/api/generate/prova', methods=['POST'])
@login_required
@limiter.limit('10 per minute')
def api_generate_prova():
    """Gera uma prova estruturada via JSON Schema.

    Entrada: { tema, ano, disciplina, num_vf?, num_mc?, num_disc? }
    Saída:   { prova: { tema, questoes_verdadeiro_falso, questoes_multipla_escolha,
                        questoes_discursivas } }
    """
    if not current_user.assinatura_ativa and not current_user.is_admin:
        geracoes = get_geracoes_mes(current_user.id)
        if geracoes >= LIMITE_GRATIS:
            return jsonify({
                'erro': 'limite_atingido',
                'geracoes': geracoes,
                'cta': '/planos',
                'mensagem': 'Você atingiu o limite do plano grátis. Faça upgrade para continuar gerando materiais!'
            }), 403

    data       = request.get_json(force=True) or {}
    tema       = str(data.get('tema', '')).strip()[:300]
    ano        = str(data.get('ano', '')).strip()[:50]
    disciplina = str(data.get('disciplina', '')).strip()[:100]

    if not tema or not ano or not disciplina:
        return jsonify({'erro': 'Campos obrigatórios: tema, ano, disciplina'}), 400

    num_vf   = max(1, min(int(data.get('num_vf',   5)), 10))
    num_mc   = max(1, min(int(data.get('num_mc',   5)), 10))
    num_disc = max(1, min(int(data.get('num_disc', 3)), 5))

    user_prompt = (
        f"Elabore uma prova completa para:\n"
        f"- Tema: {tema}\n"
        f"- Ano/Série: {ano}\n"
        f"- Disciplina: {disciplina}\n\n"
        f"Quantidade de questões:\n"
        f"- Verdadeiro ou Falso: {num_vf}\n"
        f"- Múltipla Escolha (4 alternativas A/B/C/D): {num_mc}\n"
        f"- Discursivas (com gabarito esperado): {num_disc}\n\n"
        "Garanta que as questões estejam alinhadas à BNCC e sejam adequadas ao nível escolar."
    )

    prova_json   = None
    erro_motores = []

    # ── Claude (tool_use — schema garantido) ──────────────────────────────────
    try:
        if os.environ.get('ANTHROPIC_API_KEY'):
            resp = client.messages.create(
                model='claude-sonnet-4-6',
                max_tokens=4000,
                system=SYSTEM_PROMPT_PROVA,
                tools=[PROVA_TOOL],
                tool_choice={"type": "tool", "name": "salvar_prova"},
                messages=[{"role": "user", "content": user_prompt}]
            )
            for block in resp.content:
                if block.type == 'tool_use' and block.name == 'salvar_prova':
                    prova_json = block.input
                    break
    except Exception as e:
        erro_motores.append(f'Claude: {e}')
        logger.warning('api_generate_prova — Claude falhou: %s', e)

    # ── Gemini (response_schema) ──────────────────────────────────────────────
    if prova_json is None and _gemini_disponivel():
        try:
            import google.generativeai as genai

            gemini_schema = {
                'type': 'object',
                'properties': {
                    'prova': {
                        'type': 'object',
                        'properties': {
                            'tema': {'type': 'string'},
                            'questoes_verdadeiro_falso': {
                                'type': 'array',
                                'items': {
                                    'type': 'object',
                                    'properties': {
                                        'afirmacao': {'type': 'string'},
                                        'resposta':  {'type': 'string'},
                                    }
                                }
                            },
                            'questoes_multipla_escolha': {
                                'type': 'array',
                                'items': {
                                    'type': 'object',
                                    'properties': {
                                        'pergunta':     {'type': 'string'},
                                        'alternativas': {'type': 'array', 'items': {'type': 'string'}},
                                        'correta':      {'type': 'string'},
                                    }
                                }
                            },
                            'questoes_discursivas': {
                                'type': 'array',
                                'items': {
                                    'type': 'object',
                                    'properties': {
                                        'pergunta':          {'type': 'string'},
                                        'gabarito_esperado': {'type': 'string'},
                                    }
                                }
                            },
                        }
                    }
                }
            }
            gm = genai.GenerativeModel(
                model_name='gemini-2.0-flash',
                system_instruction=SYSTEM_PROMPT_PROVA,
                generation_config=genai.GenerationConfig(
                    response_mime_type='application/json',
                    response_schema=gemini_schema
                )
            )
            resp_g = gm.generate_content(user_prompt)
            prova_json = json.loads(resp_g.text)
        except Exception as e:
            erro_motores.append(f'Gemini: {e}')
            logger.warning('api_generate_prova — Gemini falhou: %s', e)

    # ── OpenAI (json_schema) ──────────────────────────────────────────────────
    if prova_json is None:
        try:
            oai_key = os.environ.get('OPENAI_API_KEY')
            if oai_key:
                import openai as _oai
                oai_client = _oai.OpenAI(api_key=oai_key)
                resp_o = oai_client.chat.completions.create(
                    model='gpt-4o-mini',
                    messages=[
                        {'role': 'system', 'content': SYSTEM_PROMPT_PROVA},
                        {'role': 'user',   'content': user_prompt}
                    ],
                    response_format={
                        'type': 'json_schema',
                        'json_schema': {
                            'name': 'prova',
                            'strict': True,
                            'schema': _OAI_PROVA_SCHEMA
                        }
                    },
                    max_tokens=4000
                )
                prova_json = json.loads(resp_o.choices[0].message.content)
        except Exception as e:
            erro_motores.append(f'OpenAI: {e}')
            logger.warning('api_generate_prova — OpenAI falhou: %s', e)

    if prova_json is None:
        return jsonify({
            'erro': 'Todos os motores de IA falharam. Verifique as chaves de API.',
            'detalhes': erro_motores
        }), 503

    # Contabiliza geração
    conn = get_db()
    conn.execute(
        "INSERT INTO chat_messages (usuario_id, role, content, criado_em) VALUES (?, ?, ?, ?)",
        (current_user.id, 'assistant',
         f'[prova estruturada] {tema} — {disciplina} {ano}',
         datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    )
    conn.commit()
    conn.close()

    return jsonify(prova_json)


# ESTE BLOCO ABAIXO DEVE SER O FINAL ABSOLUTO DO ARQUIVO
if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5001))
    app.run(debug=False, host='0.0.0.0', port=port)