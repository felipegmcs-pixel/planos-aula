"""
gerar_template_plano.py
Gera o arquivo static/templates/plano_de_aula.docx com placeholders docxtpl.
Execute uma vez para criar/atualizar o template.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL       = RGBColor(0x1E, 0x3A, 0x5F)
AZUL_MED   = RGBColor(0x1E, 0x40, 0xAF)
BRANCO     = RGBColor(0xFF, 0xFF, 0xFF)
CINZA_TEX  = RGBColor(0x37, 0x41, 0x51)
CINZA_BG   = "EEF2FF"
AZUL_BG    = "1E3A5F"
AZUL_MED_H = "1E40AF"


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _p(para, text, bold=False, italic=False, size=10, color=None, align=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align
    return run


def criar_template():
    doc = Document()
    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width  = Cm(21.0)
        sec.top_margin = sec.bottom_margin = Cm(1.8)
        sec.left_margin = sec.right_margin  = Cm(2.0)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    def add_heading_band(text, bg_hex=AZUL_BG):
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'
        c = t.cell(0, 0)
        _set_cell_bg(c, bg_hex)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BRANCO
        return t

    # ── Cabeçalho ──────────────────────────────────────────────────────────────
    h = doc.add_table(rows=1, cols=2)
    h.style = 'Table Grid'
    h.columns[0].width = Cm(14)
    h.columns[1].width = Cm(3)
    _set_cell_bg(h.cell(0, 0), AZUL_BG)
    _set_cell_bg(h.cell(0, 1), AZUL_BG)

    lp = h.cell(0, 0).paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after  = Pt(2)
    r = lp.add_run('{{ escola | upper if escola else "ESCOLA ESTADUAL" }}')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BRANCO

    lp2 = h.cell(0, 0).add_paragraph()
    lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp2.paragraph_format.space_before = Pt(0)
    lp2.paragraph_format.space_after  = Pt(2)
    r2 = lp2.add_run('SECRETARIA DE ESTADO DA EDUCAÇÃO')
    r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)

    rp = h.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.paragraph_format.space_before = Pt(8)
    rp.add_run('ProfessorIA').font.color.rgb = RGBColor(0x81, 0x8C, 0xF8)

    doc.add_paragraph()

    # ── Título ─────────────────────────────────────────────────────────────────
    add_heading_band('PLANO DE AULA')
    doc.add_paragraph()

    # ── Ficha de identificação ─────────────────────────────────────────────────
    info = doc.add_table(rows=4, cols=2)
    info.style = 'Table Grid'
    labels = [
        ('Tema / Conteúdo',    '{{ tema_central }}'),
        ('Disciplina',         '{{ disciplina }}'),
        ('Ano / Série',        '{{ ano_escolar }}'),
        ('Tempo Estimado',     '{{ tempo_estimado }}'),
    ]
    for i, (lbl, val) in enumerate(labels):
        lc = info.cell(i, 0).paragraphs[0]
        lc.paragraph_format.space_before = Pt(3)
        lc.paragraph_format.space_after  = Pt(3)
        r = lc.add_run(lbl)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = CINZA_TEX
        _set_cell_bg(info.cell(i, 0), CINZA_BG)

        vc = info.cell(i, 1).paragraphs[0]
        vc.paragraph_format.space_before = Pt(3)
        vc.paragraph_format.space_after  = Pt(3)
        vc.add_run(val).font.size = Pt(9)

    doc.add_paragraph()

    # ── Habilidades BNCC ───────────────────────────────────────────────────────
    add_heading_band('HABILIDADES BNCC', bg_hex=AZUL_MED_H)
    bncc_tbl = doc.add_table(rows=1, cols=1)
    bncc_tbl.style = 'Table Grid'
    bp = bncc_tbl.cell(0, 0).paragraphs[0]
    bp.paragraph_format.space_before = Pt(4)
    bp.paragraph_format.space_after  = Pt(4)
    bp.add_run('{% for h in habilidades_bncc %}{{ h.codigo }}: {{ h.descricao }}{% if not loop.last %}  |  {% endif %}{% endfor %}').font.size = Pt(9)

    doc.add_paragraph()

    # ── Desenvolvimento ────────────────────────────────────────────────────────
    add_heading_band('DESENVOLVIMENTO DA AULA', bg_hex=AZUL_MED_H)
    dev_tbl = doc.add_table(rows=1, cols=4)
    dev_tbl.style = 'Table Grid'
    for j, hdr_text in enumerate(['Etapa', 'Conteúdo', 'Estratégias Didáticas', 'Recursos']):
        cp = dev_tbl.cell(0, j).paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after  = Pt(3)
        rr = cp.add_run(hdr_text)
        rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = BRANCO
        _set_cell_bg(dev_tbl.cell(0, j), AZUL_MED_H)

    # Linha de dados — docxtpl loop
    dev_tbl.add_row()
    row_cells = dev_tbl.rows[1].cells
    tags = [
        '{% for et in desenvolvimento %}{{ et.etapa }}',
        '{{ et.conteudo }}',
        '{{ et.estrategias_didaticas }}',
        '{{ et.recursos_pedagogicos | join(", ") }}{% endfor %}',
    ]
    for j, tag in enumerate(tags):
        p = row_cells[j].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.add_run(tag).font.size = Pt(9)

    doc.add_paragraph()

    # ── Avaliação e Fechamento ─────────────────────────────────────────────────
    add_heading_band('AVALIAÇÃO E FECHAMENTO', bg_hex=AZUL_MED_H)
    av_tbl = doc.add_table(rows=1, cols=2)
    av_tbl.style = 'Table Grid'
    for j, (lbl, tag) in enumerate([
        ('Método de Avaliação', '{{ avaliacao_e_fechamento.metodo }}'),
        ('Critérios',           '{{ avaliacao_e_fechamento.criterios }}'),
    ]):
        lc = av_tbl.cell(0, j).paragraphs[0]
        lc.paragraph_format.space_before = Pt(3)
        lc.paragraph_format.space_after  = Pt(3)
        r1 = lc.add_run(lbl + '\n')
        r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = CINZA_TEX
        lc.add_run(tag).font.size = Pt(9)

    doc.add_paragraph()

    # ── Rodapé ─────────────────────────────────────────────────────────────────
    rod = doc.add_paragraph()
    rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rod.paragraph_format.space_before = Pt(6)
    r = rod.add_run('Gerado por ProfessorIA™  ·  Prof(a). {{ professor }}  ·  {{ data_geracao }}')
    r.font.size = Pt(7); r.italic = True
    r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Salvar
    out = os.path.join(os.path.dirname(__file__), 'static', 'templates', 'plano_de_aula.docx')
    doc.save(out)
    print(f'Template salvo em: {out}')
    return out


if __name__ == '__main__':
    criar_template()
