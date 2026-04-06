import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell borders.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top', 'start', 'bottom', 'end'):
        if side in kwargs:
            tag = 'w:{}'.format(side)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, value in kwargs[side].items():
                element.set(qn('w:{}'.format(key)), str(value))

def gerar_template():
    doc = Document()
    
    # Configurações de página (A4)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # --- CABEÇALHO PREMIUM ---
    table = doc.add_table(rows=1, cols=2)
    table.allow_autofit = False
    table.columns[0].width = Cm(12)
    table.columns[1].width = Cm(4)
    
    cell_info = table.cell(0, 0)
    p = cell_info.paragraphs[0]
    run = p.add_run("PROFESSORIA SaaS - Avaliação Acadêmica")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x5A, 0x63) # Turquesa Profundo
    
    p2 = cell_info.add_paragraph("Sistema Inteligente de Apoio ao Docente")
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.runs[0]
    run2.font.size = Pt(9)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Espaço para identificação do aluno
    p3 = cell_info.add_paragraph("NOME: _______________________________________________________")
    p3.paragraph_format.space_before = Pt(10)
    p4 = cell_info.add_paragraph("DATA: ____/____/________   TURMA: ____________   NOTA: ________")
    
    # --- TÍTULO DA PROVA ---
    doc.add_paragraph() # Espaçador
    title_p = doc.add_paragraph("{{ titulo_prova }}")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.runs[0]
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(0x0D, 0x23, 0x27) # Turquesa Escuro

    # --- OBJETIVOS BNCC ---
    doc.add_paragraph()
    obj_title = doc.add_paragraph("Objetivos de Aprendizagem (BNCC):")
    obj_title.runs[0].bold = True
    obj_title.runs[0].font.color.rgb = RGBColor(0xD4, 0xAF, 0x37) # Dourado Envelhecido
    
    doc.add_paragraph("{% for obj in objetivos_bncc %}")
    p_obj = doc.add_paragraph("• {{ obj }}")
    p_obj.paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph("{% endfor %}")

    # --- QUESTÕES DE MÚLTIPLA ESCOLHA ---
    doc.add_paragraph()
    mc_title = doc.add_paragraph("PARTE I: QUESTÕES DE MÚLTIPLA ESCOLHA")
    mc_title.runs[0].bold = True
    mc_title.runs[0].font.size = Pt(12)
    mc_title.runs[0].font.color.rgb = RGBColor(0x1E, 0x5A, 0x63)

    doc.add_paragraph("{% for q in questoes_multipla_escolha %}")
    p_qmc = doc.add_paragraph("Questão {{ loop.index }}: {{ q.pergunta }}")
    p_qmc.paragraph_format.space_before = Pt(12)
    p_qmc.runs[0].bold = True
    
    doc.add_paragraph("{% for alt in q.alternativas %}")
    # Usando letras A, B, C, D baseadas no índice do loop interno
    p_alt = doc.add_paragraph("({{ ['A','B','C','D'][loop.index0] }}) {{ alt }}")
    p_alt.paragraph_format.left_indent = Cm(1.0)
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endfor %}")

    # --- QUESTÕES DISSERTATIVAS ---
    doc.add_paragraph()
    di_title = doc.add_paragraph("PARTE II: QUESTÕES DISSERTATIVAS")
    di_title.runs[0].bold = True
    di_title.runs[0].font.size = Pt(12)
    di_title.runs[0].font.color.rgb = RGBColor(0x1E, 0x5A, 0x63)

    doc.add_paragraph("{% for q in questoes_dissertativas %}")
    p_qdi = doc.add_paragraph("Questão {{ loop.index + questoes_multipla_escolha|length }}: {{ q.pergunta }}")
    p_qdi.paragraph_format.space_before = Pt(12)
    p_qdi.runs[0].bold = True
    
    # Linhas para resposta
    for _ in range(4):
        doc.add_paragraph("______________________________________________________________________________")
    doc.add_paragraph("{% endfor %}")

    # --- GABARITO (PÁGINA SEPARADA) ---
    doc.add_page_break()
    gab_title = doc.add_paragraph("GABARITO E ORIENTAÇÕES PEDAGÓGICAS")
    gab_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gab_title.runs[0].bold = True
    gab_title.runs[0].font.size = Pt(16)
    gab_title.runs[0].font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)

    doc.add_paragraph()
    doc.add_paragraph("Resumo Geral: {{ gabarito_geral }}")
    
    doc.add_paragraph("\nRespostas Múltipla Escolha:")
    doc.add_paragraph("{% for q in questoes_multipla_escolha %}")
    doc.add_paragraph("Q{{ loop.index }}: {{ q.resposta_correta }}")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph("\nExpectativas de Resposta (Dissertativas):")
    doc.add_paragraph("{% for q in questoes_dissertativas %}")
    doc.add_paragraph("Q{{ loop.index + questoes_multipla_escolha|length }}: {{ q.expectativa_resposta }}")
    doc.add_paragraph("{% endfor %}")

    # Salvar
    output_path = "/home/ubuntu/planos-aula/static/templates/prova_template.docx"
    doc.save(output_path)
    print(f"Template criado em: {output_path}")

if __name__ == "__main__":
    gerar_template()
